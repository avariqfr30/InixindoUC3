import base64
import io
import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from docx_rendering import DocumentBuilder
from report_finalization import ReportFinalizer

logger = logging.getLogger(__name__)


class ReportDocumentAssembler:
    def __init__(self, finalizer=None, document_builder=DocumentBuilder):
        self.finalizer = finalizer or ReportFinalizer()
        self.document_builder = document_builder

    def assemble(self, generated_content, analysis_payload, theme_color):
        document = Document()
        self.document_builder.create_cover(document, theme_color)
        self.document_builder.add_table_of_contents(document)
        self.document_builder.process_content(
            document,
            generated_content,
            theme_color,
        )
        self.embed_dashboard_screenshots(document, analysis_payload, theme_color)
        return document

    def has_dashboard_screenshots(self, analysis_payload):
        payload = self.finalizer.normalize_analysis_payload(analysis_payload)
        screenshots = payload.get("dashboard_screenshots")
        return bool(screenshots and isinstance(screenshots, list) and len(screenshots) > 0)

    def embed_dashboard_screenshots(self, document, analysis_payload, theme_color):
        payload = self.finalizer.normalize_analysis_payload(analysis_payload)
        screenshots = payload.get("dashboard_screenshots")
        if not screenshots or not isinstance(screenshots, list):
            return

        valid_screenshots = [
            shot for shot in screenshots
            if isinstance(shot, dict) and shot.get("image_base64")
        ]
        if not valid_screenshots:
            return

        heading = document.add_heading("Dashboard Cashflow Snapshot", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*theme_color)

        intro = document.add_paragraph(
            "Visual berikut diambil langsung dari dashboard operasional pada saat laporan diminta. "
            "Setiap horizon menampilkan kondisi kas, runway, coverage ratio, prediksi saldo, "
            "distribusi delay pembayaran, dan daftar akun overdue utama."
        )
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in intro.runs:
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)

        for shot in valid_screenshots:
            try:
                image_bytes = base64.b64decode(shot["image_base64"])
                image_stream = io.BytesIO(image_bytes)
                horizon_label = shot.get("horizon_label") or shot.get("horizon_key") or "Dashboard"

                image_paragraph = document.add_paragraph()
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_paragraph.add_run().add_picture(image_stream, width=Inches(6.8))

                try:
                    caption = document.add_paragraph(
                        f"Dashboard snapshot - {horizon_label}",
                        style="Caption",
                    )
                except KeyError:
                    caption = document.add_paragraph(
                        f"Dashboard snapshot - {horizon_label}"
                    )
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption.runs:
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(100, 100, 100)

            except Exception as exc:
                logger.warning("Failed to embed dashboard screenshot for %s: %s", shot.get("horizon_key"), exc)
                continue

        document.add_page_break()
