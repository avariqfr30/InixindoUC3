import base64
import concurrent.futures
import io
import json
import logging
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from ollama import Client

from cashflow_analysis import FinancialAnalyzer
from config import (
    DEFAULT_COLOR,
    FINANCE_SYSTEM_PROMPT,
    LLM_MODEL,
    OLLAMA_HOST,
    PERSONAS,
    REPORT_MIN_COMPLETENESS_SCORE,
    REPORT_NUM_CTX,
    REPORT_NUM_PREDICT,
    REPORT_REPEAT_PENALTY,
    REPORT_SECTION_SEQUENCE,
    REPORT_TEMPERATURE,
    REPORT_TOP_P,
)
from docx_rendering import DocumentBuilder
from osint_research import Researcher

logger = logging.getLogger(__name__)

class ReportGenerator:
    SECTION_PASSES = (
        {
            "sections": REPORT_SECTION_SEQUENCE[:1],
            "include_visuals": False,
            "label": "executive_confidence",
        },
        {
            "sections": REPORT_SECTION_SEQUENCE[1:3],
            "include_visuals": True,
            "label": "diagnostic_evidence",
        },
        {
            "sections": REPORT_SECTION_SEQUENCE[3:],
            "include_visuals": False,
            "label": "actions_readiness",
        },
    )

    def __init__(self, kb_instance):
        self.ollama = Client(host=OLLAMA_HOST)
        self.kb = kb_instance
        self.io_pool = concurrent.futures.ThreadPoolExecutor(max_workers=4)

    @staticmethod
    def _normalize_analysis_payload(analysis_payload):
        if isinstance(analysis_payload, dict):
            return analysis_payload
        if not analysis_payload:
            return {}
        if isinstance(analysis_payload, str):
            try:
                parsed = json.loads(analysis_payload)
            except json.JSONDecodeError:
                return {}
            return parsed if isinstance(parsed, dict) else {}
        return {}

    @staticmethod
    def _build_dashboard_visual_markers(analysis_payload):
        payload = ReportGenerator._normalize_analysis_payload(analysis_payload)
        horizon_snapshot = payload.get("horizon_snapshot") or {}
        forecasts = horizon_snapshot.get("forecasts") if isinstance(horizon_snapshot, dict) else None
        if not isinstance(forecasts, dict):
            return []

        markers = []
        for horizon_key in ("short_term", "mid_term", "long_term"):
            forecast = forecasts.get(horizon_key)
            if not isinstance(forecast, dict):
                continue
            dashboard = forecast.get("dashboard_snapshot")
            if not isinstance(dashboard, dict):
                continue
            coverage_bars = (((dashboard.get("coverage_chart") or {}).get("bars")) or [])[:4]
            compact_payload = {
                "horizon_key": dashboard.get("horizon_key") or horizon_key,
                "horizon_label": dashboard.get("horizon_label"),
                "horizon_focus": dashboard.get("horizon_focus"),
                "status": dashboard.get("status"),
                "current_cash": dashboard.get("current_cash"),
                "runway_months": dashboard.get("runway_months"),
                "coverage_ratio": dashboard.get("coverage_ratio"),
                "average_delay_days": dashboard.get("average_delay_days"),
                "balance_projection": dashboard.get("balance_projection_30d") or [],
                "coverage_bars": coverage_bars,
            }
            markers.append(f"[[DASHBOARD:{json.dumps(compact_payload, ensure_ascii=False, separators=(',', ':'))}]]")
        return markers

    @staticmethod
    def _build_operational_snapshot_block(analysis_payload):
        payload = ReportGenerator._normalize_analysis_payload(analysis_payload)
        if not payload:
            return ""

        selected_period = payload.get("selected_period") or {}
        sync_status = payload.get("sync_status") or {}
        financial_sync = sync_status.get("financialData") or {}
        cash_out_sync = sync_status.get("cashOutSource") or {}

        lines = []
        period_label = selected_period.get("label")
        if period_label:
            lines.append(f"- Periode dashboard yang diekspor ke laporan: {period_label}.")

        cash_on_hand = payload.get("cash_on_hand")
        if cash_on_hand is not None:
            lines.append(f"- Cash in hand pada saat review: Rp{int(cash_on_hand):,}.")

        if financial_sync:
            freshness = financial_sync.get("sourceAgeMinutes")
            freshness_label = "belum tersedia"
            if freshness is not None:
                freshness_label = f"{float(freshness):.1f} menit"
            lines.append(
                "- Status sinkronisasi data finansial: "
                f"{financial_sync.get('syncStatus') or '-'}; usia data {freshness_label}; "
                f"record aktif {int(financial_sync.get('recordCount') or 0)}."
            )

        if cash_out_sync:
            if cash_out_sync.get("configured"):
                lines.append(
                    "- Sumber cash out memakai komitmen aktual dengan status "
                    f"{cash_out_sync.get('syncStatus') or '-'} dan "
                    f"{int(cash_out_sync.get('recordCount') or 0)} item aktif."
                )
            else:
                lines.append("- Sumber cash out masih memakai model operating cost bulanan karena feed kewajiban aktual belum dikonfigurasi.")

        horizons = ((payload.get("horizon_snapshot") or {}).get("forecasts")) or {}
        for key in ("short_term", "mid_term", "long_term"):
            horizon_payload = horizons.get(key) or {}
            forecast = horizon_payload.get("forecast") or {}
            health = horizon_payload.get("cashflow_health") or {}
            horizon = horizon_payload.get("time_horizon") or {}
            if not forecast:
                continue
            lines.append(
                "- "
                f"{horizon.get('label') or key}: cash masuk {FinancialAnalyzer._format_currency(int((forecast.get('cash_in') or {}).get('total_predicted_cash_in') or 0))}, "
                f"cash keluar {FinancialAnalyzer._format_currency(int((forecast.get('cash_out') or {}).get('total_cash_out') or 0))}, "
                f"ending cash {FinancialAnalyzer._format_currency(int(forecast.get('ending_cash') or 0))}, "
                f"runway {float(health.get('runway_months') or 0):.1f} bulan, "
                f"coverage {float(health.get('coverage_ratio') or 0):.2f}x, "
                f"fokus {horizon.get('focus') or '-'}."
            )

        return "\n".join(lines).strip()

    @staticmethod
    def _build_user_instruction(notes, active_sections):
        notes = (notes or "").strip()
        report_sections = "\n".join(f"- {section}" for section in active_sections)
        focus_block = notes if notes else "Tidak ada fokus tambahan dari pengguna."
        return (
            "Susun laporan internal yang detail, profesional, dan siap dipakai sebagai bahan diskusi rapat manajemen.\n"
            "Jangan menulis seperti jawaban AI generik; tulis seperti memo manajemen yang berbasis data.\n"
            "Kerjakan hanya section yang diminta pada pass ini dan hentikan output setelah section terakhir selesai.\n"
            "Pastikan setiap bagian di bawah terisi secara jelas:\n"
            f"{report_sections}\n\n"
            "Setiap rekomendasi harus menyebutkan fokus tindakan dan dampak yang diharapkan.\n"
            f"Fokus pengguna:\n{focus_block}"
        )

    @staticmethod
    def _build_section_scope(active_sections):
        section_list = "\n".join(f"- {section}" for section in active_sections)
        section_headings = "\n".join(f"   # {section}" for section in active_sections)
        section_scope = (
            "Generate only the sections listed below for this pass.\n"
            "Start with the first heading listed and stop after the last heading listed.\n"
            "Do not repeat prior sections and do not preview later sections.\n"
            f"{section_list}"
        )
        return section_scope, section_headings

    def _build_report_prompt(self, report_context, notes, analysis_context, macro_osint, active_sections, include_visuals):
        persona = PERSONAS.get("default", "Chief Financial Officer")
        section_scope, section_headings = self._build_section_scope(active_sections)
        structured_context = self._format_structured_context_block(analysis_context)
        return FINANCE_SYSTEM_PROMPT.format(
            persona=persona,
            financial_summary=report_context["financial_summary"],
            management_brief=report_context["management_brief"],
            internal_evidence=report_context["evidence"],
            industry_trends=macro_osint,
            user_focus=(notes or "Tidak ada fokus tambahan."),
            cashflow_context=(structured_context or "Tidak ada konteks forecast terstruktur tambahan."),
            readiness_signals=report_context["readiness_signals"],
            section_scope=section_scope,
            section_headings=section_headings,
            visual_prompt=report_context["visual_prompt"] if include_visuals else "",
        )

    @staticmethod
    def _score_report_completeness(raw_text):
        report_text = str(raw_text or "")
        if not report_text.strip():
            return {
                "score": 0.0,
                "passed": False,
                "components": {},
                "missing": ["Dokumen kosong."],
            }

        components = {}
        missing = []

        required_headings = [f"# {section}" for section in REPORT_SECTION_SEQUENCE]
        present_headings = [heading for heading in required_headings if heading in report_text]
        heading_score = (len(present_headings) / len(required_headings)) * 30
        components["top_level_sections"] = round(heading_score, 1)
        if len(present_headings) < len(required_headings):
            missing_sections = [heading.replace("# ", "") for heading in required_headings if heading not in present_headings]
            missing.append(f"Bagian utama belum lengkap: {', '.join(missing_sections)}.")

        required_subheadings = [
            "### Dampak Bisnis",
            "### Tingkat Keyakinan dan Caveat",
            "### Batasan Data dan Asumsi",
            "### Konteks OSINT Pendukung",
            "### Risiko dan Kontrol",
            "### Prasyarat Implementasi",
            "### Kesiapan Pelaksanaan",
        ]
        present_subheadings = [subheading for subheading in required_subheadings if subheading in report_text]
        subheading_score = (len(present_subheadings) / len(required_subheadings)) * 20
        components["required_subsections"] = round(subheading_score, 1)
        if len(present_subheadings) < len(required_subheadings):
            missing_subsections = [subheading.replace("### ", "") for subheading in required_subheadings if subheading not in present_subheadings]
            missing.append(f"Subbagian wajib belum lengkap: {', '.join(missing_subsections)}.")

        table_score = 0
        scenario_table_pattern = re.compile(
            r"\|\s*Skenario\s*\|\s*Estimasi Arus Kas Masuk\s*\|\s*Gap terhadap Total Invoice\s*\|",
            re.IGNORECASE,
        )
        if scenario_table_pattern.search(report_text):
            table_score += 8
        else:
            missing.append("Tabel skenario arus kas masuk belum ditemukan.")

        priority_table_pattern = re.compile(
            r"\|\s*Prioritas\s*\|\s*Fokus\s*\|\s*Penanggung Jawab\s*\|\s*Isu Utama\s*\|\s*Aksi 30 Hari\s*\|\s*Dampak yang Diharapkan\s*\|"
        )
        if priority_table_pattern.search(report_text):
            table_score += 12
        else:
            missing.append("Tabel prioritas 30 hari belum lengkap.")
        components["tables_and_owners"] = round(table_score, 1)

        enrichment_score = 0
        if "### Konteks OSINT Pendukung" in report_text:
            enrichment_score += 4
        if "[[CHART:" in report_text:
            enrichment_score += 3
        else:
            missing.append("Visual distribusi pembayaran belum masuk.")
        if "[[FLOW:" in report_text:
            enrichment_score += 3
        else:
            missing.append("Visual alur mitigasi belum masuk.")
        components["context_and_visuals"] = round(enrichment_score, 1)

        consistency_score = 10
        contradiction_patterns = [
            r"turun[^.\n]{0,120}memburuk",
            r"menurun[^.\n]{0,120}memburuk",
            r"naik[^.\n]{0,120}membaik",
            r"meningkat[^.\n]{0,120}membaik",
        ]
        if any(re.search(pattern, report_text, re.IGNORECASE) for pattern in contradiction_patterns):
            consistency_score = 0
            missing.append("Narasi tren risiko terdeteksi kontradiktif terhadap arah metrik.")
        components["numeric_consistency"] = round(consistency_score, 1)

        density_score = 0
        if len(report_text.strip()) >= 4500:
            density_score += 4
        elif len(report_text.strip()) >= 3200:
            density_score += 2
        else:
            missing.append("Narasi laporan masih terlalu tipis untuk bahan rapat internal.")

        if report_text.count("\n- ") >= 12:
            density_score += 3
        else:
            missing.append("Rincian bullet operasional masih kurang kaya.")

        if report_text.count("\n1.") >= 1 or report_text.count("\n1. ") >= 1:
            density_score += 1
        else:
            missing.append("Daftar tindakan bernomor belum kuat.")

        if "catatan" in report_text.lower() or "bukti" in report_text.lower():
            density_score += 2
        else:
            missing.append("Rujukan bukti internal belum cukup terlihat.")
        components["narrative_density"] = round(density_score, 1)

        total_score = round(sum(components.values()), 1)
        return {
            "score": total_score,
            "passed": total_score >= REPORT_MIN_COMPLETENESS_SCORE,
            "components": components,
            "missing": missing,
        }

    @staticmethod
    def _is_acceptable_report(raw_text):
        return ReportGenerator._score_report_completeness(raw_text)["passed"]

    @staticmethod
    def _extract_visual_markers(visual_prompt):
        chart_marker = ""
        flow_marker = ""
        for line in str(visual_prompt or "").splitlines():
            stripped_line = line.strip()
            if stripped_line.startswith("[[CHART:"):
                chart_marker = stripped_line
            elif stripped_line.startswith("[[FLOW:"):
                flow_marker = stripped_line
        return chart_marker, flow_marker

    @staticmethod
    def _split_top_level_sections(raw_text):
        matches = list(re.finditer(r"(?m)^# ([^\n]+?)\s*$", raw_text or ""))
        if not matches:
            return []

        sections = []
        for index, match in enumerate(matches):
            section_title = match.group(1).strip()
            section_end = matches[index + 1].start() if index + 1 < len(matches) else len(raw_text)
            section_body = raw_text[match.end():section_end].strip()
            sections.append({"title": section_title, "body": section_body})

        return sections

    @staticmethod
    def _join_top_level_sections(sections):
        blocks = []
        for section in sections:
            section_title = section["title"].strip()
            section_body = section["body"].strip()
            if section_body:
                blocks.append(f"# {section_title}\n{section_body}")
            else:
                blocks.append(f"# {section_title}")
        return "\n\n".join(blocks).strip()

    @staticmethod
    def _inject_subheading_block(section_body, subheading, content, before_subheading=None):
        if not content or not str(content).strip():
            return section_body

        subheading_marker = f"### {subheading}"
        if subheading_marker in section_body:
            return section_body

        new_block = f"{subheading_marker}\n{str(content).strip()}"
        if before_subheading:
            before_match = re.search(rf"(?m)^### {re.escape(before_subheading)}\s*$", section_body)
            if before_match:
                section_prefix = section_body[:before_match.start()].rstrip()
                section_suffix = section_body[before_match.start():].lstrip()
                return f"{section_prefix}\n\n{new_block}\n\n{section_suffix}".strip()

        return f"{section_body.rstrip()}\n\n{new_block}".strip()

    @staticmethod
    def _append_marker_block(section_body, marker):
        marker = str(marker or "").strip()
        if not marker or marker in section_body:
            return section_body
        return f"{section_body.rstrip()}\n\n{marker}".strip()

    @staticmethod
    def _format_structured_context_block(raw_text):
        lines = []
        for raw_line in str(raw_text or "").splitlines():
            cleaned_line = raw_line.strip()
            if not cleaned_line:
                if lines and lines[-1] != "":
                    lines.append("")
                continue
            if cleaned_line.startswith("===") and cleaned_line.endswith("==="):
                cleaned_line = cleaned_line.strip("= ").strip()
            if cleaned_line.endswith(":") and not cleaned_line.startswith("-"):
                lines.append(f"#### {cleaned_line[:-1].strip()}")
            else:
                lines.append(cleaned_line)
        return "\n".join(lines).strip()

    @classmethod
    def _sanitize_generated_report_text(cls, raw_text):
        sanitized = str(raw_text or "")
        sanitized = re.sub(r"(?<!\n)(===\s*[^\n=]+?\s*===)", r"\n\n\1", sanitized)
        sanitized = re.sub(r"\n?===\s*([^\n=]+?)\s*===\n?", lambda match: f"\n\n### {match.group(1).strip()}\n", sanitized)
        sanitized = re.sub(r"[ \t]+\n", "\n", sanitized)
        sanitized = re.sub(r"\n{3,}", "\n\n", sanitized)
        return sanitized.strip()

    def _finalize_report_content(self, raw_text, report_context, macro_osint, analysis_payload=None):
        raw_text = self._sanitize_generated_report_text(raw_text)
        sections = self._split_top_level_sections(raw_text)
        if not sections:
            return raw_text

        chart_marker, flow_marker = self._extract_visual_markers(report_context.get("visual_prompt", ""))
        dashboard_markers = self._build_dashboard_visual_markers(analysis_payload)
        operational_snapshot = self._build_operational_snapshot_block(analysis_payload)
        finalized_sections = []

        for section in sections:
            section_title = section["title"]
            section_body = section["body"]

            if section_title == "Analisis Deskriptif Cashflow":
                section_body = self._append_marker_block(section_body, chart_marker)
            elif section_title == "Analisis Diagnostik Cashflow":
                section_body = self._inject_subheading_block(
                    section_body,
                    "Konteks OSINT Pendukung",
                    macro_osint or "OSINT tidak dipakai karena tidak ada konteks eksternal yang cukup sebanding dengan profil perusahaan.",
                    before_subheading="Risiko dan Kontrol",
                )
            elif section_title == "Analisis Prediktif Cashflow":
                section_body = self._inject_subheading_block(
                    section_body,
                    "Snapshot Dashboard Operasional",
                    operational_snapshot,
                    before_subheading="Skenario 1-2 Kuartal",
                )
                if dashboard_markers:
                    section_body = self._inject_subheading_block(
                        section_body,
                        "Visual Dashboard Snapshot",
                        "\n".join(dashboard_markers),
                        before_subheading="Implikasi terhadap Arus Kas Masuk dan Keluar",
                    )
            elif section_title == "Rekomendasi Preskriptif":
                section_body = self._append_marker_block(section_body, flow_marker)

            finalized_sections.append({"title": section_title, "body": section_body})

        return self._join_top_level_sections(finalized_sections)

    def _build_fallback_report(self, report_context, notes, analysis_context, macro_osint, analysis_payload=None):
        chart_marker, flow_marker = self._extract_visual_markers(report_context.get("visual_prompt", ""))
        dashboard_markers = self._build_dashboard_visual_markers(analysis_payload)
        operational_snapshot = self._build_operational_snapshot_block(analysis_payload)
        focus_block = notes.strip() if notes and notes.strip() else "Tidak ada fokus tambahan dari pengguna."
        structured_context_block = self._format_structured_context_block(analysis_context)

        lines = [
            "# Ringkasan Eksekutif",
            "### Dampak Bisnis",
            "- Laporan ini digunakan untuk membantu manajemen membaca risiko cashflow, memahami prioritas penagihan, mengendalikan tekanan cash out, dan menentukan tindakan yang paling cepat berdampak pada ending cash.",
            "- Fokus pengguna tetap dijaga dalam interpretasi, namun narasi dan prioritas diturunkan dari bukti internal, pola historis, serta konteks pelaksanaan yang tersedia saat ini.",
            "",
            "## Fakta Eksekutif",
            report_context["executive_facts"],
            "",
            "### Tingkat Keyakinan dan Caveat",
            report_context["confidence_summary"],
        ]

        lines.extend(
            [
                "",
                "# Analisis Deskriptif Cashflow",
                "### Snapshot Portofolio dan Konsentrasi Risiko",
                report_context["financial_summary"],
                "",
                "### Batasan Data dan Asumsi",
                report_context["assumptions"],
                "",
                "# Analisis Diagnostik Cashflow",
                "### Pola Hambatan Utama",
                report_context["diagnostic_breakdown"],
                "",
                "### Bukti Internal yang Mewakili",
                report_context["evidence"],
                "",
                "### Konteks OSINT Pendukung",
                macro_osint or "OSINT tidak dipakai karena tidak ada konteks eksternal yang cukup sebanding dengan profil perusahaan.",
                "",
                "### Risiko dan Kontrol",
                report_context["controls"],
                "",
                "### Fokus Pengguna",
                f"- {focus_block}",
            ]
        )

        if structured_context_block:
            lines.extend(
                [
                    "",
                    "### Parameter Forecast dan Ruang Lingkup",
                    structured_context_block,
                ]
            )

        if chart_marker:
            lines.extend(["", chart_marker])

        lines.extend(
            [
                "",
                "# Analisis Prediktif Cashflow",
                "### Dasar Proyeksi",
                "- Proyeksi menggunakan pendekatan risk-adjusted berdasarkan campuran kelas pembayaran historis, sehingga hasil harus dibaca sebagai skenario manajemen, bukan kepastian arus kas masuk.",
                "- Base case mewakili perilaku penagihan yang paling mungkin terjadi bila pola historis bertahan, sedangkan upside dan downside menunjukkan ruang perbaikan atau penurunan.",
                "",
                "### Snapshot Dashboard Operasional",
                operational_snapshot or "- Snapshot dashboard operasional belum tersedia pada saat laporan dibentuk.",
                "",
                "### Skenario 1-2 Kuartal",
                report_context["scenario_table"],
                "",
                "### Implikasi terhadap Arus Kas Masuk dan Keluar",
                report_context["cash_plan_implications"],
                "",
                "# Rekomendasi Preskriptif",
                "### Prinsip Tindakan",
                "1. Dahulukan invoice bernilai besar dengan skor risiko tinggi dan penyebab yang masih bisa dipulihkan dalam 30 hari.",
                "2. Pisahkan treatment untuk isu anggaran, approval, administrasi, likuiditas, sengketa, dan kewajiban jatuh tempo agar tindakan inflow dan outflow tidak tercampur.",
                "3. Gunakan bukti internal dan jadwal tindak lanjut yang terdokumentasi agar eskalasi ke manajemen klien lebih kuat.",
                "",
                "### Prasyarat Implementasi",
                report_context["implementation_prerequisites"],
                "",
                "### Kesiapan Pelaksanaan",
                report_context["organizational_readiness"],
                "",
                "# Prioritas Tindakan 30 Hari",
                "### Tabel Prioritas",
                report_context["priority_table"],
            ]
        )

        lines.extend(
            [
                "",
                "### Catatan Pelaksanaan",
                "- Tetapkan owner utama per akun prioritas dan review statusnya minimal mingguan.",
                "- Gunakan rapat internal untuk memastikan hambatan administratif dan eskalasi ke klien ditutup dengan tenggat yang jelas.",
                "- Konsolidasikan hasil follow-up ke finance collection, account owner, dan sponsor bisnis agar keputusan rapat langsung dapat dieksekusi.",
            ]
        )

        if dashboard_markers:
            dashboard_insert_index = lines.index("### Implikasi terhadap Arus Kas Masuk dan Keluar") + 2
            dashboard_block = ["", "### Visual Dashboard Snapshot", *dashboard_markers, ""]
            lines[dashboard_insert_index:dashboard_insert_index] = dashboard_block

        if flow_marker:
            lines.extend(["", flow_marker])

        return "\n".join(lines)

    def _run_generation_pass(self, report_context, notes, analysis_context, macro_osint, active_sections, include_visuals, label):
        prompt = self._build_report_prompt(
            report_context,
            notes,
            analysis_context,
            macro_osint,
            active_sections,
            include_visuals,
        )
        user_instruction = self._build_user_instruction(notes, active_sections)
        response = self.ollama.chat(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": user_instruction},
            ],
            options={
                "num_ctx": REPORT_NUM_CTX,
                "num_predict": REPORT_NUM_PREDICT,
                "temperature": REPORT_TEMPERATURE,
                "top_p": REPORT_TOP_P,
                "repeat_penalty": REPORT_REPEAT_PENALTY,
            },
        )
        logger.info(
            "Generation pass %s completed with done_reason=%s, eval_count=%s.",
            label,
            response.get("done_reason"),
            response.get("eval_count"),
        )
        return response["message"]["content"]

    def run(self, notes="", analysis_context="", analysis_payload=None):
        logger.info("Starting cashflow intelligence report generation.")

        osint_context = "\n".join(part for part in (notes, analysis_context) if str(part or "").strip())
        global_osint_future = self.io_pool.submit(Researcher.get_macro_finance_trends, osint_context)
        report_context = self.kb.get_report_context(notes)

        try:
            macro_osint = global_osint_future.result(timeout=45)
        except Exception:
            macro_osint = "OSINT tidak dipakai karena konteks eksternal yang cukup sebanding tidak tersedia."

        fallback_used = False
        generated_sections = []
        for section_pass in self.SECTION_PASSES:
            generated_sections.append(
                self._run_generation_pass(
                    report_context,
                    notes,
                    analysis_context,
                    macro_osint,
                    section_pass["sections"],
                    section_pass["include_visuals"],
                    section_pass["label"],
                ).strip()
            )

        generated_content = "\n\n".join(section for section in generated_sections if section).strip()
        generated_content = self._finalize_report_content(generated_content, report_context, macro_osint, analysis_payload=analysis_payload)
        completeness_result = self._score_report_completeness(generated_content)
        logger.info(
            "Report completeness score %.1f/100 before fallback.",
            completeness_result["score"],
        )
        if not completeness_result["passed"]:
            logger.warning("Generated report failed quality gate. Falling back to deterministic management draft.")
            fallback_used = True
            generated_content = self._build_fallback_report(report_context, notes, analysis_context, macro_osint, analysis_payload=analysis_payload)
            generated_content = self._finalize_report_content(generated_content, report_context, macro_osint, analysis_payload=analysis_payload)
            completeness_result = self._score_report_completeness(generated_content)
            logger.info(
                "Report completeness score %.1f/100 after fallback.",
                completeness_result["score"],
            )

        document = Document()
        DocumentBuilder.create_cover(document, DEFAULT_COLOR)
        DocumentBuilder.add_table_of_contents(document)
        self._embed_dashboard_screenshots(document, analysis_payload, DEFAULT_COLOR)
        DocumentBuilder.process_content(
            document,
            generated_content,
            DEFAULT_COLOR,
        )

        dashboard_screenshots_included = self._has_dashboard_screenshots(analysis_payload)
        run_metadata = {
            "fallback_used": fallback_used,
            "quality_gate_passed": completeness_result["passed"],
            "completeness_score": completeness_result["score"],
            "completeness_missing": completeness_result["missing"],
            "osint_available": bool(
                macro_osint
                and "tidak tersedia" not in macro_osint.lower()
                and "tidak ada data osint" not in macro_osint.lower()
                and "osint tidak dipakai" not in macro_osint.lower()
                and "tidak ada sinyal eksternal yang cukup sebanding" not in macro_osint.lower()
            ),
            "visuals_included": any(marker in generated_content for marker in ("[[CHART:", "[[FLOW:", "[[DASHBOARD:")) or dashboard_screenshots_included,
            "dashboard_screenshots_included": dashboard_screenshots_included,
            "report_length": len(generated_content),
        }

        return document, "Inixindo_Cashflow_Intelligence_Report", run_metadata

    @staticmethod
    def _has_dashboard_screenshots(analysis_payload):
        payload = ReportGenerator._normalize_analysis_payload(analysis_payload)
        screenshots = payload.get("dashboard_screenshots")
        return bool(screenshots and isinstance(screenshots, list) and len(screenshots) > 0)

    @staticmethod
    def _embed_dashboard_screenshots(document, analysis_payload, theme_color):
        payload = ReportGenerator._normalize_analysis_payload(analysis_payload)
        screenshots = payload.get("dashboard_screenshots")
        if not screenshots or not isinstance(screenshots, list):
            return

        valid_screenshots = [
            shot for shot in screenshots
            if isinstance(shot, dict) and shot.get("image_base64")
        ]
        if not valid_screenshots:
            return

        heading = document.add_heading("Dashboard Cashflow Snapshot", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*theme_color)

        intro = document.add_paragraph(
            "Visual berikut diambil langsung dari dashboard operasional pada saat laporan diminta. "
            "Setiap horizon menampilkan kondisi kas, runway, coverage ratio, prediksi saldo, "
            "distribusi delay pembayaran, dan daftar akun overdue utama."
        )
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in intro.runs:
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)

        for shot in valid_screenshots:
            try:
                image_bytes = base64.b64decode(shot["image_base64"])
                image_stream = io.BytesIO(image_bytes)
                horizon_label = shot.get("horizon_label") or shot.get("horizon_key") or "Dashboard"

                image_paragraph = document.add_paragraph()
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_paragraph.add_run().add_picture(image_stream, width=Inches(6.8))

                try:
                    caption = document.add_paragraph(
                        f"Dashboard snapshot - {horizon_label}",
                        style="Caption",
                    )
                except KeyError:
                    caption = document.add_paragraph(
                        f"Dashboard snapshot - {horizon_label}"
                    )
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption.runs:
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(100, 100, 100)

            except Exception as exc:
                logger.warning("Failed to embed dashboard screenshot for %s: %s", shot.get("horizon_key"), exc)
                continue

        document.add_page_break()
