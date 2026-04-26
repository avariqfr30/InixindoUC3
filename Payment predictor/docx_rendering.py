import io
import json
import re
import textwrap
from datetime import datetime

import markdown
import matplotlib
import matplotlib.patches as patches
import matplotlib.pyplot as plt
from bs4 import BeautifulSoup, NavigableString, Tag
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from config import DEFAULT_COLOR, WRITER_FIRM_NAME

matplotlib.use("Agg")

class StyleEngine:
    @staticmethod
    def _insert_field(paragraph, field_code, placeholder="1"):
        begin_run = paragraph.add_run()
        begin_field = OxmlElement("w:fldChar")
        begin_field.set(qn("w:fldCharType"), "begin")
        begin_run._r.append(begin_field)

        instruction_run = paragraph.add_run()
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = field_code
        instruction_run._r.append(instruction)

        separate_run = paragraph.add_run()
        separate_field = OxmlElement("w:fldChar")
        separate_field.set(qn("w:fldCharType"), "separate")
        separate_run._r.append(separate_field)

        paragraph.add_run(placeholder)

        end_run = paragraph.add_run()
        end_field = OxmlElement("w:fldChar")
        end_field.set(qn("w:fldCharType"), "end")
        end_run._r.append(end_field)

    @classmethod
    def insert_toc_field(cls, paragraph):
        cls._insert_field(
            paragraph,
            'TOC \\o "1-3" \\h \\z \\u',
            "Klik kanan lalu pilih Update Field untuk memuat daftar isi.",
        )

    @classmethod
    def apply_document_styles(cls, doc, theme_color):
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

            header = section.header
            header_paragraph = (
                header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            )
            header_paragraph.text = ""
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            header_paragraph.add_run("INIXINDO JOGJA | INTERNAL FINANCE REPORT")
            for run in header_paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(120, 120, 120)

            footer = section.footer
            footer_paragraph = (
                footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            )
            footer_paragraph.text = ""
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer_paragraph.add_run("STRICTLY CONFIDENTIAL | Page ")
            cls._insert_field(footer_paragraph, "PAGE")
            footer_paragraph.add_run(" of ")
            cls._insert_field(footer_paragraph, "NUMPAGES")
            for run in footer_paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(110, 110, 110)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = RGBColor(33, 37, 41)
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_paragraph.line_spacing = 1.15
        normal_paragraph.space_after = Pt(8)

        heading_1 = doc.styles["Heading 1"]
        heading_1.font.name = "Calibri"
        heading_1.font.size = Pt(16)
        heading_1.font.bold = True
        heading_1.font.color.rgb = RGBColor(*theme_color)
        heading_1.paragraph_format.space_before = Pt(18)
        heading_1.paragraph_format.space_after = Pt(8)
        heading_1.paragraph_format.keep_with_next = True

        heading_2 = doc.styles["Heading 2"]
        heading_2.font.name = "Calibri"
        heading_2.font.size = Pt(13)
        heading_2.font.bold = True
        heading_2.font.color.rgb = RGBColor(0, 0, 0)
        heading_2.paragraph_format.space_before = Pt(14)
        heading_2.paragraph_format.space_after = Pt(4)
        heading_2.paragraph_format.keep_with_next = True

        heading_3 = doc.styles["Heading 3"]
        heading_3.font.name = "Calibri"
        heading_3.font.size = Pt(12)
        heading_3.font.bold = True
        heading_3.font.color.rgb = RGBColor(64, 64, 64)
        heading_3.paragraph_format.space_before = Pt(10)
        heading_3.paragraph_format.space_after = Pt(4)
        heading_3.paragraph_format.keep_with_next = True

        for style_name in [
            "List Bullet",
            "List Bullet 2",
            "List Bullet 3",
            "List Number",
            "List Number 2",
            "List Number 3",
        ]:
            try:
                list_style = doc.styles[style_name]
                list_style.font.name = "Calibri"
                list_style.font.size = Pt(11)
                list_style.paragraph_format.line_spacing = 1.05
                list_style.paragraph_format.space_after = Pt(4)
            except KeyError:
                continue

        try:
            caption_style = doc.styles["Caption"]
            caption_style.font.name = "Calibri"
            caption_style.font.size = Pt(10)
            caption_style.font.italic = True
            caption_style.font.color.rgb = RGBColor(80, 80, 80)
        except KeyError:
            pass

class ChartEngine:
    @staticmethod
    def _theme_to_plt_color(theme_color):
        return tuple(component / 255 for component in theme_color)

    @staticmethod
    def _format_compact_currency(value):
        amount = float(value or 0)
        if abs(amount) >= 1_000_000_000:
            return f"Rp {amount / 1_000_000_000:.1f}M"
        if abs(amount) >= 1_000_000:
            return f"Rp {amount / 1_000_000:.0f} juta"
        return f"Rp {amount:,.0f}".replace(",", ".")

    @staticmethod
    def create_bar_chart(data_str, theme_color):
        try:
            parts = [part.strip() for part in data_str.split("|")]
            if len(parts) >= 3:
                title = parts[0]
                y_label = parts[1]
                raw_data = "|".join(parts[2:])
            else:
                title = "Distribusi Historis Kelas Pembayaran"
                y_label = "Persentase"
                raw_data = data_str

            labels = []
            values = []
            for chunk in raw_data.split(";"):
                if "," not in chunk:
                    continue
                label, value = chunk.split(",", 1)
                numeric_value = re.sub(r"[^\d.]", "", value)
                if not numeric_value:
                    continue
                labels.append(label.strip())
                values.append(float(numeric_value))

            if not labels:
                return None

            fig, axis = plt.subplots(figsize=(7, 4.5))
            axis.bar(
                labels,
                values,
                color=ChartEngine._theme_to_plt_color(theme_color),
                alpha=0.9,
                width=0.5,
            )
            axis.set_title(title, fontsize=12, fontweight="bold", pad=20)
            axis.set_ylabel(y_label, fontsize=10)
            axis.spines["top"].set_visible(False)
            axis.spines["right"].set_visible(False)

            image_stream = io.BytesIO()
            plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=150)
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception:
            return None

    @staticmethod
    def create_flowchart(data_str, theme_color):
        try:
            steps = [
                "\n".join(textwrap.wrap(step.strip(), width=18))
                for step in data_str.split("->")
                if step.strip()
            ]
            if len(steps) < 2:
                return None

            fig, axis = plt.subplots(figsize=(8, 3))
            axis.axis("off")
            x_positions = [index * 2.5 for index in range(len(steps))]

            for index in range(len(steps) - 1):
                axis.annotate(
                    "",
                    xy=(x_positions[index + 1] - 1.0, 0.5),
                    xytext=(x_positions[index] + 1.0, 0.5),
                    arrowprops={"arrowstyle": "-|>", "lw": 1.5},
                )

            for index, step in enumerate(steps):
                box = patches.FancyBboxPatch(
                    (x_positions[index] - 1.0, 0.1),
                    2.0,
                    0.8,
                    boxstyle="round,pad=0.1",
                    fc=ChartEngine._theme_to_plt_color(theme_color),
                    alpha=0.9,
                )
                axis.add_patch(box)
                axis.text(
                    x_positions[index],
                    0.5,
                    step,
                    ha="center",
                    va="center",
                    size=9,
                    color="white",
                    fontweight="bold",
                )

            axis.set_xlim(-1.2, (len(steps) - 1) * 2.5 + 1.2)
            axis.set_ylim(0, 1)

            image_stream = io.BytesIO()
            plt.savefig(
                image_stream,
                format="png",
                bbox_inches="tight",
                dpi=200,
                transparent=True,
            )
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception:
            return None

    @staticmethod
    def create_dashboard_snapshot(data_str, theme_color):
        try:
            payload = json.loads(data_str)
            horizon_label = payload.get("horizon_label") or "Dashboard Cashflow"
            horizon_focus = payload.get("horizon_focus") or "-"
            status = str(payload.get("status") or "-").upper()
            current_cash = float(payload.get("current_cash") or 0)
            runway_months = float(payload.get("runway_months") or 0)
            coverage_ratio = float(payload.get("coverage_ratio") or 0)
            average_delay_days = float(payload.get("average_delay_days") or 0)
            balance_points = payload.get("balance_projection") or []
            coverage_bars = payload.get("coverage_bars") or []

            fig = plt.figure(figsize=(9.4, 5.4))
            grid = fig.add_gridspec(2, 2, height_ratios=[1.0, 1.65], width_ratios=[1.55, 1.0], hspace=0.34, wspace=0.28)
            header_axis = fig.add_subplot(grid[0, :])
            balance_axis = fig.add_subplot(grid[1, 0])
            coverage_axis = fig.add_subplot(grid[1, 1])

            for axis in (header_axis,):
                axis.axis("off")

            fig.patch.set_facecolor("white")
            theme_rgb = ChartEngine._theme_to_plt_color(theme_color)
            safe_color = "#16a34a"
            watch_color = "#d97706"
            risk_color = "#dc2626"
            status_color = safe_color if status == "AMAN" else watch_color if status == "WASPADA" else risk_color

            header_axis.text(
                0.0,
                0.98,
                "Cashflow Health Dashboard",
                fontsize=16,
                fontweight="bold",
                color="#203152",
                va="top",
                transform=header_axis.transAxes,
            )
            header_axis.text(
                0.0,
                0.76,
                horizon_label,
                fontsize=12,
                fontweight="bold",
                color="#0f172a",
                va="top",
                transform=header_axis.transAxes,
            )
            header_axis.text(
                0.0,
                0.62,
                horizon_focus,
                fontsize=10,
                color="#475569",
                va="top",
                transform=header_axis.transAxes,
            )

            status_box = patches.FancyBboxPatch(
                (0.76, 0.72),
                0.2,
                0.16,
                boxstyle="round,pad=0.02,rounding_size=0.02",
                fc=status_color,
                ec="none",
                alpha=0.95,
                transform=header_axis.transAxes,
            )
            header_axis.add_patch(status_box)
            header_axis.text(
                0.86,
                0.8,
                f"STATUS: {status}",
                fontsize=11,
                fontweight="bold",
                color="white",
                ha="center",
                va="center",
                transform=header_axis.transAxes,
            )

            metric_cards = [
                ("Cash", ChartEngine._format_compact_currency(current_cash)),
                ("Runway", f"{runway_months:.1f} bulan"),
                ("Coverage", f"{coverage_ratio:.2f}x"),
                ("Delay", f"{average_delay_days:.0f} hari"),
            ]
            card_width = 0.22
            card_gap = 0.025
            card_y = 0.08
            for index, (label, value) in enumerate(metric_cards):
                x = index * (card_width + card_gap)
                card = patches.FancyBboxPatch(
                    (x, card_y),
                    card_width,
                    0.34,
                    boxstyle="round,pad=0.015,rounding_size=0.02",
                    fc="#f8fafc" if index != 0 else "#eef2ff",
                    ec="#cbd5e1",
                    lw=1,
                    transform=header_axis.transAxes,
                )
                header_axis.add_patch(card)
                header_axis.text(x + 0.02, card_y + 0.23, label, fontsize=9.5, fontweight="bold", color="#475569", transform=header_axis.transAxes)
                header_axis.text(x + 0.02, card_y + 0.08, value, fontsize=13.5, fontweight="bold", color="#111827", transform=header_axis.transAxes)

            balance_axis.set_title("Prediksi Saldo", fontsize=11, fontweight="bold", loc="left", color="#334155")
            if balance_points:
                x_values = list(range(len(balance_points)))
                balances = [float(point.get("balance") or 0) for point in balance_points]
                labels = [str(point.get("label") or "") for point in balance_points]
                balance_axis.fill_between(x_values, balances, color=theme_rgb, alpha=0.18)
                balance_axis.plot(x_values, balances, color=theme_rgb, marker="o", linewidth=2)
                balance_axis.set_xticks(x_values)
                balance_axis.set_xticklabels(labels, fontsize=8)
                min_balance = min(balances)
                threshold = 100_000_000
                if min_balance <= threshold:
                    balance_axis.axhline(threshold, color=risk_color, linestyle="--", linewidth=1.2)
                balance_axis.grid(axis="y", linestyle="--", alpha=0.25)
            else:
                balance_axis.text(0.5, 0.5, "Tidak ada proyeksi saldo.", ha="center", va="center", transform=balance_axis.transAxes)
            balance_axis.spines["top"].set_visible(False)
            balance_axis.spines["right"].set_visible(False)
            balance_axis.spines["left"].set_color("#cbd5e1")
            balance_axis.spines["bottom"].set_color("#cbd5e1")
            balance_axis.tick_params(axis="y", labelsize=8)

            coverage_axis.set_title("Coverage & Runway", fontsize=11, fontweight="bold", loc="left", color="#334155")
            if coverage_bars:
                labels = [str(bar.get("label") or "") for bar in coverage_bars]
                values = [float(bar.get("value") or 0) for bar in coverage_bars]
                colors = []
                for bar in coverage_bars:
                    variant = str(bar.get("variant") or "")
                    if variant == "danger":
                        colors.append(risk_color)
                    elif variant == "target":
                        colors.append(watch_color)
                    elif variant == "current":
                        colors.append("#2563eb")
                    else:
                        colors.append("#0f766e")
                y_positions = list(range(len(values)))
                coverage_axis.barh(y_positions, values, color=colors, alpha=0.9)
                coverage_axis.set_yticks(y_positions)
                coverage_axis.set_yticklabels(labels, fontsize=8)
                coverage_axis.invert_yaxis()
                for index, value in enumerate(values):
                    coverage_axis.text(value + 0.03, index, f"{value:.2f}", va="center", fontsize=8)
                coverage_axis.axvline(1.0, color=risk_color, linestyle="--", linewidth=1)
                coverage_axis.axvline(1.2, color=watch_color, linestyle="--", linewidth=1)
                coverage_axis.grid(axis="x", linestyle="--", alpha=0.2)
            else:
                coverage_axis.text(0.5, 0.5, "Tidak ada data coverage.", ha="center", va="center", transform=coverage_axis.transAxes)
            coverage_axis.spines["top"].set_visible(False)
            coverage_axis.spines["right"].set_visible(False)
            coverage_axis.spines["left"].set_color("#cbd5e1")
            coverage_axis.spines["bottom"].set_color("#cbd5e1")
            coverage_axis.tick_params(axis="x", labelsize=8)

            image_stream = io.BytesIO()
            plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=170)
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception:
            return None

class DocumentBuilder:
    VISUAL_MARKER_PREFIXES = {
        "CHART": "[[CHART:",
        "DASHBOARD": "[[DASHBOARD:",
        "FLOW": "[[FLOW:",
    }

    @staticmethod
    def _append_inline_text(paragraph, node, bold=False, italic=False, underline=False, monospace=False):
        if isinstance(node, NavigableString):
            text = str(node)
            if not text:
                return
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if monospace:
                run.font.name = "Consolas"
            return

        if not isinstance(node, Tag):
            return

        next_bold = bold or node.name in {"strong", "b"}
        next_italic = italic or node.name in {"em", "i"}
        next_underline = underline or node.name == "u"
        next_monospace = monospace or node.name == "code"

        if node.name == "br":
            paragraph.add_run("\n")
            return

        if node.name == "a":
            for child in node.children:
                DocumentBuilder._append_inline_text(
                    paragraph,
                    child,
                    bold=next_bold,
                    italic=next_italic,
                    underline=True,
                    monospace=next_monospace,
                )
            return

        for child in node.children:
            DocumentBuilder._append_inline_text(
                paragraph,
                child,
                bold=next_bold,
                italic=next_italic,
                underline=next_underline,
                monospace=next_monospace,
            )

    @staticmethod
    def _set_cell_shading(cell, fill):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def _format_table_cell(cell, header=False):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(8.5 if not header else 9)
                if header:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

    @classmethod
    def _format_table(cls, table, header=True):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                if header and row_index == 0:
                    cls._set_cell_shading(cell, "C00000")
                    cls._format_table_cell(cell, header=True)
                else:
                    cls._format_table_cell(cell, header=False)

    @staticmethod
    def _resolve_list_style(doc, ordered, level):
        ordered_styles = ["List Number", "List Number 2", "List Number 3"]
        bullet_styles = ["List Bullet", "List Bullet 2", "List Bullet 3"]
        style_names = ordered_styles if ordered else bullet_styles
        preferred_style = style_names[min(level, len(style_names) - 1)]

        try:
            doc.styles[preferred_style]
            return preferred_style
        except KeyError:
            return "List Number" if ordered else "List Bullet"

    @classmethod
    def _add_list(cls, doc, list_tag, level=0, ordered=False):
        style_name = cls._resolve_list_style(doc, ordered, level)

        for list_item in list_tag.find_all("li", recursive=False):
            inline_nodes = []
            nested_lists = []

            for child in list_item.children:
                if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                    nested_lists.append(child)
                else:
                    inline_nodes.append(child)

            if inline_nodes:
                paragraph = doc.add_paragraph(style=style_name)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for node in inline_nodes:
                    cls._append_inline_text(paragraph, node)

            for nested_list in nested_lists:
                cls._add_list(
                    doc,
                    nested_list,
                    level=level + 1,
                    ordered=nested_list.name == "ol",
                )

    @staticmethod
    def _add_table(doc, table_tag):
        rows = table_tag.find_all("tr")
        if not rows:
            return

        header_cells = rows[0].find_all(["th", "td"])
        if not header_cells:
            return

        column_count = len(header_cells)
        table = doc.add_table(rows=1, cols=column_count)
        table.style = "Table Grid"

        for column_index, cell in enumerate(header_cells):
            paragraph = table.rows[0].cells[column_index].paragraphs[0]
            paragraph.text = cell.get_text(" ", strip=True)
            if paragraph.runs:
                paragraph.runs[0].bold = True

        for html_row in rows[1:]:
            html_cells = html_row.find_all(["th", "td"])
            table_cells = table.add_row().cells
            for column_index in range(column_count):
                value = ""
                if column_index < len(html_cells):
                    value = html_cells[column_index].get_text(" ", strip=True)
                table_cells[column_index].text = value
        DocumentBuilder._format_table(table)

    @classmethod
    def parse_html_to_docx(cls, doc, html_content, theme_color):
        soup = BeautifulSoup(html_content, "html.parser")

        for element in soup.contents:
            if isinstance(element, NavigableString):
                if not str(element).strip():
                    continue
                paragraph = doc.add_paragraph(str(element).strip())
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                continue

            if not isinstance(element, Tag):
                continue

            if element.name in {"h1", "h2", "h3", "h4"}:
                level = min(max(int(element.name[1]), 1), 3)
                heading = doc.add_heading(level=level)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cls._append_inline_text(heading, element)
                if level == 1:
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(*theme_color)
                continue

            if element.name == "p":
                if not element.get_text(" ", strip=True):
                    continue
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for child in element.children:
                    cls._append_inline_text(paragraph, child)
                continue

            if element.name in {"ul", "ol"}:
                cls._add_list(doc, element, level=0, ordered=element.name == "ol")
                continue

            if element.name == "table":
                cls._add_table(doc, element)

    @classmethod
    def _flush_markdown_block(cls, doc, lines, theme_color):
        if not lines:
            return

        markdown_text = "\n".join(lines).strip()
        if not markdown_text:
            return

        markdown_text = re.sub(r"\n{3,}", "\n\n", markdown_text)
        html_content = markdown.markdown(markdown_text, extensions=["tables", "sane_lists"])
        cls.parse_html_to_docx(doc, html_content, theme_color)

    @staticmethod
    def _add_visual(doc, marker_type, marker_payload, theme_color):
        if marker_type == "CHART":
            image = ChartEngine.create_bar_chart(marker_payload, theme_color)
            width = Inches(5.8)
            caption = "Grafik distribusi historis kelas pembayaran"
        elif marker_type == "DASHBOARD":
            image = ChartEngine.create_dashboard_snapshot(marker_payload, theme_color)
            try:
                payload = json.loads(marker_payload)
                caption_suffix = payload.get("horizon_label") or "Horizon aktif"
            except Exception:
                caption_suffix = "Horizon aktif"
            width = Inches(6.6)
            caption = f"Dashboard cashflow snapshot - {caption_suffix}"
        else:
            image = ChartEngine.create_flowchart(marker_payload, theme_color)
            width = Inches(6.3)
            caption = "Diagram alur rekomendasi mitigasi"

        if image is None:
            return

        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(image, width=width)

        try:
            caption_paragraph = doc.add_paragraph(caption, style="Caption")
        except KeyError:
            caption_paragraph = doc.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @classmethod
    def process_content(cls, doc, raw_text, theme_color=DEFAULT_COLOR):
        markdown_buffer = []

        for raw_line in raw_text.splitlines():
            stripped_line = raw_line.strip()
            marker_type, marker_payload = cls._parse_visual_marker(stripped_line)
            if marker_type:
                cls._flush_markdown_block(doc, markdown_buffer, theme_color)
                markdown_buffer = []
                cls._add_visual(doc, marker_type, marker_payload, theme_color)
                continue

            markdown_buffer.append(raw_line.rstrip())

        cls._flush_markdown_block(doc, markdown_buffer, theme_color)

    @classmethod
    def _parse_visual_marker(cls, stripped_line):
        if not stripped_line.endswith("]]"):
            return None, None
        for marker_type, prefix in cls.VISUAL_MARKER_PREFIXES.items():
            if stripped_line.startswith(prefix):
                return marker_type, stripped_line.replace(prefix, "", 1).rsplit("]]", 1)[0].strip()
        return None, None

    @staticmethod
    def create_cover(doc, theme_color=DEFAULT_COLOR):
        StyleEngine.apply_document_styles(doc, theme_color)

        properties = doc.core_properties
        properties.title = "Inixindo Cashflow Intelligence Report"
        properties.subject = "Internal Cashflow Intelligence Report"
        properties.author = WRITER_FIRM_NAME
        properties.category = "Finance"

        for _ in range(4):
            doc.add_paragraph()

        confidentiality = doc.add_paragraph("STRICTLY CONFIDENTIAL")
        confidentiality.alignment = WD_ALIGN_PARAGRAPH.CENTER
        confidentiality.runs[0].font.name = "Calibri"
        confidentiality.runs[0].font.size = Pt(10)
        confidentiality.runs[0].font.bold = True
        confidentiality.runs[0].font.color.rgb = RGBColor(120, 120, 120)

        doc.add_paragraph()

        title = doc.add_paragraph("CASHFLOW INTELLIGENCE REPORT")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.name = "Calibri"
        title.runs[0].font.size = Pt(22)
        title.runs[0].font.bold = True

        organization = doc.add_paragraph("INIXINDO JOGJA")
        organization.alignment = WD_ALIGN_PARAGRAPH.CENTER
        organization.runs[0].font.name = "Calibri"
        organization.runs[0].font.size = Pt(34)
        organization.runs[0].font.bold = True
        organization.runs[0].font.color.rgb = RGBColor(*theme_color)

        doc.add_paragraph()

        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = "Table Grid"
        metadata = [
            ("Cakupan Data", "Seluruh histori invoice dan catatan penagihan"),
            ("Tipe Laporan", "Analisis deskriptif, diagnostik, prediktif, dan preskriptif cashflow"),
            ("Tanggal Generasi", datetime.now().strftime("%d %B %Y")),
            ("Disusun Oleh", WRITER_FIRM_NAME),
        ]

        for row_index, (label, value) in enumerate(metadata):
            left_cell = metadata_table.rows[row_index].cells[0]
            right_cell = metadata_table.rows[row_index].cells[1]
            left_cell.text = label
            right_cell.text = value
            if left_cell.paragraphs[0].runs:
                left_cell.paragraphs[0].runs[0].bold = True
        DocumentBuilder._format_table(metadata_table, header=False)

        doc.add_page_break()

    @staticmethod
    def add_table_of_contents(doc):
        doc.add_heading("Daftar Isi", level=1)
        toc_paragraph = doc.add_paragraph()
        StyleEngine.insert_toc_field(toc_paragraph)

        note = doc.add_paragraph(
            "Catatan: jika daftar isi belum muncul, klik kanan pada area daftar isi lalu pilih Update Field."
        )
        note.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(10)

        doc.add_page_break()
