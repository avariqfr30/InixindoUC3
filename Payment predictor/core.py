import base64
import copy
import concurrent.futures
import io
import json
import logging
import os
import re
import statistics
import textwrap
import threading
import time
from datetime import datetime, timedelta
from pathlib import Path
from urllib.parse import urljoin, urlparse

import chromadb
import diskcache as dc
import markdown
import matplotlib
import matplotlib.patches as patches
import matplotlib.pyplot as plt
import pandas as pd
import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from chromadb.config import Settings
from chromadb.utils import embedding_functions
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from pydantic import BaseModel, Field
from ollama import Client
from sqlalchemy import create_engine

from config import (
    APP_SERVER,
    CASH_OUT_API_AUTH_TOKEN,
    CASH_OUT_API_BASIC_PASSWORD,
    CASH_OUT_API_BASIC_USERNAME,
    CASH_OUT_API_BODY_JSON,
    CASH_OUT_API_ENDPOINT_URL,
    CASH_OUT_API_HEADERS_JSON,
    CASH_OUT_API_METHOD,
    CASH_OUT_API_QUERY_PARAMS_JSON,
    CASH_OUT_API_RECORDS_KEY,
    CASH_OUT_API_TIMEOUT,
    CASH_OUT_API_VERIFY_SSL,
    CASH_OUT_FIELD_MAP_JSON,
    DATA_ACQUISITION_MODE,
    DATA_SOURCE_ACTIVE_STATE_PATH,
    DATA_SOURCE_DEMO_PROFILE_PATH,
    DATA_SOURCE_PRODUCTION_PROFILE_PATH,
    DATA_DIR,
    DEFAULT_COLOR,
    DEMO_CSV_PATH,
    EMBED_MODEL,
    FINANCE_SYSTEM_PROMPT,
    INTERNAL_API_AUTH_TOKEN,
    INTERNAL_API_BASE_URL,
    INTERNAL_API_BASIC_PASSWORD,
    INTERNAL_API_BASIC_USERNAME,
    INTERNAL_API_BODY_JSON,
    INTERNAL_API_CONFIG_FILE,
    INTERNAL_API_DATASET_PATH,
    INTERNAL_API_ENDPOINT_URL,
    INTERNAL_API_FIELD_MAP_JSON,
    INTERNAL_API_HEADERS_JSON,
    INTERNAL_API_MAX_RETRIES,
    INTERNAL_API_METHOD,
    INTERNAL_API_PAGE_SIZE,
    INTERNAL_API_PAGINATION_CURSOR_KEY,
    INTERNAL_API_PAGINATION_LIMIT_PARAM,
    INTERNAL_API_PAGINATION_MAX_PAGES,
    INTERNAL_API_PAGINATION_MODE,
    INTERNAL_API_PAGINATION_OFFSET_PARAM,
    INTERNAL_API_QUERY_PARAMS_JSON,
    INTERNAL_API_RECORDS_KEY,
    INTERNAL_API_RETRY_BACKOFF_BASE,
    INTERNAL_API_TIMEOUT,
    INTERNAL_API_VERIFY_SSL,
    LLM_MODEL,
    OLLAMA_HOST,
    PERSONAS,
    REPORT_NUM_CTX,
    REPORT_MAX_CONCURRENT_JOBS,
    REPORT_MIN_COMPLETENESS_SCORE,
    REPORT_NUM_PREDICT,
    REPORT_REPEAT_PENALTY,
    REPORT_SECTION_SEQUENCE,
    REPORT_TEMPERATURE,
    REPORT_TOP_P,
    SERPER_API_KEY,
    WAITRESS_THREADS,
    WRITER_FIRM_NAME,
)
from data_contract import (
    build_internal_data_summary,
    extract_records_from_payload,
    get_internal_api_contract,
    normalize_financial_dataframe,
    parse_internal_api_field_map,
)
from data_sources import (
    load_available_source_profiles,
    resolve_active_source_profile,
    summarize_source_profile,
    write_active_source_key,
)
from forecast_engine import parse_idr_amount

matplotlib.use("Agg")
logger = logging.getLogger(__name__)


# ==========================================
# PYDANTIC SCHEMAS & FAST CACHING
# ==========================================
class InsightSchema(BaseModel):
    insight: str = Field(description="The extracted insight in Indonesian. 'NOT_FOUND' if missing.")

# Initialize ultra-fast disk caching for OSINT (survives server restarts)
osint_cache_dir = Path(DATA_DIR) / '.osint_cache' if DATA_DIR else Path('./.osint_cache')
osint_cache = dc.Cache(str(osint_cache_dir))
# ==========================================


class InternalAPIClient:
    def __init__(self, source_profile=None):
        profile = copy.deepcopy(source_profile) if source_profile else {
            "type": "json_api",
            "endpoint": {
                "url": INTERNAL_API_ENDPOINT_URL.strip(),
                "base_url": INTERNAL_API_BASE_URL.rstrip("/"),
                "path": INTERNAL_API_DATASET_PATH.strip() or "/api/finance/invoices",
                "method": (INTERNAL_API_METHOD or "GET").strip().upper(),
                "timeout": INTERNAL_API_TIMEOUT,
                "verify_ssl": INTERNAL_API_VERIFY_SSL,
                "records_key": INTERNAL_API_RECORDS_KEY.strip(),
            },
            "auth": {
                "bearer_token": INTERNAL_API_AUTH_TOKEN.strip(),
                "basic_username": INTERNAL_API_BASIC_USERNAME.strip(),
                "basic_password": INTERNAL_API_BASIC_PASSWORD,
            },
            "request": {
                "headers": self._parse_json_object(INTERNAL_API_HEADERS_JSON, "headers"),
                "query_params": self._parse_json_object(
                    INTERNAL_API_QUERY_PARAMS_JSON,
                    "query params",
                ),
                "body": self._parse_optional_json_value(INTERNAL_API_BODY_JSON, "body"),
            },
            "field_map": parse_internal_api_field_map(INTERNAL_API_FIELD_MAP_JSON),
            "pagination": {
                "mode": INTERNAL_API_PAGINATION_MODE,
                "page_size": INTERNAL_API_PAGE_SIZE,
                "cursor_key": INTERNAL_API_PAGINATION_CURSOR_KEY,
                "offset_param": INTERNAL_API_PAGINATION_OFFSET_PARAM,
                "limit_param": INTERNAL_API_PAGINATION_LIMIT_PARAM,
                "max_pages": INTERNAL_API_PAGINATION_MAX_PAGES,
            },
            "retry": {
                "max_retries": INTERNAL_API_MAX_RETRIES,
                "backoff_base": INTERNAL_API_RETRY_BACKOFF_BASE,
            },
        }
        endpoint = profile.get("endpoint", {}) or {}
        auth = profile.get("auth", {}) or {}
        request_config = profile.get("request", {}) or {}
        pagination_config = profile.get("pagination", {}) or {}
        retry_config = profile.get("retry", {}) or {}

        self.source_profile = profile
        self.endpoint_url = str(endpoint.get("url") or "").strip()
        self.base_url = str(endpoint.get("base_url") or "").strip().rstrip("/")
        self.dataset_path = str(endpoint.get("path") or INTERNAL_API_DATASET_PATH or "/api/finance/invoices").strip()
        self.method = str(endpoint.get("method") or "GET").strip().upper()
        self.records_key = str(endpoint.get("records_key") or "").strip()
        self.auth_token = str(auth.get("bearer_token") or "").strip()
        self.basic_username = str(auth.get("basic_username") or "").strip()
        self.basic_password = str(auth.get("basic_password") or "")
        self.timeout = int(endpoint.get("timeout") or INTERNAL_API_TIMEOUT)
        verify_ssl_value = endpoint.get("verify_ssl")
        if isinstance(verify_ssl_value, str):
            self.verify_ssl = verify_ssl_value.strip().lower() not in {"0", "false", "no"}
        elif verify_ssl_value is None:
            self.verify_ssl = INTERNAL_API_VERIFY_SSL
        else:
            self.verify_ssl = bool(verify_ssl_value)
        self.headers = dict(request_config.get("headers") or {})
        self.body = request_config.get("body")
        self.query_params = dict(request_config.get("query_params") or {})
        self.field_map = parse_internal_api_field_map(profile.get("field_map") or {})

        # Pagination config
        self.pagination_mode = str(pagination_config.get("mode") or "").strip().lower()
        self.page_size = int(pagination_config.get("page_size") or 0)
        self.pagination_cursor_key = str(pagination_config.get("cursor_key") or "").strip()
        self.pagination_offset_param = str(pagination_config.get("offset_param") or "offset").strip()
        self.pagination_limit_param = str(pagination_config.get("limit_param") or "limit").strip()
        self.pagination_max_pages = int(pagination_config.get("max_pages") or 50)

        # Retry config
        self.max_retries = max(int(retry_config.get("max_retries") or 3), 1)
        self.retry_backoff_base = max(float(retry_config.get("backoff_base") or 1.0), 0.1)

    @staticmethod
    def _parse_json_object(raw_value, label):
        if not raw_value:
            return {}

        try:
            parsed = json.loads(raw_value)
        except json.JSONDecodeError as exc:
            raise ValueError(f"Invalid INTERNAL_API_{label.upper().replace(' ', '_')}_JSON: {exc}") from exc

        if not isinstance(parsed, dict):
            raise ValueError(f"INTERNAL_API_{label.upper().replace(' ', '_')}_JSON must be a JSON object.")

        return parsed

    @staticmethod
    def _parse_optional_json_value(raw_value, label):
        if not raw_value:
            return None
        try:
            return json.loads(raw_value)
        except json.JSONDecodeError as exc:
            raise ValueError(f"Invalid INTERNAL_API_{label.upper()}_JSON: {exc}") from exc

    def is_configured(self):
        return bool(self.endpoint_url or self.base_url)

    def get_dataset_url(self):
        if self.endpoint_url:
            return self.endpoint_url
        return urljoin(f"{self.base_url}/", self.dataset_path.lstrip("/"))

    def validate_endpoint_url(self):
        """Boot-time connectivity pre-check. Returns (ok, message)."""
        if not self.is_configured():
            return False, "Internal API endpoint is not configured (INTERNAL_API_ENDPOINT_URL or INTERNAL_API_BASE_URL is empty)."
        dataset_url = self.get_dataset_url()
        parsed = urlparse(dataset_url)
        if parsed.scheme not in {"http", "https"}:
            return False, f"Internal API URL has invalid scheme '{parsed.scheme}': {dataset_url}"
        if not parsed.hostname:
            return False, f"Internal API URL has no hostname: {dataset_url}"
        try:
            requests.head(dataset_url, timeout=5, verify=self.verify_ssl)
            return True, f"Internal API endpoint is reachable: {dataset_url}"
        except requests.ConnectionError:
            return False, f"Internal API endpoint is unreachable (connection refused or DNS failure): {dataset_url}"
        except requests.Timeout:
            return False, f"Internal API endpoint timed out after 5s: {dataset_url}"
        except Exception as exc:
            return False, f"Internal API endpoint pre-check failed: {dataset_url} — {exc}"

    def _build_request_kwargs(self, extra_params=None):
        headers = {"Accept": "application/json"}
        headers.update(self.headers)
        auth = None
        if self.basic_username:
            auth = (self.basic_username, self.basic_password)
        if self.auth_token:
            headers.setdefault("Authorization", f"Bearer {self.auth_token}")

        params = dict(self.query_params)
        if extra_params:
            params.update(extra_params)

        request_kwargs = {
            "headers": headers,
            "params": params,
            "timeout": self.timeout,
            "verify": self.verify_ssl,
        }
        if auth:
            request_kwargs["auth"] = auth
        if self.method != "GET" and self.body is not None:
            request_kwargs["json"] = self.body
            headers.setdefault("Content-Type", "application/json")

        return request_kwargs, auth

    def _execute_request_with_retry(self, dataset_url, request_kwargs):
        """Execute HTTP request with exponential backoff retry."""
        last_exc = None
        for attempt in range(self.max_retries):
            try:
                response = requests.request(
                    self.method,
                    dataset_url,
                    **request_kwargs,
                )
                if response.status_code >= 500 and attempt < self.max_retries - 1:
                    logger.warning(
                        "Internal API returned HTTP %s on attempt %s/%s for %s %s. Retrying...",
                        response.status_code, attempt + 1, self.max_retries, self.method, dataset_url,
                    )
                    time.sleep(self.retry_backoff_base * (2 ** attempt))
                    continue
                if response.status_code >= 400:
                    body_preview = (response.text or "")[:300]
                    raise RuntimeError(
                        f"Internal API returned HTTP {response.status_code} for {self.method} {dataset_url}. "
                        f"Response body (truncated): {body_preview}"
                    )
                return response
            except requests.ConnectionError as exc:
                last_exc = exc
                if attempt < self.max_retries - 1:
                    wait = self.retry_backoff_base * (2 ** attempt)
                    logger.warning(
                        "Connection to %s failed on attempt %s/%s. Retrying in %.1fs...",
                        dataset_url, attempt + 1, self.max_retries, wait,
                    )
                    time.sleep(wait)
                    continue
                raise RuntimeError(
                    f"Internal API at {dataset_url} is unreachable after {self.max_retries} attempts. "
                    f"Check INTERNAL_API_ENDPOINT_URL and network connectivity. Last error: {exc}"
                ) from exc
            except requests.Timeout as exc:
                last_exc = exc
                if attempt < self.max_retries - 1:
                    wait = self.retry_backoff_base * (2 ** attempt)
                    logger.warning(
                        "Request to %s timed out on attempt %s/%s. Retrying in %.1fs...",
                        dataset_url, attempt + 1, self.max_retries, wait,
                    )
                    time.sleep(wait)
                    continue
                raise RuntimeError(
                    f"Internal API at {dataset_url} timed out after {self.timeout}s "
                    f"({self.max_retries} attempts). Increase INTERNAL_API_TIMEOUT or check endpoint performance."
                ) from exc
        raise RuntimeError(f"Internal API request failed after {self.max_retries} retries.") from last_exc

    def _parse_response_payload(self, response, dataset_url):
        """Parse and validate the JSON response, returning the raw payload."""
        try:
            payload = response.json()
        except ValueError:
            body_preview = (response.text or "")[:300]
            raise RuntimeError(
                f"Internal API at {dataset_url} returned non-JSON response "
                f"(Content-Type: {response.headers.get('Content-Type', 'unknown')}). "
                f"Body (truncated): {body_preview}"
            )

        if isinstance(payload, dict) and payload.get("success") is False:
            message = payload.get("message") or payload.get("error") or "unknown error"
            raise RuntimeError(
                f"Internal API returned a business error for {self.method} {dataset_url}: {message}"
            )
        if (
            isinstance(payload, dict)
            and payload.get("success") is True
            and "data" in payload
            and payload.get("data") is None
        ):
            raise RuntimeError(
                f"Internal API at {dataset_url} returned `data: null`. The endpoint is reachable and "
                "credentials are valid, but it still needs the correct POST payload. "
                "Set INTERNAL_API_BODY_JSON once the company shares the required request body."
            )
        return payload

    def _fetch_single_page(self, extra_params=None):
        """Fetch a single page of records."""
        dataset_url = self.get_dataset_url()
        request_kwargs, auth = self._build_request_kwargs(extra_params)
        response = self._execute_request_with_retry(dataset_url, request_kwargs)
        payload = self._parse_response_payload(response, dataset_url)
        return payload, response, auth

    def _is_pagination_enabled(self):
        return self.pagination_mode in {"offset", "cursor", "link"} and self.page_size > 0

    def _fetch_with_pagination(self):
        """Fetch all records across multiple pages."""
        dataset_url = self.get_dataset_url()
        all_records = []
        total_pages = 0

        if self.pagination_mode == "offset":
            offset = 0
            for page in range(self.pagination_max_pages):
                total_pages += 1
                extra_params = {
                    self.pagination_offset_param: offset,
                    self.pagination_limit_param: self.page_size,
                }
                payload, response, auth = self._fetch_single_page(extra_params)
                page_records, _ = extract_records_from_payload(
                    payload, explicit_records_path=self.records_key or None,
                )
                all_records.extend(page_records)
                if len(page_records) < self.page_size:
                    break
                offset += self.page_size

        elif self.pagination_mode == "cursor":
            cursor = None
            for page in range(self.pagination_max_pages):
                total_pages += 1
                extra_params = {self.pagination_limit_param: self.page_size}
                if cursor:
                    extra_params[self.pagination_cursor_key or "cursor"] = cursor
                payload, response, auth = self._fetch_single_page(extra_params)
                page_records, _ = extract_records_from_payload(
                    payload, explicit_records_path=self.records_key or None,
                )
                all_records.extend(page_records)
                next_cursor = None
                if isinstance(payload, dict):
                    cursor_path = self.pagination_cursor_key or "next_cursor"
                    next_cursor = payload.get(cursor_path) or payload.get("meta", {}).get(cursor_path)
                if not next_cursor or len(page_records) < self.page_size:
                    break
                cursor = next_cursor

        elif self.pagination_mode == "link":
            next_url = dataset_url
            for page in range(self.pagination_max_pages):
                total_pages += 1
                extra_params = {self.pagination_limit_param: self.page_size} if page == 0 else {}
                if page == 0:
                    payload, response, auth = self._fetch_single_page(extra_params)
                else:
                    request_kwargs, auth = self._build_request_kwargs()
                    request_kwargs["params"] = {}
                    response = self._execute_request_with_retry(next_url, request_kwargs)
                    payload = self._parse_response_payload(response, next_url)
                page_records, _ = extract_records_from_payload(
                    payload, explicit_records_path=self.records_key or None,
                )
                all_records.extend(page_records)
                link_header = response.headers.get("Link", "")
                next_url = self._parse_link_next(link_header)
                if not next_url or len(page_records) < self.page_size:
                    break

        logger.info("Pagination completed: %s records across %s pages.", len(all_records), total_pages)
        return all_records, auth, total_pages

    @staticmethod
    def _parse_link_next(link_header):
        """Parse RFC 5988 Link header for rel='next'."""
        if not link_header:
            return None
        for part in link_header.split(","):
            if 'rel="next"' in part or "rel='next'" in part:
                match = re.search(r"<([^>]+)>", part)
                if match:
                    return match.group(1)
        return None

    def fetch_records(self, preview_limit=0):
        if not self.is_configured():
            raise RuntimeError(
                "Internal API endpoint is not configured. "
                "Set INTERNAL_API_ENDPOINT_URL or INTERNAL_API_BASE_URL."
            )

        dataset_url = self.get_dataset_url()

        if self._is_pagination_enabled() and preview_limit <= 0:
            all_records, auth, total_pages = self._fetch_with_pagination()
            extraction_summary = {
                "strategy": "paginated",
                "paginationMode": self.pagination_mode,
                "pageSize": self.page_size,
                "totalPages": total_pages,
                "recordCount": len(all_records),
            }
        else:
            extra_params = {}
            if preview_limit > 0 and self.page_size > 0:
                extra_params[self.pagination_limit_param] = preview_limit
            payload, response, auth = self._fetch_single_page(extra_params)
            all_records, extraction_summary = extract_records_from_payload(
                payload,
                explicit_records_path=self.records_key or None,
            )

        extraction_summary["datasetUrl"] = dataset_url
        extraction_summary["requestMethod"] = self.method
        extraction_summary["authMode"] = "basic" if auth else ("bearer" if self.auth_token else "none")
        return all_records, extraction_summary


class CashOutAPIClient(InternalAPIClient):
    FIELD_ALIASES = {
        "amount": (
            "amount",
            "amount_idr",
            "nominal",
            "nilai",
            "cash_out",
            "outflow",
            "expense_amount",
            "planned_amount",
            "total_amount",
            "value",
        ),
        "due_date": (
            "due_date",
            "payment_date",
            "planned_date",
            "tanggal_bayar",
            "tanggal_jatuh_tempo",
            "jatuh_tempo",
            "due",
            "date",
            "tanggal",
        ),
        "category": (
            "category",
            "expense_category",
            "cash_out_category",
            "jenis_biaya",
            "kategori",
            "cost_center",
        ),
        "reference": (
            "reference",
            "reference_code",
            "code",
            "id",
            "reference_id",
            "document_no",
            "invoice_no",
            "nama",
            "name",
        ),
        "status": (
            "status",
            "payment_status",
            "approval_status",
            "state",
        ),
        "description": (
            "description",
            "note",
            "notes",
            "memo",
            "remarks",
            "keterangan",
        ),
    }

    REQUIRED_FIELDS = ("amount", "due_date")

    def __init__(self):
        self.endpoint_url = CASH_OUT_API_ENDPOINT_URL.strip()
        self.base_url = ""
        self.dataset_path = ""
        self.method = (CASH_OUT_API_METHOD or "GET").strip().upper()
        self.records_key = CASH_OUT_API_RECORDS_KEY.strip()
        self.auth_token = CASH_OUT_API_AUTH_TOKEN.strip()
        self.basic_username = CASH_OUT_API_BASIC_USERNAME.strip()
        self.basic_password = CASH_OUT_API_BASIC_PASSWORD
        self.timeout = CASH_OUT_API_TIMEOUT
        self.verify_ssl = CASH_OUT_API_VERIFY_SSL
        self.headers = self._parse_json_object(CASH_OUT_API_HEADERS_JSON, "headers")
        self.body = self._parse_optional_json_value(CASH_OUT_API_BODY_JSON, "body")
        self.field_map = self._parse_field_map(CASH_OUT_FIELD_MAP_JSON)
        self.query_params = self._parse_json_object(
            CASH_OUT_API_QUERY_PARAMS_JSON,
            "query params",
        )

    @classmethod
    def _parse_field_map(cls, raw_value):
        if not raw_value:
            return {}
        try:
            candidate = json.loads(raw_value)
        except json.JSONDecodeError as exc:
            raise ValueError(f"Invalid CASH_OUT_FIELD_MAP_JSON: {exc}") from exc
        if not isinstance(candidate, dict):
            raise ValueError("CASH_OUT_FIELD_MAP_JSON must be a JSON object.")

        normalized_map = {}
        for key, value in candidate.items():
            normalized_key = cls._normalize_key(key)
            if normalized_key not in cls.FIELD_ALIASES:
                raise ValueError(
                    "CASH_OUT_FIELD_MAP_JSON contains an unknown field key: "
                    f"{key}. Use amount, due_date, category, reference, status, or description."
                )
            normalized_map[normalized_key] = str(value).strip()
        return normalized_map

    @staticmethod
    def _normalize_key(value):
        return re.sub(r"[^a-z0-9]+", "_", str(value or "").strip().lower()).strip("_")

    @classmethod
    def _resolve_columns(cls, data_frame, explicit_field_map=None):
        explicit_field_map = explicit_field_map or {}
        columns = list(data_frame.columns)
        normalized_columns = {
            cls._normalize_key(column): column
            for column in columns
        }

        resolved = {}
        missing = []
        for canonical_field, aliases in cls.FIELD_ALIASES.items():
            explicit_column = explicit_field_map.get(canonical_field)
            if explicit_column:
                if explicit_column in data_frame.columns:
                    resolved[canonical_field] = explicit_column
                    continue
                normalized_explicit = cls._normalize_key(explicit_column)
                if normalized_explicit in normalized_columns:
                    resolved[canonical_field] = normalized_columns[normalized_explicit]
                    continue

            for alias in aliases:
                normalized_alias = cls._normalize_key(alias)
                if normalized_alias in normalized_columns:
                    resolved[canonical_field] = normalized_columns[normalized_alias]
                    break

            if canonical_field in cls.REQUIRED_FIELDS and canonical_field not in resolved:
                missing.append(canonical_field)

        return resolved, missing

    @staticmethod
    def _normalize_records(records):
        data_frame = pd.json_normalize(records, sep="_")
        if data_frame.empty:
            return data_frame

        for column in data_frame.columns:
            data_frame[column] = data_frame[column].apply(
                lambda value: json.dumps(value, ensure_ascii=False)
                if isinstance(value, (dict, list))
                else value
            )

        data_frame.columns = [str(column).strip() for column in data_frame.columns]
        return data_frame

    @staticmethod
    def _parse_due_date(raw_value):
        if raw_value is None or (isinstance(raw_value, float) and pd.isna(raw_value)):
            return None
        parsed_value = pd.to_datetime(raw_value, errors="coerce")
        if pd.isna(parsed_value):
            return None
        return parsed_value.to_pydatetime()

    @staticmethod
    def _is_open_status(raw_status):
        status = str(raw_status or "").strip().lower()
        if not status:
            return True
        return status not in {"paid", "settled", "closed", "cancelled", "canceled", "done"}

    def fetch_commitments(self):
        if not self.is_configured():
            return [], self.get_summary()

        records, extraction_summary = self.fetch_records()
        raw_data_frame = self._normalize_records(records)
        if raw_data_frame.empty:
            return [], {
                **self.get_summary(),
                "datasetUrl": self.get_dataset_url(),
                "requestMethod": self.method,
                "recordCount": 0,
                "missingRequiredFields": list(self.REQUIRED_FIELDS),
            }

        resolved_columns, missing_fields = self._resolve_columns(raw_data_frame, self.field_map)
        normalized_records = []
        for _, row in raw_data_frame.iterrows():
            amount_column = resolved_columns.get("amount")
            due_date_column = resolved_columns.get("due_date")
            try:
                amount = parse_idr_amount(row.get(amount_column, 0)) if amount_column else 0
            except ValueError:
                continue
            due_date = self._parse_due_date(row.get(due_date_column)) if due_date_column else None
            if amount <= 0 or due_date is None:
                continue
            normalized_records.append(
                {
                    "amount": amount,
                    "due_date": due_date,
                    "category": str(row.get(resolved_columns.get("category"), "")).strip(),
                    "reference": str(row.get(resolved_columns.get("reference"), "")).strip(),
                    "status": str(row.get(resolved_columns.get("status"), "")).strip(),
                    "description": str(row.get(resolved_columns.get("description"), "")).strip(),
                    "is_open": self._is_open_status(row.get(resolved_columns.get("status"), "")),
                }
            )

        summary = {
            **self.get_summary(),
            **extraction_summary,
            "recordCount": len(normalized_records),
            "resolvedFields": resolved_columns,
            "missingRequiredFields": missing_fields,
        }
        return normalized_records, summary

    def get_summary(self):
        return {
            "configured": self.is_configured(),
            "datasetUrl": self.get_dataset_url() if self.is_configured() else None,
            "requestMethod": self.method,
        }


class CashOutStore:
    def __init__(self):
        self.client = CashOutAPIClient()
        self.records = []
        self.summary = self.client.get_summary()
        self.lock = threading.Lock()
        self.refresh_lock = threading.Lock()
        self.version = 0
        self.sync_status = "not_configured" if not self.client.is_configured() else "not_loaded"
        self.last_sync_started_at = None
        self.last_sync_at = None
        self.last_success_at = None
        self.last_sync_duration_seconds = None
        self.last_sync_error = None
        if self.client.is_configured():
            self.refresh_data()

    def refresh_data(self):
        if not self.client.is_configured():
            with self.lock:
                self.records = []
                self.summary = self.client.get_summary()
                self.sync_status = "not_configured"
                self.last_sync_error = None
            return False

        with self.refresh_lock:
            started_at = datetime.now()
            with self.lock:
                self.sync_status = "refreshing"
                self.last_sync_started_at = started_at
                self.last_sync_error = None

            try:
                records, summary = self.client.fetch_commitments()
                completed_at = datetime.now()
                with self.lock:
                    self.records = records
                    self.summary = summary
                    self.version += 1
                    self.sync_status = "ready"
                    self.last_sync_at = completed_at
                    self.last_success_at = completed_at
                    self.last_sync_duration_seconds = round((completed_at - started_at).total_seconds(), 2)
                    self.last_sync_error = None
                return True
            except Exception as exc:
                completed_at = datetime.now()
                logger.error("Cash-out source sync failed: %s", exc)
                with self.lock:
                    self.records = []
                    self.summary = {
                        **self.client.get_summary(),
                        "recordCount": 0,
                        "missingRequiredFields": list(self.client.REQUIRED_FIELDS),
                    }
                    self.sync_status = "error"
                    self.last_sync_at = completed_at
                    self.last_sync_duration_seconds = round((completed_at - started_at).total_seconds(), 2)
                    self.last_sync_error = str(exc)
                return False

    def get_records(self):
        with self.lock:
            return copy.deepcopy(self.records)

    def get_status(self, refresh_interval_seconds=0):
        with self.lock:
            last_success_at = self.last_success_at
            next_refresh_at = None
            if refresh_interval_seconds > 0 and last_success_at is not None:
                next_refresh_at = last_success_at + timedelta(seconds=refresh_interval_seconds)
            source_age_minutes = None
            if last_success_at is not None:
                source_age_minutes = round((datetime.now() - last_success_at).total_seconds() / 60, 1)
            return {
                "configured": self.client.is_configured(),
                "syncStatus": self.sync_status,
                "lastSyncStartedAt": self.last_sync_started_at.isoformat() if self.last_sync_started_at else None,
                "lastSyncAt": self.last_sync_at.isoformat() if self.last_sync_at else None,
                "lastSuccessAt": last_success_at.isoformat() if last_success_at else None,
                "lastSyncDurationSeconds": self.last_sync_duration_seconds,
                "lastSyncError": self.last_sync_error,
                "sourceAgeMinutes": source_age_minutes,
                "nextRefreshAt": next_refresh_at.isoformat() if next_refresh_at else None,
                "recordCount": len(self.records),
                "version": self.version,
                "summary": copy.deepcopy(self.summary),
            }


class KnowledgeBase:
    def __init__(self, db_uri):
        os.makedirs(DATA_DIR, exist_ok=True)
        self.engine = create_engine(db_uri)
        self.source_registry = {}
        self.source_registry_issues = []
        self.active_source_state_path = DATA_SOURCE_ACTIVE_STATE_PATH
        self.active_source_key = "demo"
        self.source_profile = {}
        self.data_mode = "demo"
        self.table_name = "invoices_demo"
        self.internal_api_client = None
        self.chroma = chromadb.Client(Settings(anonymized_telemetry=False))
        self.embed_fn = embedding_functions.OllamaEmbeddingFunction(
            url=f"{OLLAMA_HOST}/api/embeddings",
            model_name=EMBED_MODEL,
        )
        self.collection = self.chroma.get_or_create_collection(
            name="finance_holistic_db",
            embedding_function=self.embed_fn,
        )
        self.df = None
        self.report_context_cache = None
        self.data_contract_summary = build_internal_data_summary(None)
        self.cache_lock = threading.Lock()
        self.refresh_lock = threading.Lock()
        self.sync_status = "not_loaded"
        self.last_sync_started_at = None
        self.last_sync_at = None
        self.last_success_at = None
        self.last_sync_duration_seconds = None
        self.last_sync_error = None
        self.data_version = 0
        self._reload_source_registry()
        self.refresh_data()

    @staticmethod
    def _normalize_records(records):
        data_frame = pd.json_normalize(records, sep="_")
        if data_frame.empty:
            return data_frame

        for column in data_frame.columns:
            data_frame[column] = data_frame[column].apply(
                lambda value: json.dumps(value, ensure_ascii=False)
                if isinstance(value, (dict, list))
                else value
            )

        data_frame.columns = [column.strip() for column in data_frame.columns]
        return data_frame

    @staticmethod
    def _build_table_name(source_key):
        normalized_key = re.sub(r"[^a-z0-9_]+", "_", str(source_key or "demo").strip().lower()).strip("_")
        return f"invoices_{normalized_key or 'demo'}"

    def _reload_source_registry(self):
        profiles, issues, default_key = load_available_source_profiles(
            demo_csv_path=DEMO_CSV_PATH,
            legacy_data_mode=DATA_ACQUISITION_MODE,
            internal_api_endpoint_url=INTERNAL_API_ENDPOINT_URL,
            internal_api_base_url=INTERNAL_API_BASE_URL,
            internal_api_dataset_path=INTERNAL_API_DATASET_PATH,
            demo_profile_path=DATA_SOURCE_DEMO_PROFILE_PATH,
            production_profile_path=DATA_SOURCE_PRODUCTION_PROFILE_PATH,
            config_file_path=INTERNAL_API_CONFIG_FILE,
        )
        self.source_registry = profiles
        self.source_registry_issues = issues
        selected_key, selected_profile = resolve_active_source_profile(
            profiles=profiles,
            state_path=self.active_source_state_path,
            legacy_default_key=default_key,
        )
        self._set_active_source(selected_key, selected_profile, persist=False)

    def _set_active_source(self, source_key, source_profile, persist=True):
        self.active_source_key = source_key
        self.source_profile = copy.deepcopy(source_profile or {})
        self.data_mode = str(self.source_profile.get("mode") or "demo").strip().lower() or "demo"
        self.table_name = self._build_table_name(source_key)
        if self.source_profile.get("type") == "json_api":
            self.internal_api_client = InternalAPIClient(source_profile=self.source_profile)
        else:
            self.internal_api_client = None
        if persist:
            write_active_source_key(self.active_source_state_path, source_key)

    def _load_demo_data(self, profile=None):
        active_profile = profile or self.source_profile
        csv_path = str((active_profile or {}).get("path") or DEMO_CSV_PATH)
        try:
            data_frame = pd.read_sql(f"SELECT * FROM {self.table_name}", self.engine)
        except Exception:
            if not os.path.exists(csv_path):
                raise FileNotFoundError(f"Demo CSV source is unavailable: {csv_path}")

            raw_df = pd.read_csv(csv_path)
            raw_df.columns = [column.strip() for column in raw_df.columns]
            data_frame = raw_df

        normalized_df, data_summary = normalize_financial_dataframe(data_frame)
        return normalized_df, data_summary

    def _load_internal_api_data(self, profile=None):
        client = InternalAPIClient(source_profile=profile or self.source_profile)
        if not client.is_configured():
            raise RuntimeError("Internal data source is not configured.")

        records, extraction_summary = client.fetch_records()

        raw_data_frame = self._normalize_records(records)
        if raw_data_frame.empty:
            raise RuntimeError("Internal data source returned no records.")

        normalized_df, _ = normalize_financial_dataframe(
            raw_data_frame,
            explicit_field_map=client.field_map,
        )
        data_summary = build_internal_data_summary(
            normalized_df,
            explicit_field_map=client.field_map,
            extraction_summary=extraction_summary,
        )

        if data_summary["missingRequiredFields"]:
            logger.warning(
                "Internal data source is missing required fields after normalization: %s",
                ", ".join(data_summary["missingRequiredFields"]),
            )

        return normalized_df, data_summary

    def _load_source_data(self, profile=None):
        active_profile = profile or self.source_profile
        source_type = str((active_profile or {}).get("type") or "demo_csv").strip().lower()
        if source_type == "json_api":
            return self._load_internal_api_data(profile=active_profile)
        return self._load_demo_data(profile=active_profile)

    def _rebuild_embeddings(self, data_frame):
        if data_frame is None or data_frame.empty:
            return False

        existing_ids = self.collection.get().get("ids", [])
        if existing_ids:
            self.collection.delete(ids=existing_ids)

        ids = []
        documents = []
        metadatas = []

        for index, row in data_frame.iterrows():
            text_representation = " | ".join(
                f"{column}: {value}" for column, value in row.items()
            )
            ids.append(str(index))
            documents.append(text_representation)
            metadatas.append(row.astype(str).to_dict())

        if not ids:
            return False

        try:
            logger.info(
                "Syncing %s financial records to the embedding store.",
                len(ids),
            )
            self.collection.add(
                documents=documents,
                metadatas=metadatas,
                ids=ids,
            )
        except Exception as exc:
            logger.error("Embedding sync failed: %s", exc)
            return False

        return True

    def refresh_data(self):
        with self.refresh_lock:
            started_at = datetime.now()
            self.sync_status = "refreshing"
            self.last_sync_started_at = started_at
            self.last_sync_error = None

            try:
                loaded_df, loaded_summary = self._load_source_data()
            except Exception as exc:
                completed_at = datetime.now()
                self.sync_status = "error"
                self.last_sync_at = completed_at
                self.last_sync_duration_seconds = round((completed_at - started_at).total_seconds(), 2)
                self.last_sync_error = str(exc)
                return False

            rebuilt = self._rebuild_embeddings(loaded_df)
            completed_at = datetime.now()
            self.df = loaded_df
            self.data_contract_summary = loaded_summary
            self.df.to_sql(self.table_name, self.engine, index=False, if_exists="replace")
            with self.cache_lock:
                self.report_context_cache = None
            self.sync_status = "ready" if rebuilt else "degraded"
            self.last_sync_at = completed_at
            self.last_success_at = completed_at
            self.last_sync_duration_seconds = round((completed_at - started_at).total_seconds(), 2)
            self.last_sync_error = None if rebuilt else "Embedding store gagal diperbarui. Dashboard tetap memakai data finansial terbaru."
            self.data_version += 1
            return True

    def validate_source(self, source_key):
        self._reload_source_registry()
        profile = self.source_registry.get(source_key)
        if not profile:
            raise ValueError(f"Sumber data `{source_key}` tidak tersedia.")

        summary = summarize_source_profile(profile)
        validation = {
            "source": summary,
            "ready": False,
            "message": "",
            "recordCount": None,
            "missingRequiredFields": [],
            "contractSummary": None,
        }

        try:
            data_frame, data_summary = self._load_source_data(profile=profile)
        except Exception as exc:
            validation["message"] = str(exc)
            return validation

        validation["ready"] = bool(data_summary.get("isReady"))
        validation["message"] = "Sumber data valid dan siap diaktifkan." if validation["ready"] else "Sumber data terbaca, tetapi field wajib belum lengkap."
        validation["recordCount"] = int(len(data_frame))
        validation["missingRequiredFields"] = list(data_summary.get("missingRequiredFields") or [])
        validation["contractSummary"] = data_summary
        return validation

    def activate_source(self, source_key):
        validation = self.validate_source(source_key)
        if not validation["ready"]:
            return {**validation, "activated": False}

        profile = self.source_registry[source_key]
        previous_key = self.active_source_key
        previous_profile = copy.deepcopy(self.source_profile)
        self._set_active_source(source_key, profile, persist=False)
        if self.refresh_data():
            write_active_source_key(self.active_source_state_path, source_key)
            return {**validation, "activated": True, "activeSourceKey": source_key}

        error_message = self.last_sync_error or "Aktivasi gagal."
        self._set_active_source(previous_key, previous_profile, persist=False)
        self.refresh_data()
        return {
            **validation,
            "activated": False,
            "message": error_message,
            "activeSourceKey": previous_key,
        }

    def query(self, context_keywords="", max_results=12):
        query_text = (
            "Historical invoice delays, payment behavior class A-E, "
            "systemic financial risk, collection bottlenecks. "
            f"{context_keywords or ''}"
        )
        if self.df is None or self.df.empty:
            return "Tidak ada data finansial internal yang dapat dipakai."

        max_results = min(max_results, len(self.df))
        collection_size = self.collection.count()
        if collection_size <= 0:
            return "Tidak ada data finansial internal yang dapat dipakai."
        max_results = min(max_results, collection_size)

        try:
            result = self.collection.query(query_texts=[query_text], n_results=max_results)
            documents = result.get("documents", [])
            if documents and documents[0]:
                return "\n---\n".join(documents[0])
        except Exception as exc:
            logger.error("Query error: %s", exc)

        return "Tidak ada data finansial internal yang dapat dipakai."

    def get_report_context(self, notes=""):
        with self.cache_lock:
            if self.report_context_cache is None:
                self.report_context_cache = FinancialAnalyzer.build_report_context(
                    self.df,
                    data_mode=self.data_mode,
                )

        context = copy.deepcopy(self.report_context_cache)
        notes = (notes or "").strip()
        if notes:
            focused_evidence = self.query(notes, max_results=10) or context["evidence"]
            context["evidence"] = FinancialAnalyzer.normalize_evidence_text(focused_evidence)
        context.update(FinancialAnalyzer.apply_silent_assessment(context, notes))
        return context

    def get_review_context(self):
        report_context = self.get_report_context("")
        return report_context.get("review_context", {})

    def get_internal_data_contract(self):
        contract = get_internal_api_contract()
        contract["currentSummary"] = self.data_contract_summary
        contract["dataMode"] = self.data_mode
        contract["internalApiConfigured"] = bool(self.internal_api_client and self.internal_api_client.is_configured())
        contract["datasetUrl"] = self.internal_api_client.get_dataset_url() if self.internal_api_client and self.internal_api_client.is_configured() else None
        contract["activeSourceKey"] = self.active_source_key
        contract["activeSource"] = summarize_source_profile(self.source_profile)
        contract["availableSources"] = [
            summarize_source_profile(profile)
            for _, profile in sorted(self.source_registry.items(), key=lambda item: item[0])
        ]
        contract["registryIssues"] = list(self.source_registry_issues)
        return contract

    def get_sync_status(self, refresh_interval_seconds=0):
        next_refresh_at = None
        if refresh_interval_seconds > 0 and self.last_success_at is not None:
            next_refresh_at = self.last_success_at + timedelta(seconds=refresh_interval_seconds)

        source_age_minutes = None
        if self.last_success_at is not None:
            source_age_minutes = round((datetime.now() - self.last_success_at).total_seconds() / 60, 1)

        return {
            "dataMode": self.data_mode,
            "syncStatus": self.sync_status,
            "lastSyncStartedAt": self.last_sync_started_at.isoformat() if self.last_sync_started_at else None,
            "lastSyncAt": self.last_sync_at.isoformat() if self.last_sync_at else None,
            "lastSuccessAt": self.last_success_at.isoformat() if self.last_success_at else None,
            "lastSyncDurationSeconds": self.last_sync_duration_seconds,
            "lastSyncError": self.last_sync_error,
            "sourceAgeMinutes": source_age_minutes,
            "nextRefreshAt": next_refresh_at.isoformat() if next_refresh_at else None,
            "recordCount": 0 if self.df is None else int(len(self.df)),
            "dataVersion": self.data_version,
            "contractReady": bool(self.data_contract_summary.get("isReady")),
            "activeSourceKey": self.active_source_key,
            "activeSource": summarize_source_profile(self.source_profile),
            "availableSources": [
                summarize_source_profile(profile)
                for _, profile in sorted(self.source_registry.items(), key=lambda item: item[0])
            ],
            "sourceRegistryIssues": list(self.source_registry_issues),
        }


class FinancialAnalyzer:
    CLASS_SCORE_MAP = {
        "kelas a": 1,
        "kelas b": 2,
        "kelas c": 3,
        "kelas d": 4,
        "kelas e": 5,
    }
    PAYMENT_CLASS_ORDER = {
        "Kelas A": 1,
        "Kelas B": 2,
        "Kelas C": 3,
        "Kelas D": 4,
        "Kelas E": 5,
        "Tidak Diketahui": 6,
    }
    REALIZATION_RATE_MAP = {
        "Kelas A": 0.98,
        "Kelas B": 0.90,
        "Kelas C": 0.72,
        "Kelas D": 0.45,
        "Kelas E": 0.18,
        "Tidak Diketahui": 0.65,
    }
    UPSIDE_RATE_MAP = {
        "Kelas A": 1.00,
        "Kelas B": 0.95,
        "Kelas C": 0.82,
        "Kelas D": 0.60,
        "Kelas E": 0.28,
        "Tidak Diketahui": 0.72,
    }
    DOWNSIDE_RATE_MAP = {
        "Kelas A": 0.95,
        "Kelas B": 0.82,
        "Kelas C": 0.60,
        "Kelas D": 0.30,
        "Kelas E": 0.05,
        "Tidak Diketahui": 0.50,
    }
    CONFIDENCE_LABELS = {
        1: "rendah",
        2: "rendah",
        3: "menengah",
        4: "tinggi",
        5: "tinggi",
    }
    OWNERSHIP_KEYWORDS = (
        "owner",
        "pic",
        "account manager",
        "finance manager",
        "finance",
        "collection",
        "cfo",
        "direktur",
        "direksi",
        "sponsor",
        "project lead",
        "project manager",
        "kepala",
        "manajer",
        "tim",
        "business translator",
        "ai engineer",
        "it",
    )
    ADOPTION_KEYWORDS = (
        "pilot",
        "uat",
        "testing",
        "rollout",
        "implementasi",
        "adopsi",
        "pelatihan",
        "sosialisasi",
        "workshop",
        "change management",
        "manajemen",
        "governance",
        "kontrol",
        "roadmap",
    )
    CAUTION_KEYWORDS = (
        "asumsi",
        "batasan",
        "risiko",
        "kontrol",
        "governance",
        "pilot",
        "owner",
        "pic",
        "siap",
        "implementasi",
    )

    DELAY_THEME_KEYWORDS = {
        "Siklus anggaran": ("dipa", "anggaran", "apbn", "apbd", "pencairan", "budget"),
        "Persetujuan internal klien": ("direksi", "approval", "persetujuan", "otorisasi", "tanda tangan"),
        "Dokumen dan administrasi": ("bast", "dokumen", "administrasi", "berita acara", "po", "invoice revisi"),
        "Likuiditas pelanggan": ("cashflow", "likuiditas", "arus kas", "pending dana", "menunggu dana"),
        "Sengketa atau klarifikasi": ("dispute", "klarifikasi", "komplain", "revisi", "sengketa"),
    }
    THEME_ACTION_MAP = {
        "Siklus anggaran": {
            "action": "Validasi posisi anggaran dan jadwal pencairan dengan PIC klien, lalu siapkan eskalasi sebelum cut-off termin.",
            "impact": "Memperjelas kapan invoice realistis dapat direalisasikan dan menurunkan slip karena siklus anggaran.",
            "owner": "Finance Collection + Account Manager",
        },
        "Persetujuan internal klien": {
            "action": "Kunci jalur approval, percepat BAST/BA final, dan pastikan sponsor internal klien ikut mendorong sign-off.",
            "impact": "Memangkas waktu tunggu approval yang selama ini menahan pelepasan pembayaran.",
            "owner": "Account Manager + Project Lead",
        },
        "Dokumen dan administrasi": {
            "action": "Selesaikan kelengkapan dokumen, revisi invoice/PO, dan lakukan quality check sebelum follow-up penagihan berikutnya.",
            "impact": "Mengurangi alasan administratif yang membuat invoice technically ready tetapi belum bisa dibayar.",
            "owner": "Project Admin + Finance Collection",
        },
        "Likuiditas pelanggan": {
            "action": "Negosiasikan skema termin, minta komitmen tanggal bayar tertulis, dan siapkan eskalasi manajemen untuk akun rentan.",
            "impact": "Meningkatkan peluang pemulihan arus kas masuk pada akun yang menghadapi tekanan likuiditas.",
            "owner": "Finance Manager + Account Manager",
        },
        "Sengketa atau klarifikasi": {
            "action": "Tutup ruang lingkup yang disengketakan, dokumentasikan keputusan bersama, dan tahan ekspansi pekerjaan sampai dispute selesai.",
            "impact": "Membuka blokir pembayaran yang tertahan karena ketidakjelasan ruang lingkup atau deliverable.",
            "owner": "Project Lead + Finance Manager",
        },
        "Follow-up umum": {
            "action": "Lakukan follow-up berbasis bukti dengan PIC AP/finance dan pastikan komitmen bayar terbaru tercatat.",
            "impact": "Menjaga ritme collection pada akun yang belum menunjukkan isu dominan tertentu.",
            "owner": "Finance Collection",
        },
    }

    COLUMN_ALIASES = {
        "period": ("periode laporan", "period", "report period", "invoice period"),
        "partner": ("tipe partner", "partner type", "partner_type", "customer segment", "segment"),
        "service": ("layanan", "service", "product", "offering"),
        "payment_class": ("kelas pembayaran", "payment class", "payment_class", "collection class"),
        "invoice_value": ("nilai invoice", "invoice value", "invoice_amount", "amount", "outstanding"),
        "notes": (
            "catatan historis keterlambatan",
            "collection note",
            "collection notes",
            "notes",
            "delay note",
        ),
    }

    @staticmethod
    def _normalize_column_name(column_name):
        return re.sub(r"[^a-z0-9]+", " ", str(column_name).strip().lower()).strip()

    @classmethod
    def _find_column(cls, df, alias_key):
        normalized_map = {
            cls._normalize_column_name(column): column
            for column in df.columns
        }
        for candidate in cls.COLUMN_ALIASES[alias_key]:
            column = normalized_map.get(candidate)
            if column:
                return column
        return None

    @staticmethod
    def _parse_currency(value):
        return parse_idr_amount(value)

    @classmethod
    def _detect_payment_class(cls, raw_value):
        text = str(raw_value or "").lower()
        for class_name, score in cls.CLASS_SCORE_MAP.items():
            if class_name in text:
                return class_name.upper().replace("KELAS", "Kelas"), score
        return "Tidak Diketahui", 3

    @staticmethod
    def _format_currency(amount):
        return f"Rp {amount:,.0f}".replace(",", ".")

    @staticmethod
    def _format_percentage(value):
        return f"{value:.1f}%"

    @staticmethod
    def _normalize_report_text_fragment(text):
        normalized = re.sub(r"\s+", " ", str(text or "")).strip()
        normalized = normalized.replace("…", " ")
        normalized = re.sub(r"\.{3,}", " ", normalized)
        normalized = re.sub(r"\s+([,.;:!?])", r"\1", normalized)
        normalized = re.sub(r"[\"'`]+$", "", normalized).strip(" -–—")
        return normalized.strip()

    @classmethod
    def _trim_note_for_report(cls, note, max_length=220):
        normalized = cls._normalize_report_text_fragment(note)
        if not normalized:
            return ""
        if len(normalized) <= max_length:
            return normalized

        candidate = normalized[: max_length + 1]
        sentence_breaks = [candidate.rfind(marker) for marker in (".", "!", "?", ";", ":")]
        best_sentence_break = max(sentence_breaks)
        if best_sentence_break >= int(max_length * 0.55):
            candidate = candidate[: best_sentence_break + 1]
        else:
            last_space = candidate.rfind(" ")
            if last_space > int(max_length * 0.55):
                candidate = candidate[:last_space]
            else:
                candidate = candidate[:max_length]

        return candidate.strip(" ,;:-")

    @staticmethod
    def _parse_period_sort_key(period_label):
        text = str(period_label or "").strip().lower()
        match = re.search(r"q([1-4])\s*(20\d{2})", text)
        if match:
            return int(match.group(2)), int(match.group(1))
        year_match = re.search(r"(20\d{2})", text)
        if year_match:
            return int(year_match.group(1)), 0
        return 0, 0

    @classmethod
    def _extract_delay_themes(cls, notes_series):
        counts = {theme: 0 for theme in cls.DELAY_THEME_KEYWORDS}
        for note in notes_series.dropna().astype(str):
            lowered_note = note.lower()
            for theme, keywords in cls.DELAY_THEME_KEYWORDS.items():
                if any(keyword in lowered_note for keyword in keywords):
                    counts[theme] += 1

        ranked = [(theme, count) for theme, count in counts.items() if count > 0]
        ranked.sort(key=lambda item: item[1], reverse=True)
        return ranked[:4]

    @classmethod
    def _detect_note_themes(cls, note):
        matched_themes = []
        lowered_note = str(note or "").lower()
        for theme, keywords in cls.DELAY_THEME_KEYWORDS.items():
            if any(keyword in lowered_note for keyword in keywords):
                matched_themes.append(theme)
        return matched_themes

    @classmethod
    def _get_action_plan(cls, note):
        themes = cls._detect_note_themes(note)
        selected_theme = themes[0] if themes else "Follow-up umum"
        action_plan = cls.THEME_ACTION_MAP.get(selected_theme, cls.THEME_ACTION_MAP["Follow-up umum"])
        return selected_theme, action_plan

    @staticmethod
    def _format_evidence_chunk(chunk):
        normalized = re.sub(r"\s+", " ", str(chunk or "")).strip()
        if not normalized:
            return ""

        pattern = re.compile(
            r"Periode Laporan:\s*(?P<period>.*?)\s*\|\s*"
            r"Tipe Partner:\s*(?P<partner>.*?)\s*\|\s*"
            r"Layanan:\s*(?P<service>.*?)\s*\|\s*"
            r"Kelas Pembayaran:\s*(?P<payment_class>.*?)\s*\|\s*"
            r"Nilai Invoice:\s*(?P<invoice_value>.*?)\s*\|\s*"
            r"Catatan Historis Keterlambatan:\s*(?P<note>.*)$"
        )
        match = pattern.match(normalized)
        if not match:
            return f"- {normalized}"

        parts = match.groupdict()
        return "\n".join(
            [
                f"- {parts['period']} | {parts['partner']} | {parts['service']}",
                f"  - Kelas pembayaran: {parts['payment_class']}",
                f"  - Nilai invoice: {parts['invoice_value']}",
                f"  - Catatan utama: {parts['note']}",
            ]
        )

    @staticmethod
    def normalize_evidence_text(raw_text):
        lines = []
        chunks = [chunk.strip() for chunk in str(raw_text or "").split("\n---\n") if chunk.strip()]

        if len(chunks) > 1 or any("Periode Laporan:" in chunk for chunk in chunks):
            for chunk in chunks[:10]:
                lines.append(FinancialAnalyzer._format_evidence_chunk(chunk))
            return "\n".join(lines) if lines else "- Tidak ada catatan historis yang cukup untuk dikutip."

        for raw_line in str(raw_text or "").splitlines():
            cleaned_line = raw_line.strip()
            if not cleaned_line or cleaned_line == "---":
                continue
            if cleaned_line.startswith("- ") or re.match(r"^\d+\.", cleaned_line):
                lines.append(cleaned_line)
            else:
                lines.append(f"- {cleaned_line}")

        return "\n".join(lines) if lines else "- Tidak ada catatan historis yang cukup untuk dikutip."

    @staticmethod
    def _clamp_score(score):
        return max(1, min(5, int(round(score))))

    @classmethod
    def _score_to_confidence(cls, score):
        return cls.CONFIDENCE_LABELS.get(cls._clamp_score(score), "menengah")

    @staticmethod
    def _format_list_as_sentence(items):
        filtered_items = [str(item).strip() for item in items if str(item).strip()]
        if not filtered_items:
            return "-"
        if len(filtered_items) == 1:
            return filtered_items[0]
        if len(filtered_items) == 2:
            return f"{filtered_items[0]} dan {filtered_items[1]}"
        return f"{', '.join(filtered_items[:-1])}, dan {filtered_items[-1]}"

    @classmethod
    def _count_keyword_hits(cls, raw_text, keywords):
        lowered_text = str(raw_text or "").lower()
        return sum(1 for keyword in keywords if keyword in lowered_text)

    @classmethod
    def _build_hidden_dimensions(cls, base_profile, notes):
        notes = (notes or "").strip()
        missing_core_fields = len(base_profile.get("missing_core_fields", []))
        total_invoices = base_profile.get("total_invoices", 0)
        data_mode = base_profile.get("data_mode", "demo")
        owner_hits = cls._count_keyword_hits(notes, cls.OWNERSHIP_KEYWORDS)
        adoption_hits = cls._count_keyword_hits(notes, cls.ADOPTION_KEYWORDS)

        business_value_score = 4 if total_invoices > 0 and base_profile.get("high_risk_invoices", 0) > 0 else 3
        if notes:
            business_value_score += 1

        data_model_score = 4 if missing_core_fields <= 1 and total_invoices >= 25 else 3
        if data_mode == "demo":
            data_model_score -= 1

        infrastructure_score = 4 if APP_SERVER == "waitress" else 2
        if REPORT_MAX_CONCURRENT_JOBS >= 4:
            infrastructure_score += 1
        if APP_SERVER == "waitress" and WAITRESS_THREADS < 8:
            infrastructure_score -= 1

        people_score = 2
        if owner_hits >= 3:
            people_score = 4
        elif owner_hits >= 1:
            people_score = 3

        governance_score = 4
        if data_mode == "demo":
            governance_score -= 1
        if missing_core_fields >= 2:
            governance_score -= 1

        adoption_score = 2
        if adoption_hits >= 3:
            adoption_score = 4
        elif adoption_hits >= 1 or owner_hits >= 2:
            adoption_score = 3

        dimensions = {
            "business_value_clarity": cls._clamp_score(business_value_score),
            "data_model_readiness": cls._clamp_score(data_model_score),
            "infrastructure_readiness": cls._clamp_score(infrastructure_score),
            "people_ownership_readiness": cls._clamp_score(people_score),
            "governance_control_readiness": cls._clamp_score(governance_score),
            "organizational_adoption_readiness": cls._clamp_score(adoption_score),
        }
        note_profile = {
            "owner_hits": owner_hits,
            "adoption_hits": adoption_hits,
            "has_caution_prompt": bool(cls._count_keyword_hits(notes, cls.CAUTION_KEYWORDS)),
        }
        return dimensions, note_profile

    @classmethod
    def _build_readiness_outputs(cls, base_profile, hidden_dimensions, note_profile):
        data_mode = base_profile.get("data_mode", "demo")
        total_invoices = base_profile.get("total_invoices", 0)
        data_source_label = "dataset demo lokal" if data_mode == "demo" else "API internal perusahaan"
        core_field_summary = (
            f"{base_profile.get('core_fields_available', 0)}/{base_profile.get('core_fields_expected', 6)} atribut inti tersedia"
        )
        missing_fields_sentence = cls._format_list_as_sentence(base_profile.get("missing_core_fields", []))
        partner_focus_sentence = cls._format_list_as_sentence(base_profile.get("top_risk_partners", []))

        data_confidence = cls._score_to_confidence(hidden_dimensions["data_model_readiness"])
        deployment_confidence = cls._score_to_confidence(hidden_dimensions["infrastructure_readiness"])
        ownership_confidence = cls._score_to_confidence(hidden_dimensions["people_ownership_readiness"])
        control_confidence = cls._score_to_confidence(hidden_dimensions["governance_control_readiness"])
        adoption_confidence = cls._score_to_confidence(hidden_dimensions["organizational_adoption_readiness"])
        expected_gap_base = base_profile.get("expected_gap_base", 0)

        if data_mode == "demo":
            data_mode_line = (
                "Data masih berasal dari dataset demo lokal, sehingga laporan cocok untuk simulasi diskusi internal "
                "dan pengujian alur, bukan untuk komitmen operasional final."
            )
        else:
            data_mode_line = (
                "Data sudah ditarik dari API internal, sehingga laporan lebih layak dipakai sebagai dasar prioritas operasional "
                "selama sinkronisasi dan validasi sumber tetap terjaga."
            )

        readiness_signals = [
            f"- Business value clarity: {hidden_dimensions['business_value_clarity']}/5. Use case cashflow jelas dan langsung terkait percepatan realisasi invoice, pengendalian cash out, pengurangan risiko keterlambatan, dan prioritas follow-up.",
            f"- Data/model readiness: {hidden_dimensions['data_model_readiness']}/5. Sumber data saat ini adalah {data_source_label} dengan {total_invoices} invoice dan {core_field_summary}.",
            f"- Infrastructure/deployment readiness: {hidden_dimensions['infrastructure_readiness']}/5. Runtime aktif adalah {APP_SERVER} dengan queue {REPORT_MAX_CONCURRENT_JOBS} job dan thread Waitress {WAITRESS_THREADS}.",
            f"- People/ownership readiness: {hidden_dimensions['people_ownership_readiness']}/5. Sinyal owner atau sponsor eksplisit dari catatan pengguna = {note_profile['owner_hits']}.",
            f"- Governance/risk control: {hidden_dimensions['governance_control_readiness']}/5. Internal data tetap menjadi sumber fakta utama, OSINT hanya pendukung, dan fallback menjaga konsistensi struktur laporan.",
            f"- Organizational adoption readiness: {hidden_dimensions['organizational_adoption_readiness']}/5. Sinyal kesiapan adopsi atau pilot dari catatan pengguna = {note_profile['adoption_hits']}.",
            "- Gunakan sinyal ini untuk mengatur tingkat keyakinan, caveat, risiko kontrol, kepemilikan tindakan, dan prasyarat implementasi secara halus tanpa menyebut kerangka internal apa pun.",
        ]

        confidence_lines = [
            f"- Tingkat keyakinan data saat ini {data_confidence} karena laporan memakai {data_source_label} dengan {core_field_summary}.",
            f"- Kesiapan operasional aplikasi berada pada tingkat {deployment_confidence}; jalur deployment saat ini {APP_SERVER} dengan dukungan antrean {REPORT_MAX_CONCURRENT_JOBS} pekerjaan paralel.",
            f"- Kepastian penanggung jawab lintas fungsi berada pada tingkat {ownership_confidence}; sinyal owner eksplisit yang tertangkap dari catatan pengguna = {note_profile['owner_hits']}.",
            data_mode_line,
        ]

        assumption_lines = [
            f"- Analisis ini mengasumsikan {total_invoices} invoice yang tersedia sudah mewakili pola penagihan utama pada portofolio yang dibahas.",
            f"- Atribut inti yang terbaca saat ini adalah {core_field_summary}; atribut yang belum kuat atau belum tersedia: {missing_fields_sentence}.",
            "- Proyeksi arus kas masuk dibaca sebagai skenario manajemen berbasis histori kelas pembayaran, bukan kepastian realisasi kas.",
        ]
        if data_mode == "demo":
            assumption_lines.append("- Karena masih demo mode, angka dan narasi diposisikan sebagai bahan kalibrasi diskusi sebelum integrasi source-of-truth internal.")
        if note_profile["owner_hits"] == 0:
            assumption_lines.append("- Catatan pengguna belum menyebut penanggung jawab spesifik, sehingga penetapan owner tindakan masih perlu divalidasi saat rapat.")

        control_lines = [
            f"- Postur kontrol saat ini berada pada tingkat {control_confidence}; invoice prioritas tetap memerlukan verifikasi manual sebelum komitmen eskalasi atau forecast dinaikkan.",
            f"- Fokus kontrol utama saat ini berada pada konsentrasi risiko di {partner_focus_sentence} dan invoice Kelas D/E yang bernilai besar.",
            "- Sebelum eskalasi ke klien, verifikasi ulang invoice prioritas, kelengkapan dokumen, dan status komitmen bayar terbaru.",
            "- OSINT hanya dipakai sebagai konteks eksternal; fakta invoice, nilai, dan prioritas tetap harus mengikuti data internal yang tersedia.",
        ]
        if data_mode == "demo":
            control_lines.append("- Hindari menjadikan hasil demo sebagai dasar komitmen eksternal atau forecast final tanpa verifikasi terhadap sistem internal.")
        if hidden_dimensions["governance_control_readiness"] <= 2:
            control_lines.append("- Risiko salah tafsir naik bila data tambahan dan kontrol review manual tidak disiapkan sebelum sesi tindak lanjut.")

        implementation_lines = [
            f"- Untuk penggunaan operasional yang lebih kuat, pertahankan minimal {REPORT_MAX_CONCURRENT_JOBS} slot antrean dan gunakan runtime {APP_SERVER if APP_SERVER == 'waitress' else 'Waitress pada uji bersama'} untuk akses internal bersama.",
            "- Pastikan ritme refresh data, sumber invoice prioritas, dan jalur eskalasi ke account owner diputuskan sebelum laporan dipakai sebagai dasar action plan mingguan.",
        ]
        if data_mode == "demo":
            implementation_lines.append("- Langkah implementasi terdekat adalah memetakan endpoint API internal, autentikasi, dan struktur data source-of-truth agar prioritas penagihan tidak lagi bergantung pada dataset simulasi.")
        else:
            implementation_lines.append("- Pertahankan monitoring sinkronisasi API, validasi schema, dan fallback dokumen agar kualitas laporan tetap stabil saat dipakai banyak pengguna.")

        organizational_lines = [
            f"- Kesiapan pelaksanaan saat ini berada pada tingkat {adoption_confidence}; laporan akan lebih mudah dieksekusi bila sponsor bisnis, finance collection, dan account owner duduk pada forum review yang sama.",
            "- Gunakan laporan ini sebagai bahan keputusan internal: apa yang harus ditagih lebih dulu, siapa yang memimpin eskalasi, dan kontrol apa yang wajib ditutup sebelum follow-up berikutnya.",
        ]
        if note_profile["owner_hits"] == 0:
            organizational_lines.append("- Tetapkan minimal satu owner bisnis dan satu owner operasional untuk setiap cluster invoice prioritas agar keputusan rapat langsung bisa dijalankan.")
        if note_profile["adoption_hits"] == 0:
            organizational_lines.append("- Siapkan ritme pilot atau review berkala agar penggunaan laporan tidak berhenti di tahap analisis saja.")

        review_context = {
            "dataSource": "Demo dataset lokal" if data_mode == "demo" else "API internal perusahaan",
            "dataStatus": f"{total_invoices} invoice siap dianalisis",
            "operationalScope": "Analisis cashflow (arus kas masuk dan arus kas keluar), risiko realisasi invoice, dan prioritas tindak lanjut 30 hari.",
            "reportPurpose": "Bahan diskusi internal manajemen untuk keputusan penagihan, kontrol risiko, dan kesiapan pelaksanaan.",
            "readinessCaveat": data_mode_line,
            "controlNote": "Fakta internal tetap menjadi sumber utama; konteks eksternal hanya dipakai untuk memperkaya pembacaan risiko.",
        }
        cash_plan_implications = [
            "- Base case sebaiknya dipakai sebagai jangkar pembacaan rencana kas jangka pendek, sedangkan upside dan downside dipakai untuk menguji kebutuhan eskalasi dan ruang koreksi target.",
            "- Semakin besar eksposur Kelas D/E pada partner bernilai tinggi, semakin besar kebutuhan buffer keputusan, ritme follow-up, dan verifikasi dokumen sebelum asumsi arus kas masuk dinaikkan.",
        ]
        if data_mode == "demo":
            cash_plan_implications.append("- Karena masih demo mode, implikasi rencana kas diposisikan sebagai arah diskusi internal, bukan angka forecast final.")
        if expected_gap_base > 0:
            cash_plan_implications.append("- Gap arus kas masuk pada base case harus dibaca sebagai ruang risiko yang perlu diperkecil lewat penagihan prioritas, bukan langsung diasumsikan akan pulih otomatis.")
        cash_plan_implications.append("- Tekanan cash out harus dibaca bersama ending cash, runway, dan coverage ratio; kenaikan arus kas masuk saja tidak cukup bila kewajiban jatuh tempo menumpuk pada horizon yang sama.")

        return {
            "readiness_signals": "\n".join(readiness_signals),
            "confidence_summary": "\n".join(confidence_lines),
            "assumptions": "\n".join(assumption_lines),
            "controls": "\n".join(control_lines),
            "implementation_prerequisites": "\n".join(implementation_lines),
            "organizational_readiness": "\n".join(organizational_lines),
            "cash_plan_implications": "\n".join(cash_plan_implications),
            "review_context": review_context,
        }

    @classmethod
    def _build_visual_prompt(cls, class_distribution):
        labels = []
        for payment_class, metrics in class_distribution.items():
            labels.append(f"{payment_class},{round(metrics['share'], 1)}")
        chart_marker = (
            "[[CHART: Distribusi Historis Kelas Pembayaran | Persentase Invoice | "
            + "; ".join(labels)
            + "]]"
        )
        flow_marker = (
            "[[FLOW: Prioritisasi Invoice Risiko Tinggi -> Penagihan Berbasis Bukti -> "
            "Eskalasi Manajemen -> Pemulihan Cashflow]]"
        )
        return f"{chart_marker}\n{flow_marker}"

    @classmethod
    def apply_silent_assessment(cls, context, notes=""):
        base_profile = context.get("base_profile", {})
        hidden_dimensions, note_profile = cls._build_hidden_dimensions(base_profile, notes)
        visible_outputs = cls._build_readiness_outputs(base_profile, hidden_dimensions, note_profile)
        return {
            **visible_outputs,
            "hidden_dimensions": hidden_dimensions,
            "note_profile": note_profile,
        }

    @classmethod
    def build_report_context(cls, df, data_mode="demo"):
        if df is None or df.empty:
            return {
                "financial_summary": "Tidak ada data finansial internal yang tersedia.",
                "evidence": "Tidak ada catatan historis yang tersedia.",
                "diagnostic_breakdown": "- Belum ada pola hambatan yang dapat dijelaskan karena data masih kosong.",
                "management_brief": "Tidak ada management brief yang dapat disusun dari data kosong.",
                "executive_facts": "- Tidak ada fakta eksekutif yang tersedia.",
                "scenario_table": "| Skenario | Estimasi Arus Kas Masuk | Gap terhadap Total Invoice | Narasi Manajemen |\n|---|---:|---:|---|\n| Base Case | Rp 0 | Rp 0 | Data kosong. |",
                "priority_table": "| Prioritas | Fokus | Penanggung Jawab | Isu Utama | Aksi 30 Hari | Dampak yang Diharapkan |\n|---:|---|---|---|---|---|\n| 1 | Tidak ada data | Finance Collection | - | Lengkapi data terlebih dahulu. | Memberi dasar analisis yang layak. |",
                "meeting_agenda": "1. Pastikan data internal tersedia sebelum rapat dilanjutkan.",
                "base_profile": {
                    "data_mode": data_mode,
                    "total_invoices": 0,
                    "total_invoice_value": 0,
                    "delayed_invoice_value": 0,
                    "high_risk_invoices": 0,
                    "high_risk_invoice_value": 0,
                    "expected_realization_base": 0,
                    "expected_gap_base": 0,
                    "core_fields_available": 0,
                    "core_fields_expected": 6,
                    "missing_core_fields": [
                        "periode",
                        "partner",
                        "layanan",
                        "kelas pembayaran",
                        "nilai invoice",
                        "catatan keterlambatan",
                    ],
                    "top_risk_partners": [],
                },
                "visual_prompt": "Do not force visuals.",
            }

        working_df = df.copy()
        period_column = cls._find_column(working_df, "period")
        partner_column = cls._find_column(working_df, "partner")
        service_column = cls._find_column(working_df, "service")
        payment_class_column = cls._find_column(working_df, "payment_class")
        invoice_value_column = cls._find_column(working_df, "invoice_value")
        notes_column = cls._find_column(working_df, "notes")
        core_field_map = {
            "periode": period_column,
            "partner": partner_column,
            "layanan": service_column,
            "kelas pembayaran": payment_class_column,
            "nilai invoice": invoice_value_column,
            "catatan keterlambatan": notes_column,
        }
        missing_core_fields = [
            label for label, column in core_field_map.items() if not column
        ]

        working_df["__period"] = (
            working_df[period_column].astype(str).fillna("Tidak Diketahui")
            if period_column
            else "Tidak Diketahui"
        )
        working_df["__partner"] = (
            working_df[partner_column].astype(str).fillna("Tidak Diketahui")
            if partner_column
            else "Tidak Diketahui"
        )
        working_df["__service"] = (
            working_df[service_column].astype(str).fillna("Tidak Diketahui")
            if service_column
            else "Tidak Diketahui"
        )
        working_df["__note"] = (
            working_df[notes_column].astype(str).fillna("")
            if notes_column
            else ""
        )
        working_df["__invoice_value"] = (
            working_df[invoice_value_column].apply(cls._parse_currency)
            if invoice_value_column
            else 0
        )
        class_labels = []
        class_scores = []
        source_series = working_df[payment_class_column] if payment_class_column else pd.Series([""] * len(working_df))
        for raw_value in source_series:
            payment_class, score = cls._detect_payment_class(raw_value)
            class_labels.append(payment_class)
            class_scores.append(score)
        working_df["__payment_class"] = class_labels
        working_df["__payment_score"] = class_scores
        working_df["__base_realization"] = working_df.apply(
            lambda row: row["__invoice_value"] * cls.REALIZATION_RATE_MAP.get(row["__payment_class"], 0.65),
            axis=1,
        )
        working_df["__upside_realization"] = working_df.apply(
            lambda row: row["__invoice_value"] * cls.UPSIDE_RATE_MAP.get(row["__payment_class"], 0.72),
            axis=1,
        )
        working_df["__downside_realization"] = working_df.apply(
            lambda row: row["__invoice_value"] * cls.DOWNSIDE_RATE_MAP.get(row["__payment_class"], 0.50),
            axis=1,
        )

        total_invoices = len(working_df)
        total_invoice_value = int(working_df["__invoice_value"].sum())
        delayed_invoices = int((working_df["__payment_score"] > 1).sum())
        high_risk_invoices = int((working_df["__payment_score"] >= 4).sum())
        delayed_invoice_value = int(working_df.loc[working_df["__payment_score"] > 1, "__invoice_value"].sum())
        high_risk_invoice_value = int(working_df.loc[working_df["__payment_score"] >= 4, "__invoice_value"].sum())
        weighted_risk_score = statistics.fmean(working_df["__payment_score"]) if total_invoices else 0
        expected_realization_base = int(round(working_df["__base_realization"].sum()))
        expected_realization_upside = int(round(working_df["__upside_realization"].sum()))
        expected_realization_downside = int(round(working_df["__downside_realization"].sum()))
        expected_gap_base = max(total_invoice_value - expected_realization_base, 0)
        expected_gap_upside = max(total_invoice_value - expected_realization_upside, 0)
        expected_gap_downside = max(total_invoice_value - expected_realization_downside, 0)

        class_summary_df = (
            working_df.groupby("__payment_class", dropna=False)
            .agg(
                invoice_count=("__payment_class", "size"),
                invoice_value=("__invoice_value", "sum"),
                risk_score=("__payment_score", "mean"),
            )
        )
        class_summary_df["__sort_key"] = [
            cls.PAYMENT_CLASS_ORDER.get(index_value, 99) for index_value in class_summary_df.index
        ]
        class_summary_df = class_summary_df.sort_values("__sort_key")
        class_distribution = {}
        for payment_class, row in class_summary_df.iterrows():
            class_distribution[payment_class] = {
                "count": int(row["invoice_count"]),
                "value": int(row["invoice_value"]),
                "share": (row["invoice_count"] / total_invoices) * 100 if total_invoices else 0,
            }

        partner_summary_df = (
            working_df.groupby("__partner", dropna=False)
            .agg(
                invoice_count=("__partner", "size"),
                invoice_value=("__invoice_value", "sum"),
                avg_risk_score=("__payment_score", "mean"),
            )
            .sort_values(["avg_risk_score", "invoice_value"], ascending=[False, False])
            .head(5)
        )
        service_summary_df = (
            working_df.groupby("__service", dropna=False)
            .agg(
                invoice_count=("__service", "size"),
                invoice_value=("__invoice_value", "sum"),
                avg_risk_score=("__payment_score", "mean"),
            )
            .sort_values(["avg_risk_score", "invoice_value"], ascending=[False, False])
            .head(5)
        )
        period_summary_df = (
            working_df.groupby("__period", dropna=False)
            .agg(
                invoice_count=("__period", "size"),
                invoice_value=("__invoice_value", "sum"),
                avg_risk_score=("__payment_score", "mean"),
            )
            .reset_index()
        )
        period_summary_df["__sort_key"] = period_summary_df["__period"].apply(cls._parse_period_sort_key)
        period_summary_df = period_summary_df.sort_values("__sort_key").tail(6)

        high_risk_subset_df = working_df[working_df["__payment_score"] >= 4].copy()
        high_risk_partner_df = (
            high_risk_subset_df.groupby("__partner", dropna=False)
            .agg(
                invoice_count=("__partner", "size"),
                invoice_value=("__invoice_value", "sum"),
                avg_risk_score=("__payment_score", "mean"),
            )
            .sort_values(["invoice_value", "invoice_count"], ascending=[False, False])
            .head(5)
            if not high_risk_subset_df.empty
            else pd.DataFrame(columns=["invoice_count", "invoice_value", "avg_risk_score"])
        )
        high_risk_service_df = (
            high_risk_subset_df.groupby("__service", dropna=False)
            .agg(
                invoice_count=("__service", "size"),
                invoice_value=("__invoice_value", "sum"),
                avg_risk_score=("__payment_score", "mean"),
            )
            .sort_values(["invoice_value", "invoice_count"], ascending=[False, False])
            .head(5)
            if not high_risk_subset_df.empty
            else pd.DataFrame(columns=["invoice_count", "invoice_value", "avg_risk_score"])
        )

        top_themes = cls._extract_delay_themes(working_df["__note"])
        evidence_rows = working_df.sort_values(
            ["__payment_score", "__invoice_value"],
            ascending=[False, False],
        ).head(8)
        priority_rows = []
        for row_index, (_, row) in enumerate(evidence_rows.iterrows(), start=1):
            theme, action_plan = cls._get_action_plan(row["__note"])
            priority_rows.append(
                {
                    "priority": row_index,
                    "focus": f"{row['__partner']} / {row['__service']}",
                    "issue": theme,
                    "payment_class": row["__payment_class"],
                    "invoice_value": int(row["__invoice_value"]),
                    "action": action_plan["action"],
                    "impact": action_plan["impact"],
                    "owner": action_plan["owner"],
                }
            )

        latest_period_summary = period_summary_df.iloc[-1] if not period_summary_df.empty else None
        previous_period_summary = period_summary_df.iloc[-2] if len(period_summary_df) >= 2 else None

        recent_trend_line = "- Belum ada cukup periode untuk membaca perubahan terbaru."
        recent_value_change_line = "- Perubahan nilai invoice belum dapat dibandingkan antarperiode."
        recent_risk_change_line = "- Perubahan skor risiko belum dapat dibandingkan antarperiode."
        recent_period_label = latest_period_summary["__period"] if latest_period_summary is not None else "periode terbaru"

        if latest_period_summary is not None and previous_period_summary is not None:
            value_delta = int(latest_period_summary["invoice_value"] - previous_period_summary["invoice_value"])
            prior_value = int(previous_period_summary["invoice_value"])
            value_delta_pct = (value_delta / prior_value) * 100 if prior_value else 0
            risk_delta = float(latest_period_summary["avg_risk_score"] - previous_period_summary["avg_risk_score"])
            count_delta = int(latest_period_summary["invoice_count"] - previous_period_summary["invoice_count"])

            value_direction = "naik" if value_delta > 0 else "turun" if value_delta < 0 else "relatif stabil"
            risk_direction = "naik" if risk_delta > 0.05 else "turun" if risk_delta < -0.05 else "relatif stabil"
            risk_interpretation = (
                "memburuk"
                if risk_delta > 0.05
                else "membaik"
                if risk_delta < -0.05
                else "stabil"
            )
            recent_trend_line = (
                f"- Periode terbaru yang diamati adalah {previous_period_summary['__period']} ke {latest_period_summary['__period']} "
                f"dengan jumlah invoice berubah {count_delta:+d} dan nilai invoice {value_direction} "
                f"{abs(value_delta_pct):.1f}%."
            )
            recent_value_change_line = (
                f"- Nilai invoice {previous_period_summary['__period']} = {cls._format_currency(int(previous_period_summary['invoice_value']))}; "
                f"{latest_period_summary['__period']} = {cls._format_currency(int(latest_period_summary['invoice_value']))}."
            )
            recent_risk_change_line = (
                f"- Skor risiko rata-rata {risk_direction} {abs(risk_delta):.2f} poin dari "
                f"{previous_period_summary['avg_risk_score']:.2f} ke {latest_period_summary['avg_risk_score']:.2f}, "
                f"yang berarti kondisi penagihan {risk_interpretation}."
            )

        top_risk_partner_names = ", ".join(high_risk_partner_df.index.tolist()[:3]) if not high_risk_partner_df.empty else "-"
        top_risk_service_names = ", ".join(high_risk_service_df.index.tolist()[:3]) if not high_risk_service_df.empty else "-"
        top_theme_map = {theme: count for theme, count in top_themes}
        process_issue_count = top_theme_map.get("Dokumen dan administrasi", 0) + top_theme_map.get("Sengketa atau klarifikasi", 0)
        budget_issue_count = top_theme_map.get("Siklus anggaran", 0) + top_theme_map.get("Persetujuan internal klien", 0)
        liquidity_issue_count = top_theme_map.get("Likuiditas pelanggan", 0)

        diagnostic_breakdown_lines = []
        if process_issue_count:
            diagnostic_breakdown_lines.append(
                f"1. Hambatan proses, dokumen, dan klarifikasi masih dominan dengan {process_issue_count} sinyal historis; ini biasanya menahan invoice yang sebetulnya sudah siap ditagih tetapi belum lolos kelengkapan atau sign-off."
            )
        if budget_issue_count:
            diagnostic_breakdown_lines.append(
                f"2. Hambatan anggaran dan persetujuan internal klien muncul pada {budget_issue_count} catatan; dampaknya paling terasa pada akun pemerintah, BUMN, dan partner dengan approval berlapis."
            )
        if liquidity_issue_count:
            diagnostic_breakdown_lines.append(
                f"3. Tekanan likuiditas pelanggan muncul pada {liquidity_issue_count} catatan; risiko ini perlu dibedakan dari isu administratif karena membutuhkan pola negosiasi dan komitmen bayar yang lebih aktif."
            )
        if top_risk_partner_names != "-":
            diagnostic_breakdown_lines.append(
                f"4. Eksposur dampak terbesar saat ini terkonsentrasi pada {top_risk_partner_names}, sehingga setiap bottleneck di segmen tersebut memberi pengaruh paling besar ke arus kas masuk dan kualitas ending cash."
            )
        if not diagnostic_breakdown_lines:
            diagnostic_breakdown_lines.append("1. Belum ada pola hambatan dominan yang cukup kuat untuk dipisahkan dari catatan historis.")

        financial_summary_lines = [
            "## Snapshot Arus Kas Masuk",
            f"- Total invoice dianalisis: {total_invoices}",
            f"- Total nilai invoice: {cls._format_currency(total_invoice_value)}",
            f"- Porsi invoice terlambat: {cls._format_percentage((delayed_invoices / total_invoices) * 100 if total_invoices else 0)}",
            f"- Nilai invoice terlambat: {cls._format_currency(delayed_invoice_value)}",
            f"- Porsi invoice risiko tinggi (Kelas D/E): {cls._format_percentage((high_risk_invoices / total_invoices) * 100 if total_invoices else 0)}",
            f"- Nilai invoice risiko tinggi (Kelas D/E): {cls._format_currency(high_risk_invoice_value)}",
            f"- Skor risiko penagihan rata-rata: {weighted_risk_score:.2f} dari 5.00",
            f"- Estimasi arus kas masuk risk-adjusted (base case): {cls._format_currency(expected_realization_base)} atau {cls._format_percentage((expected_realization_base / total_invoice_value) * 100 if total_invoice_value else 0)} dari total nilai invoice",
            f"- Gap arus kas masuk pada base case: {cls._format_currency(expected_gap_base)}",
            "",
            "## Pergerakan Periode Terbaru",
            recent_trend_line,
            recent_value_change_line,
            recent_risk_change_line,
            "",
            "## Distribusi Kelas Pembayaran",
            "| Kelas | Jumlah Invoice | Nilai Invoice | Porsi Invoice |",
            "|---|---:|---:|---:|",
        ]
        for payment_class, metrics in class_distribution.items():
            financial_summary_lines.append(
                f"| {payment_class} | {metrics['count']} | {cls._format_currency(metrics['value'])} | {cls._format_percentage(metrics['share'])} |"
            )

        financial_summary_lines.extend(
            [
                "",
                "## Segmentasi Risiko per Tipe Partner",
                "| Tipe Partner | Jumlah Invoice | Nilai Invoice | Skor Risiko Rata-rata |",
                "|---|---:|---:|---:|",
            ]
        )
        for partner, row in partner_summary_df.iterrows():
            financial_summary_lines.append(
                f"| {partner} | {int(row['invoice_count'])} | {cls._format_currency(int(row['invoice_value']))} | {row['avg_risk_score']:.2f} |"
            )

        financial_summary_lines.extend(
            [
                "",
                "## Segmentasi Risiko per Layanan",
                "| Layanan | Jumlah Invoice | Nilai Invoice | Skor Risiko Rata-rata |",
                "|---|---:|---:|---:|",
            ]
        )
        for service, row in service_summary_df.iterrows():
            financial_summary_lines.append(
                f"| {service} | {int(row['invoice_count'])} | {cls._format_currency(int(row['invoice_value']))} | {row['avg_risk_score']:.2f} |"
            )

        financial_summary_lines.extend(
            [
                "",
                "## Tren Ringkas per Periode",
                "| Periode | Jumlah Invoice | Nilai Invoice | Skor Risiko Rata-rata |",
                "|---|---:|---:|---:|",
            ]
        )
        for _, row in period_summary_df.iterrows():
            financial_summary_lines.append(
                f"| {row['__period']} | {int(row['invoice_count'])} | {cls._format_currency(int(row['invoice_value']))} | {row['avg_risk_score']:.2f} |"
            )

        financial_summary_lines.extend(
            [
                "",
                "## Konsentrasi Eksposur Risiko Tinggi",
                "| Tipe Partner | Jumlah Invoice D/E | Nilai Invoice D/E | Skor Risiko Rata-rata |",
                "|---|---:|---:|---:|",
            ]
        )
        if not high_risk_partner_df.empty:
            for partner, row in high_risk_partner_df.iterrows():
                financial_summary_lines.append(
                    f"| {partner} | {int(row['invoice_count'])} | {cls._format_currency(int(row['invoice_value']))} | {row['avg_risk_score']:.2f} |"
                )
        else:
            financial_summary_lines.append("| Tidak ada konsentrasi D/E | 0 | Rp 0 | 0.00 |")

        financial_summary_lines.extend(["", "## Sinyal Diagnostik Utama"])
        if top_themes:
            for index, (theme, count) in enumerate(top_themes, start=1):
                financial_summary_lines.append(f"{index}. {theme} muncul pada {count} catatan historis.")
        else:
            financial_summary_lines.append("1. Belum ada tema dominan yang dapat diambil dari catatan historis.")

        evidence_lines = []
        for _, row in evidence_rows.iterrows():
            note = str(row["__note"]).strip()
            if not note:
                continue
            trimmed_note = cls._trim_note_for_report(note, max_length=220)
            evidence_lines.append(
                "\n".join(
                    [
                        f"- {row['__period']} | {row['__partner']} | {row['__service']}",
                        f"  - Kelas pembayaran: {row['__payment_class']}",
                        f"  - Nilai invoice: {cls._format_currency(int(row['__invoice_value']))}",
                        f"  - Catatan utama: {trimmed_note}",
                    ]
                )
            )
        if not evidence_lines:
            evidence_lines.append("- Tidak ada catatan historis yang cukup untuk dikutip.")

        executive_fact_lines = [
            f"- Portofolio yang dianalisis mencakup {total_invoices} invoice dengan total nilai {cls._format_currency(total_invoice_value)}.",
            f"- Invoice terlambat mencapai {delayed_invoices} kasus atau {cls._format_percentage((delayed_invoices / total_invoices) * 100 if total_invoices else 0)} dari populasi, dengan nilai tertunda {cls._format_currency(delayed_invoice_value)}.",
            f"- Eksposur risiko tinggi Kelas D/E mencapai {high_risk_invoices} invoice senilai {cls._format_currency(high_risk_invoice_value)} dan terkonsentrasi pada {top_risk_partner_names}.",
            f"- Estimasi arus kas masuk risk-adjusted base case adalah {cls._format_currency(expected_realization_base)} dengan gap {cls._format_currency(expected_gap_base)} terhadap total nilai invoice.",
            recent_trend_line,
            recent_risk_change_line,
            f"- Layanan dengan eksposur risiko tinggi paling dominan saat ini: {top_risk_service_names}.",
        ]
        scenario_lines = [
            "| Skenario | Estimasi Arus Kas Masuk | Gap terhadap Total Invoice | Narasi Manajemen |",
            "|---|---:|---:|---|",
            f"| Upside | {cls._format_currency(expected_realization_upside)} | {cls._format_currency(expected_gap_upside)} | Terjadi bila perbaikan approval, kelengkapan dokumen, dan penagihan pada akun prioritas dijalankan cepat. |",
            f"| Base Case | {cls._format_currency(expected_realization_base)} | {cls._format_currency(expected_gap_base)} | Menggambarkan realisasi yang paling realistis bila perilaku penagihan tetap mengikuti pola historis saat ini. |",
            f"| Downside | {cls._format_currency(expected_realization_downside)} | {cls._format_currency(expected_gap_downside)} | Terjadi bila siklus anggaran, likuiditas pelanggan, atau dispute memburuk sehingga semakin banyak invoice bergeser ke kelas risiko tinggi. |",
        ]
        priority_table_lines = [
            "| Prioritas | Fokus | Penanggung Jawab | Isu Utama | Aksi 30 Hari | Dampak yang Diharapkan |",
            "|---:|---|---|---|---|---|",
        ]
        management_priority_lines = [
            "| Prioritas | Fokus | Kelas | Nilai Invoice | Isu Utama | Aksi Awal | Fungsi Utama |",
            "|---:|---|---|---:|---|---|---|",
        ]
        for item in priority_rows[:6]:
            management_priority_lines.append(
                f"| {item['priority']} | {item['focus']} | {item['payment_class']} | {cls._format_currency(item['invoice_value'])} | {item['issue']} | {item['action']} | {item['owner']} |"
            )
            priority_table_lines.append(
                f"| {item['priority']} | {item['focus']} | {item['owner']} | {item['issue']} | {item['action']} | {item['impact']} |"
            )
        meeting_questions = [
            "Segmen partner mana yang paling layak mendapat eskalasi manajemen karena menggabungkan nilai invoice besar dan skor risiko tinggi?",
            "Hambatan apa yang paling sering berulang: anggaran, approval, dokumen, likuiditas, atau sengketa ruang lingkup?",
            f"Apakah strategi 30 hari ke depan harus difokuskan pada {recent_period_label} dan akun-akun D/E agar gap arus kas masuk dapat ditekan tanpa memperburuk tekanan cash out?",
        ]
        agenda_lines = [f"{index}. {question}" for index, question in enumerate(meeting_questions, start=1)]
        management_brief_lines = (
            ["## Fakta Eksekutif yang Wajib Dijaga Konsisten"]
            + executive_fact_lines
            + ["", "## Skenario 1-2 Kuartal"]
            + scenario_lines
            + ["", "## Prioritas Penagihan untuk Pembahasan Internal"]
            + management_priority_lines
            + ["", "## Agenda Diskusi Manajemen"]
            + agenda_lines
        )

        return {
            "financial_summary": "\n".join(financial_summary_lines),
            "evidence": "\n".join(evidence_lines),
            "diagnostic_breakdown": "\n".join(diagnostic_breakdown_lines),
            "management_brief": "\n".join(management_brief_lines),
            "executive_facts": "\n".join(executive_fact_lines),
            "scenario_table": "\n".join(scenario_lines),
            "priority_table": "\n".join(priority_table_lines),
            "meeting_agenda": "\n".join(agenda_lines),
            "base_profile": {
                "data_mode": data_mode,
                "total_invoices": total_invoices,
                "total_invoice_value": total_invoice_value,
                "delayed_invoices": delayed_invoices,
                "delayed_invoice_value": delayed_invoice_value,
                "high_risk_invoices": high_risk_invoices,
                "high_risk_invoice_value": high_risk_invoice_value,
                "expected_realization_base": expected_realization_base,
                "expected_gap_base": expected_gap_base,
                "core_fields_available": len(core_field_map) - len(missing_core_fields),
                "core_fields_expected": len(core_field_map),
                "missing_core_fields": missing_core_fields,
                "top_risk_partners": high_risk_partner_df.index.tolist()[:3],
                "top_risk_services": high_risk_service_df.index.tolist()[:3],
            },
            "visual_prompt": cls._build_visual_prompt(class_distribution),
        }


class Researcher:
    _SERPER_ENDPOINTS = {
        "search": "https://google.serper.dev/search",
        "news": "https://google.serper.dev/news",
    }

    _OSINT_TOPICS = [
        {
            "topic": "Siklus Anggaran Pemerintah",
            "query": "siklus pencairan APBN APBD termin pembayaran vendor Indonesia",
        },
        {
            "topic": "Perilaku Pembayaran BUMN dan Korporasi",
            "query": "tren keterlambatan pembayaran invoice BUMN swasta Indonesia",
        },
        {
            "topic": "Likuiditas dan Piutang Bisnis",
            "query": "risiko likuiditas perusahaan jasa Indonesia karena piutang tertunda",
        },
        {
            "topic": "Regulasi Pengadaan dan Kontrak",
            "query": "regulasi terbaru pengadaan pemerintah termin pembayaran penyedia Indonesia",
        },
    ]

    _DELAY_FACTOR_TOPICS = [
        {
            "factor": "Siklus anggaran pemerintah",
            "query": "siklus pencairan APBN APBD keterlambatan pembayaran vendor Indonesia",
            "delay_days": (10, 30),
            "impact": "Potensi mundur tambahan ketika termin pembayaran bergantung pada pencairan anggaran.",
        },
        {
            "factor": "Approval korporasi dan BUMN",
            "query": "approval internal BUMN korporasi keterlambatan pembayaran invoice Indonesia",
            "delay_days": (7, 21),
            "impact": "Potensi penambahan hari tunggu karena approval berlapis, BAST, atau verifikasi akhir.",
        },
        {
            "factor": "Likuiditas pelanggan",
            "query": "tekanan likuiditas perusahaan Indonesia keterlambatan pembayaran invoice jasa",
            "delay_days": (14, 45),
            "impact": "Potensi penundaan tambahan ketika pelanggan sedang menjaga kas atau menahan pengeluaran.",
        },
        {
            "factor": "Regulasi pengadaan dan administrasi kontrak",
            "query": "regulasi pengadaan pemerintah administrasi kontrak termin pembayaran vendor Indonesia",
            "delay_days": (5, 20),
            "impact": "Potensi penambahan waktu akibat revisi dokumen, termin, atau penyesuaian administrasi kontrak.",
        },
    ]

    _PROFILE_KEYWORD_GROUPS = {
        "government": ("pemerintah", "pemda", "kementerian", "dinas", "instansi", "apbn", "apbd", "pengadaan"),
        "bumn": ("bumn", "bumd", "persero", "holding negara"),
        "corporate": ("korporasi", "swasta", "enterprise", "perusahaan"),
        "training": ("pelatihan", "training", "sertifikasi", "academy", "bootcamp"),
        "consulting": ("konsultan", "consulting", "implementasi", "proyek", "jasa"),
        "payment_ops": ("invoice", "termin", "tagihan", "piutang", "pembayaran", "approval", "bast", "vendor"),
        "liquidity": ("cashflow", "arus kas", "likuiditas", "pencairan", "dana"),
    }

    _STRICT_PROFILE_TAGS = {"government", "bumn", "corporate", "training", "consulting"}

    @staticmethod
    def _is_serper_available():
        return bool(
            SERPER_API_KEY
            and SERPER_API_KEY.strip()
            and SERPER_API_KEY != "masukkan_api_key_serper_anda_disini"
        )

    @staticmethod
    def _normalize_osint_fragment(text, max_length=240):
        normalized = re.sub(r"\s+", " ", str(text or "")).strip()
        normalized = normalized.replace("…", " ")
        normalized = re.sub(r"\.{3,}", " ", normalized)
        normalized = re.sub(r"\s+([,.;:!?])", r"\1", normalized)
        normalized = normalized.strip(" \"'`-–—")
        if not normalized:
            return ""
        if len(normalized) <= max_length:
            return normalized

        candidate = normalized[: max_length + 1]
        sentence_breaks = [candidate.rfind(marker) for marker in (".", "!", "?", ";", ":")]
        best_sentence_break = max(sentence_breaks)
        if best_sentence_break >= int(max_length * 0.5):
            candidate = candidate[: best_sentence_break + 1]
        else:
            last_space = candidate.rfind(" ")
            candidate = candidate[: last_space if last_space > 0 else max_length]
        return candidate.strip(" ,;:-")

    @classmethod
    def _extract_profile_tags(cls, text):
        lowered = str(text or "").lower()
        tags = set()
        for tag, keywords in cls._PROFILE_KEYWORD_GROUPS.items():
            if any(keyword in lowered for keyword in keywords):
                tags.add(tag)
        return tags

    @staticmethod
    def _is_low_signal_fragment(entry):
        title = str((entry or {}).get("title") or "")
        snippet = str((entry or {}).get("snippet") or "")
        combined = f"{title} {snippet}"
        if combined.count("...") >= 2 or combined.count("…") >= 2:
            return True
        if combined.count('"') % 2 == 1:
            return True
        cleaned = Researcher._normalize_osint_fragment(combined, max_length=120)
        return len(cleaned) < 35

    @classmethod
    def _is_company_comparable_entry(cls, entry, extra_context=""):
        context_tags = cls._extract_profile_tags(extra_context)
        profile_tags = context_tags & cls._STRICT_PROFILE_TAGS
        if not profile_tags:
            return False

        entry_text = " ".join(
            [
                str(entry.get("title") or ""),
                str(entry.get("snippet") or ""),
                str(entry.get("domain") or ""),
            ]
        )
        entry_tags = cls._extract_profile_tags(entry_text)
        if not (entry_tags & profile_tags):
            return False
        if not (entry_tags & {"payment_ops", "liquidity"}):
            return False
        return not cls._is_low_signal_fragment(entry)

    @classmethod
    def _filter_company_comparable_entries(cls, entries, extra_context=""):
        return [
            entry for entry in entries
            if cls._is_company_comparable_entry(entry, extra_context)
        ]

    @classmethod
    def _build_entry_summary(cls, entry):
        raw_title = str(entry.get("title") or "")
        raw_snippet = str(entry.get("snippet") or "")
        source = entry.get("domain") or "-"
        date = f" ({entry['date']})" if entry.get("date") else ""

        use_title = raw_title and "..." not in raw_title and "…" not in raw_title
        headline = cls._normalize_osint_fragment(raw_title if use_title else raw_snippet, max_length=120)
        summary = cls._normalize_osint_fragment(raw_snippet or raw_title, max_length=220)

        lines = []
        if headline:
            lines.append(headline)
        if summary and summary != headline:
            lines.append(f"  Ringkasan: {summary}")
        lines.append(f"  Sumber: {source}{date}")
        return "\n".join(lines)

    @staticmethod
    def fetch_full_markdown(url):
        """Fetches the clean markdown text of any URL using Jina Reader."""
        if not url: return ""
        try:
            jina_url = f"https://r.jina.ai/{url}"
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(jina_url, headers=headers, timeout=12)
            if response.status_code == 200:
                return response.text[:6000] 
            return ""
        except Exception as e:
            logger.warning("Failed to fetch full markdown for %s: %s", url, e)
            return ""

    @classmethod
    def extract_insight_with_llm(cls, url, extraction_goal):
        """Universal Deep Scraper: Reads a URL and extracts a specific qualitative insight via Pydantic/LLM."""
        markdown_text = cls.fetch_full_markdown(url)
        if not markdown_text:
            return ""
            
        prompt = f"""
        You are an expert business researcher. Read the following source text.
        Your goal is to extract: {extraction_goal}
        
        SOURCE TEXT:
        {markdown_text}
        
        Respond ONLY with a valid JSON object using this schema. If the information is not present, use "NOT_FOUND".
        {{
            "insight": "<concise professional summary in Indonesian>"
        }}
        """
        try:
            client = Client(host=OLLAMA_HOST)
            res = client.chat(
                model=LLM_MODEL,
                messages=[{'role': 'user', 'content': prompt}],
                options={'temperature': 0.0}
            )
            raw_text = res['message']['content']
            match = re.search(r'\{.*\}', raw_text, re.DOTALL)
            parsed_dict = json.loads(match.group(0)) if match else json.loads(raw_text)
            
            # Pydantic validation
            data = InsightSchema.model_validate(parsed_dict)
            
            if "NOT_FOUND" in data.insight.upper() or not data.insight:
                return ""
            return data.insight
        except Exception as e:
            logger.warning("Insight extraction failed for %s: %s", url, e)
            return ""

    @classmethod
    def _execute_serper_query(cls, query, mode="search", num_results=6):
        if not cls._is_serper_available():
            return []

        endpoint = cls._SERPER_ENDPOINTS.get(mode, cls._SERPER_ENDPOINTS["search"])
        headers = {
            "X-API-KEY": SERPER_API_KEY,
            "Content-Type": "application/json",
        }
        payload = {
            "q": query,
            "num": num_results,
            "gl": "id",
            "hl": "id",
        }

        try:
            response = requests.post(endpoint, headers=headers, data=json.dumps(payload), timeout=10)
            response.raise_for_status()
            body = response.json()
        except Exception as exc:
            logger.warning("Serper %s request failed: %s", mode, exc)
            return []

        result_key = "organic" if mode == "search" else "news"
        rows = body.get(result_key, [])

        normalized = []
        for row in rows:
            title = (row.get("title") or "").strip()
            snippet = (row.get("snippet") or "").strip()
            link = (row.get("link") or "").strip()
            date = (row.get("date") or "").strip()

            if not title and not snippet:
                continue

            domain = urlparse(link).netloc.replace("www.", "") if link else "-"
            normalized.append(
                {
                    "title": title,
                    "snippet": snippet,
                    "link": link,
                    "domain": domain,
                    "date": date,
                }
            )

        return normalized

    @staticmethod
    def _deduplicate(items):
        deduplicated = []
        seen = set()

        for item in items:
            fingerprint = "|".join(
                [
                    item.get("domain", "").lower(),
                    item.get("title", "").lower(),
                    item.get("snippet", "")[:120].lower(),
                ]
            )
            if fingerprint in seen:
                continue
            seen.add(fingerprint)
            deduplicated.append(item)

        return deduplicated

    @staticmethod
    def _format_topic(topic, entries):
        lines = [f"[{topic}]"]

        if not entries:
            lines.append("- Tidak ada sinyal eksternal yang cukup sebanding dengan profil perusahaan saat ini.")
            return "\n".join(lines)

        for index, entry in enumerate(entries[:3], start=1):
            lines.append(f"{index}.")
            lines.append(Researcher._build_entry_summary(entry))

        return "\n".join(lines)

    @classmethod
    @osint_cache.memoize(expire=86400)
    def get_macro_finance_trends(cls, extra_context=""):
        if not cls._is_serper_available():
            return "Data OSINT eksternal tidak tersedia (SERPER_API_KEY belum dikonfigurasi)."

        context_snippet = (extra_context or "").strip()
        if not (cls._extract_profile_tags(context_snippet) & cls._STRICT_PROFILE_TAGS):
            return "OSINT tidak dipakai karena konteks perusahaan yang sebanding belum cukup jelas."
        search_jobs = []
        for topic_config in cls._OSINT_TOPICS:
            query = topic_config["query"]
            if context_snippet:
                query = f"{query} {context_snippet[:180]}"

            search_jobs.append((topic_config["topic"], query, "search"))
            search_jobs.append((topic_config["topic"], query, "news"))

        topic_results = {topic["topic"]: [] for topic in cls._OSINT_TOPICS}

        with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
            future_map = {
                executor.submit(cls._execute_serper_query, query, mode, 6): topic
                for topic, query, mode in search_jobs
            }
            for future, topic in future_map.items():
                try:
                    topic_results[topic].extend(future.result())
                except Exception as exc:
                    logger.warning("OSINT future failed for %s: %s", topic, exc)

        # --- DEEP SCRAPE THE #1 RESULT ---
        deep_insight_text = ""
        top_link = None
        for topic, entries in topic_results.items():
            comparable_entries = cls._filter_company_comparable_entries(entries, extra_context=context_snippet)
            if comparable_entries and comparable_entries[0].get("link"):
                top_link = comparable_entries[0]["link"]
                break
                
        if top_link:
            logger.info("Deep scraping OSINT for macro finance trends: %s", top_link)
            goal = "What are the latest macro trends regarding B2B payment behavior, budget cycles, or invoice collection challenges in Indonesia?"
            insight = cls.extract_insight_with_llm(top_link, goal)
            if insight:
                source_domain = urlparse(top_link).netloc.replace("www.", "")
                deep_insight_text = f"**Insight Mendalam (via {source_domain}):** {insight}\n\n"

        blocks = []
        for topic_config in cls._OSINT_TOPICS:
            topic_name = topic_config["topic"]
            unique_entries = cls._filter_company_comparable_entries(
                cls._deduplicate(topic_results.get(topic_name, [])),
                extra_context=context_snippet,
            )
            if unique_entries:
                blocks.append(cls._format_topic(topic_name, unique_entries))

        combined = "\n\n".join(blocks).strip()
        if not blocks:
            combined = "OSINT tidak dipakai karena tidak ada sinyal eksternal yang cukup sebanding dengan kondisi perusahaan."

        return deep_insight_text + combined

    @classmethod
    @osint_cache.memoize(expire=86400)
    def get_chapter_signal(cls, chapter_keywords, notes=""):
        if not cls._is_serper_available():
            return "Sinyal OSINT per bab tidak tersedia."

        query = (
            "Indonesia payment behavior invoice collection risk "
            f"{chapter_keywords or ''} {notes or ''}"
        ).strip()

        results = cls._execute_serper_query(query, mode="search", num_results=5)
        unique_results = cls._filter_company_comparable_entries(
            cls._deduplicate(results),
            extra_context=f"{chapter_keywords or ''} {notes or ''}",
        )
        if not unique_results:
            return "OSINT bab ini tidak dipakai karena belum ada sinyal eksternal yang cukup sebanding."

        lines = []
        for index, entry in enumerate(unique_results[:3], start=1):
            lines.append(f"{index}.")
            lines.append(cls._build_entry_summary(entry))

        return "\n".join(lines)

    @classmethod
    @osint_cache.memoize(expire=86400)
    def get_payment_delay_risks(cls, extra_context=""):
        if not cls._is_serper_available():
            return []

        context_snippet = (extra_context or "").strip()
        if not (cls._extract_profile_tags(context_snippet) & cls._STRICT_PROFILE_TAGS):
            return []
        factors = []

        for topic in cls._DELAY_FACTOR_TOPICS:
            query = topic["query"]
            if context_snippet:
                query = f"{query} {context_snippet[:180]}"

            search_results = cls._execute_serper_query(query, mode="search", num_results=4)
            news_results = cls._execute_serper_query(query, mode="news", num_results=4)
            combined = cls._filter_company_comparable_entries(
                cls._deduplicate(search_results + news_results),
                extra_context=context_snippet,
            )
            if not combined:
                continue

            sources = []
            snippets = []
            for item in combined[:2]:
                source = item.get("domain") or "-"
                if source not in sources:
                    sources.append(source)
                snippet = cls._normalize_osint_fragment(item.get("snippet") or item.get("title"), max_length=180)
                if snippet:
                    snippets.append(snippet)

            factors.append(
                {
                    "factor": topic["factor"],
                    "potential_delay_days": {
                        "min": topic["delay_days"][0],
                        "max": topic["delay_days"][1],
                    },
                    "impact": topic["impact"],
                    "summary": " ".join(snippets[:2]).strip(),
                    "source_domains": sources,
                }
            )

        return factors


class StyleEngine:
    @staticmethod
    def _insert_field(paragraph, field_code, placeholder="1"):
        begin_run = paragraph.add_run()
        begin_field = OxmlElement("w:fldChar")
        begin_field.set(qn("w:fldCharType"), "begin")
        begin_run._r.append(begin_field)

        instruction_run = paragraph.add_run()
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = field_code
        instruction_run._r.append(instruction)

        separate_run = paragraph.add_run()
        separate_field = OxmlElement("w:fldChar")
        separate_field.set(qn("w:fldCharType"), "separate")
        separate_run._r.append(separate_field)

        paragraph.add_run(placeholder)

        end_run = paragraph.add_run()
        end_field = OxmlElement("w:fldChar")
        end_field.set(qn("w:fldCharType"), "end")
        end_run._r.append(end_field)

    @classmethod
    def insert_toc_field(cls, paragraph):
        cls._insert_field(
            paragraph,
            'TOC \\o "1-3" \\h \\z \\u',
            "Klik kanan lalu pilih Update Field untuk memuat daftar isi.",
        )

    @classmethod
    def apply_document_styles(cls, doc, theme_color):
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

            header = section.header
            header_paragraph = (
                header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            )
            header_paragraph.text = ""
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            header_paragraph.add_run("INIXINDO JOGJA | INTERNAL FINANCE REPORT")
            for run in header_paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(120, 120, 120)

            footer = section.footer
            footer_paragraph = (
                footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            )
            footer_paragraph.text = ""
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer_paragraph.add_run("STRICTLY CONFIDENTIAL | Page ")
            cls._insert_field(footer_paragraph, "PAGE")
            footer_paragraph.add_run(" of ")
            cls._insert_field(footer_paragraph, "NUMPAGES")
            for run in footer_paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(110, 110, 110)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = RGBColor(33, 37, 41)
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_paragraph.line_spacing = 1.15
        normal_paragraph.space_after = Pt(8)

        heading_1 = doc.styles["Heading 1"]
        heading_1.font.name = "Calibri"
        heading_1.font.size = Pt(16)
        heading_1.font.bold = True
        heading_1.font.color.rgb = RGBColor(*theme_color)
        heading_1.paragraph_format.space_before = Pt(18)
        heading_1.paragraph_format.space_after = Pt(8)

        heading_2 = doc.styles["Heading 2"]
        heading_2.font.name = "Calibri"
        heading_2.font.size = Pt(13)
        heading_2.font.bold = True
        heading_2.font.color.rgb = RGBColor(0, 0, 0)
        heading_2.paragraph_format.space_before = Pt(14)
        heading_2.paragraph_format.space_after = Pt(4)

        heading_3 = doc.styles["Heading 3"]
        heading_3.font.name = "Calibri"
        heading_3.font.size = Pt(12)
        heading_3.font.bold = True
        heading_3.font.color.rgb = RGBColor(64, 64, 64)
        heading_3.paragraph_format.space_before = Pt(10)
        heading_3.paragraph_format.space_after = Pt(4)

        for style_name in [
            "List Bullet",
            "List Bullet 2",
            "List Bullet 3",
            "List Number",
            "List Number 2",
            "List Number 3",
        ]:
            try:
                list_style = doc.styles[style_name]
                list_style.font.name = "Calibri"
                list_style.font.size = Pt(11)
                list_style.paragraph_format.space_after = Pt(4)
            except KeyError:
                continue

        try:
            caption_style = doc.styles["Caption"]
            caption_style.font.name = "Calibri"
            caption_style.font.size = Pt(10)
            caption_style.font.italic = True
            caption_style.font.color.rgb = RGBColor(80, 80, 80)
        except KeyError:
            pass


class ChartEngine:
    @staticmethod
    def _theme_to_plt_color(theme_color):
        return tuple(component / 255 for component in theme_color)

    @staticmethod
    def _format_compact_currency(value):
        amount = float(value or 0)
        if abs(amount) >= 1_000_000_000:
            return f"Rp {amount / 1_000_000_000:.1f}M"
        if abs(amount) >= 1_000_000:
            return f"Rp {amount / 1_000_000:.0f} juta"
        return f"Rp {amount:,.0f}".replace(",", ".")

    @staticmethod
    def create_bar_chart(data_str, theme_color):
        try:
            parts = [part.strip() for part in data_str.split("|")]
            if len(parts) >= 3:
                title = parts[0]
                y_label = parts[1]
                raw_data = "|".join(parts[2:])
            else:
                title = "Distribusi Historis Kelas Pembayaran"
                y_label = "Persentase"
                raw_data = data_str

            labels = []
            values = []
            for chunk in raw_data.split(";"):
                if "," not in chunk:
                    continue
                label, value = chunk.split(",", 1)
                numeric_value = re.sub(r"[^\d.]", "", value)
                if not numeric_value:
                    continue
                labels.append(label.strip())
                values.append(float(numeric_value))

            if not labels:
                return None

            fig, axis = plt.subplots(figsize=(7, 4.5))
            axis.bar(
                labels,
                values,
                color=ChartEngine._theme_to_plt_color(theme_color),
                alpha=0.9,
                width=0.5,
            )
            axis.set_title(title, fontsize=12, fontweight="bold", pad=20)
            axis.set_ylabel(y_label, fontsize=10)
            axis.spines["top"].set_visible(False)
            axis.spines["right"].set_visible(False)

            image_stream = io.BytesIO()
            plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=150)
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception:
            return None

    @staticmethod
    def create_flowchart(data_str, theme_color):
        try:
            steps = [
                "\n".join(textwrap.wrap(step.strip(), width=18))
                for step in data_str.split("->")
                if step.strip()
            ]
            if len(steps) < 2:
                return None

            fig, axis = plt.subplots(figsize=(8, 3))
            axis.axis("off")
            x_positions = [index * 2.5 for index in range(len(steps))]

            for index in range(len(steps) - 1):
                axis.annotate(
                    "",
                    xy=(x_positions[index + 1] - 1.0, 0.5),
                    xytext=(x_positions[index] + 1.0, 0.5),
                    arrowprops={"arrowstyle": "-|>", "lw": 1.5},
                )

            for index, step in enumerate(steps):
                box = patches.FancyBboxPatch(
                    (x_positions[index] - 1.0, 0.1),
                    2.0,
                    0.8,
                    boxstyle="round,pad=0.1",
                    fc=ChartEngine._theme_to_plt_color(theme_color),
                    alpha=0.9,
                )
                axis.add_patch(box)
                axis.text(
                    x_positions[index],
                    0.5,
                    step,
                    ha="center",
                    va="center",
                    size=9,
                    color="white",
                    fontweight="bold",
                )

            axis.set_xlim(-1.2, (len(steps) - 1) * 2.5 + 1.2)
            axis.set_ylim(0, 1)

            image_stream = io.BytesIO()
            plt.savefig(
                image_stream,
                format="png",
                bbox_inches="tight",
                dpi=200,
                transparent=True,
            )
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception:
            return None

    @staticmethod
    def create_dashboard_snapshot(data_str, theme_color):
        try:
            payload = json.loads(data_str)
            horizon_label = payload.get("horizon_label") or "Dashboard Cashflow"
            horizon_focus = payload.get("horizon_focus") or "-"
            status = str(payload.get("status") or "-").upper()
            current_cash = float(payload.get("current_cash") or 0)
            runway_months = float(payload.get("runway_months") or 0)
            coverage_ratio = float(payload.get("coverage_ratio") or 0)
            average_delay_days = float(payload.get("average_delay_days") or 0)
            balance_points = payload.get("balance_projection") or []
            coverage_bars = payload.get("coverage_bars") or []

            fig = plt.figure(figsize=(9.4, 5.4))
            grid = fig.add_gridspec(2, 2, height_ratios=[1.0, 1.65], width_ratios=[1.55, 1.0], hspace=0.34, wspace=0.28)
            header_axis = fig.add_subplot(grid[0, :])
            balance_axis = fig.add_subplot(grid[1, 0])
            coverage_axis = fig.add_subplot(grid[1, 1])

            for axis in (header_axis,):
                axis.axis("off")

            fig.patch.set_facecolor("white")
            theme_rgb = ChartEngine._theme_to_plt_color(theme_color)
            safe_color = "#16a34a"
            watch_color = "#d97706"
            risk_color = "#dc2626"
            status_color = safe_color if status == "AMAN" else watch_color if status == "WASPADA" else risk_color

            header_axis.text(
                0.0,
                0.98,
                "Cashflow Health Dashboard",
                fontsize=16,
                fontweight="bold",
                color="#203152",
                va="top",
                transform=header_axis.transAxes,
            )
            header_axis.text(
                0.0,
                0.76,
                horizon_label,
                fontsize=12,
                fontweight="bold",
                color="#0f172a",
                va="top",
                transform=header_axis.transAxes,
            )
            header_axis.text(
                0.0,
                0.62,
                horizon_focus,
                fontsize=10,
                color="#475569",
                va="top",
                transform=header_axis.transAxes,
            )

            status_box = patches.FancyBboxPatch(
                (0.76, 0.72),
                0.2,
                0.16,
                boxstyle="round,pad=0.02,rounding_size=0.02",
                fc=status_color,
                ec="none",
                alpha=0.95,
                transform=header_axis.transAxes,
            )
            header_axis.add_patch(status_box)
            header_axis.text(
                0.86,
                0.8,
                f"STATUS: {status}",
                fontsize=11,
                fontweight="bold",
                color="white",
                ha="center",
                va="center",
                transform=header_axis.transAxes,
            )

            metric_cards = [
                ("Cash", ChartEngine._format_compact_currency(current_cash)),
                ("Runway", f"{runway_months:.1f} bulan"),
                ("Coverage", f"{coverage_ratio:.2f}x"),
                ("Delay", f"{average_delay_days:.0f} hari"),
            ]
            card_width = 0.22
            card_gap = 0.025
            card_y = 0.08
            for index, (label, value) in enumerate(metric_cards):
                x = index * (card_width + card_gap)
                card = patches.FancyBboxPatch(
                    (x, card_y),
                    card_width,
                    0.34,
                    boxstyle="round,pad=0.015,rounding_size=0.02",
                    fc="#f8fafc" if index != 0 else "#eef2ff",
                    ec="#cbd5e1",
                    lw=1,
                    transform=header_axis.transAxes,
                )
                header_axis.add_patch(card)
                header_axis.text(x + 0.02, card_y + 0.23, label, fontsize=9.5, fontweight="bold", color="#475569", transform=header_axis.transAxes)
                header_axis.text(x + 0.02, card_y + 0.08, value, fontsize=13.5, fontweight="bold", color="#111827", transform=header_axis.transAxes)

            balance_axis.set_title("Prediksi Saldo", fontsize=11, fontweight="bold", loc="left", color="#334155")
            if balance_points:
                x_values = list(range(len(balance_points)))
                balances = [float(point.get("balance") or 0) for point in balance_points]
                labels = [str(point.get("label") or "") for point in balance_points]
                balance_axis.fill_between(x_values, balances, color=theme_rgb, alpha=0.18)
                balance_axis.plot(x_values, balances, color=theme_rgb, marker="o", linewidth=2)
                balance_axis.set_xticks(x_values)
                balance_axis.set_xticklabels(labels, fontsize=8)
                min_balance = min(balances)
                threshold = 100_000_000
                if min_balance <= threshold:
                    balance_axis.axhline(threshold, color=risk_color, linestyle="--", linewidth=1.2)
                balance_axis.grid(axis="y", linestyle="--", alpha=0.25)
            else:
                balance_axis.text(0.5, 0.5, "Tidak ada proyeksi saldo.", ha="center", va="center", transform=balance_axis.transAxes)
            balance_axis.spines["top"].set_visible(False)
            balance_axis.spines["right"].set_visible(False)
            balance_axis.spines["left"].set_color("#cbd5e1")
            balance_axis.spines["bottom"].set_color("#cbd5e1")
            balance_axis.tick_params(axis="y", labelsize=8)

            coverage_axis.set_title("Coverage & Runway", fontsize=11, fontweight="bold", loc="left", color="#334155")
            if coverage_bars:
                labels = [str(bar.get("label") or "") for bar in coverage_bars]
                values = [float(bar.get("value") or 0) for bar in coverage_bars]
                colors = []
                for bar in coverage_bars:
                    variant = str(bar.get("variant") or "")
                    if variant == "danger":
                        colors.append(risk_color)
                    elif variant == "target":
                        colors.append(watch_color)
                    elif variant == "current":
                        colors.append("#2563eb")
                    else:
                        colors.append("#0f766e")
                y_positions = list(range(len(values)))
                coverage_axis.barh(y_positions, values, color=colors, alpha=0.9)
                coverage_axis.set_yticks(y_positions)
                coverage_axis.set_yticklabels(labels, fontsize=8)
                coverage_axis.invert_yaxis()
                for index, value in enumerate(values):
                    coverage_axis.text(value + 0.03, index, f"{value:.2f}", va="center", fontsize=8)
                coverage_axis.axvline(1.0, color=risk_color, linestyle="--", linewidth=1)
                coverage_axis.axvline(1.2, color=watch_color, linestyle="--", linewidth=1)
                coverage_axis.grid(axis="x", linestyle="--", alpha=0.2)
            else:
                coverage_axis.text(0.5, 0.5, "Tidak ada data coverage.", ha="center", va="center", transform=coverage_axis.transAxes)
            coverage_axis.spines["top"].set_visible(False)
            coverage_axis.spines["right"].set_visible(False)
            coverage_axis.spines["left"].set_color("#cbd5e1")
            coverage_axis.spines["bottom"].set_color("#cbd5e1")
            coverage_axis.tick_params(axis="x", labelsize=8)

            image_stream = io.BytesIO()
            plt.savefig(image_stream, format="png", bbox_inches="tight", dpi=170)
            plt.close(fig)
            image_stream.seek(0)
            return image_stream
        except Exception:
            return None


class DocumentBuilder:
    @staticmethod
    def _append_inline_text(paragraph, node, bold=False, italic=False, underline=False, monospace=False):
        if isinstance(node, NavigableString):
            text = str(node)
            if not text:
                return
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if monospace:
                run.font.name = "Consolas"
            return

        if not isinstance(node, Tag):
            return

        next_bold = bold or node.name in {"strong", "b"}
        next_italic = italic or node.name in {"em", "i"}
        next_underline = underline or node.name == "u"
        next_monospace = monospace or node.name == "code"

        if node.name == "br":
            paragraph.add_run("\n")
            return

        if node.name == "a":
            for child in node.children:
                DocumentBuilder._append_inline_text(
                    paragraph,
                    child,
                    bold=next_bold,
                    italic=next_italic,
                    underline=True,
                    monospace=next_monospace,
                )
            return

        for child in node.children:
            DocumentBuilder._append_inline_text(
                paragraph,
                child,
                bold=next_bold,
                italic=next_italic,
                underline=next_underline,
                monospace=next_monospace,
            )

    @staticmethod
    def _resolve_list_style(doc, ordered, level):
        ordered_styles = ["List Number", "List Number 2", "List Number 3"]
        bullet_styles = ["List Bullet", "List Bullet 2", "List Bullet 3"]
        style_names = ordered_styles if ordered else bullet_styles
        preferred_style = style_names[min(level, len(style_names) - 1)]

        try:
            doc.styles[preferred_style]
            return preferred_style
        except KeyError:
            return "List Number" if ordered else "List Bullet"

    @classmethod
    def _add_list(cls, doc, list_tag, level=0, ordered=False):
        style_name = cls._resolve_list_style(doc, ordered, level)

        for list_item in list_tag.find_all("li", recursive=False):
            inline_nodes = []
            nested_lists = []

            for child in list_item.children:
                if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                    nested_lists.append(child)
                else:
                    inline_nodes.append(child)

            if inline_nodes:
                paragraph = doc.add_paragraph(style=style_name)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for node in inline_nodes:
                    cls._append_inline_text(paragraph, node)

            for nested_list in nested_lists:
                cls._add_list(
                    doc,
                    nested_list,
                    level=level + 1,
                    ordered=nested_list.name == "ol",
                )

    @staticmethod
    def _add_table(doc, table_tag):
        rows = table_tag.find_all("tr")
        if not rows:
            return

        header_cells = rows[0].find_all(["th", "td"])
        if not header_cells:
            return

        column_count = len(header_cells)
        table = doc.add_table(rows=1, cols=column_count)
        table.style = "Table Grid"

        for column_index, cell in enumerate(header_cells):
            paragraph = table.rows[0].cells[column_index].paragraphs[0]
            paragraph.text = cell.get_text(" ", strip=True)
            if paragraph.runs:
                paragraph.runs[0].bold = True

        for html_row in rows[1:]:
            html_cells = html_row.find_all(["th", "td"])
            table_cells = table.add_row().cells
            for column_index in range(column_count):
                value = ""
                if column_index < len(html_cells):
                    value = html_cells[column_index].get_text(" ", strip=True)
                table_cells[column_index].text = value

    @classmethod
    def parse_html_to_docx(cls, doc, html_content, theme_color):
        soup = BeautifulSoup(html_content, "html.parser")

        for element in soup.contents:
            if isinstance(element, NavigableString):
                if not str(element).strip():
                    continue
                paragraph = doc.add_paragraph(str(element).strip())
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                continue

            if not isinstance(element, Tag):
                continue

            if element.name in {"h1", "h2", "h3", "h4"}:
                level = min(max(int(element.name[1]), 1), 3)
                heading = doc.add_heading(level=level)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cls._append_inline_text(heading, element)
                if level == 1:
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(*theme_color)
                continue

            if element.name == "p":
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for child in element.children:
                    cls._append_inline_text(paragraph, child)
                continue

            if element.name in {"ul", "ol"}:
                cls._add_list(doc, element, level=0, ordered=element.name == "ol")
                continue

            if element.name == "table":
                cls._add_table(doc, element)

    @classmethod
    def _flush_markdown_block(cls, doc, lines, theme_color):
        if not lines:
            return

        markdown_text = "\n".join(lines).strip()
        if not markdown_text:
            return

        html_content = markdown.markdown(markdown_text, extensions=["tables"])
        cls.parse_html_to_docx(doc, html_content, theme_color)

    @staticmethod
    def _add_visual(doc, marker_type, marker_payload, theme_color):
        if marker_type == "CHART":
            image = ChartEngine.create_bar_chart(marker_payload, theme_color)
            width = Inches(5.8)
            caption = "Grafik distribusi historis kelas pembayaran"
        elif marker_type == "DASHBOARD":
            image = ChartEngine.create_dashboard_snapshot(marker_payload, theme_color)
            try:
                payload = json.loads(marker_payload)
                caption_suffix = payload.get("horizon_label") or "Horizon aktif"
            except Exception:
                caption_suffix = "Horizon aktif"
            width = Inches(6.6)
            caption = f"Dashboard cashflow snapshot - {caption_suffix}"
        else:
            image = ChartEngine.create_flowchart(marker_payload, theme_color)
            width = Inches(6.3)
            caption = "Diagram alur rekomendasi mitigasi"

        if image is None:
            return

        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(image, width=width)

        try:
            caption_paragraph = doc.add_paragraph(caption, style="Caption")
        except KeyError:
            caption_paragraph = doc.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @classmethod
    def process_content(cls, doc, raw_text, theme_color=DEFAULT_COLOR):
        markdown_buffer = []

        for raw_line in raw_text.splitlines():
            stripped_line = raw_line.strip()

            if stripped_line.startswith("[[CHART:") and stripped_line.endswith("]]"
            ):
                cls._flush_markdown_block(doc, markdown_buffer, theme_color)
                markdown_buffer = []
                payload = stripped_line.replace("[[CHART:", "", 1).rsplit("]]", 1)[0].strip()
                cls._add_visual(doc, "CHART", payload, theme_color)
                continue

            if stripped_line.startswith("[[DASHBOARD:") and stripped_line.endswith("]]"
            ):
                cls._flush_markdown_block(doc, markdown_buffer, theme_color)
                markdown_buffer = []
                payload = stripped_line.replace("[[DASHBOARD:", "", 1).rsplit("]]", 1)[0].strip()
                cls._add_visual(doc, "DASHBOARD", payload, theme_color)
                continue

            if stripped_line.startswith("[[FLOW:") and stripped_line.endswith("]]"
            ):
                cls._flush_markdown_block(doc, markdown_buffer, theme_color)
                markdown_buffer = []
                payload = stripped_line.replace("[[FLOW:", "", 1).rsplit("]]", 1)[0].strip()
                cls._add_visual(doc, "FLOW", payload, theme_color)
                continue

            markdown_buffer.append(raw_line.rstrip())

        cls._flush_markdown_block(doc, markdown_buffer, theme_color)

    @staticmethod
    def create_cover(doc, theme_color=DEFAULT_COLOR):
        StyleEngine.apply_document_styles(doc, theme_color)

        properties = doc.core_properties
        properties.title = "Inixindo Cashflow Intelligence Report"
        properties.subject = "Internal Cashflow Intelligence Report"
        properties.author = WRITER_FIRM_NAME
        properties.category = "Finance"

        for _ in range(4):
            doc.add_paragraph()

        confidentiality = doc.add_paragraph("STRICTLY CONFIDENTIAL")
        confidentiality.alignment = WD_ALIGN_PARAGRAPH.CENTER
        confidentiality.runs[0].font.name = "Calibri"
        confidentiality.runs[0].font.size = Pt(10)
        confidentiality.runs[0].font.bold = True
        confidentiality.runs[0].font.color.rgb = RGBColor(120, 120, 120)

        doc.add_paragraph()

        title = doc.add_paragraph("CASHFLOW INTELLIGENCE REPORT")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.name = "Calibri"
        title.runs[0].font.size = Pt(22)
        title.runs[0].font.bold = True

        organization = doc.add_paragraph("INIXINDO JOGJA")
        organization.alignment = WD_ALIGN_PARAGRAPH.CENTER
        organization.runs[0].font.name = "Calibri"
        organization.runs[0].font.size = Pt(34)
        organization.runs[0].font.bold = True
        organization.runs[0].font.color.rgb = RGBColor(*theme_color)

        doc.add_paragraph()

        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = "Table Grid"
        metadata = [
            ("Cakupan Data", "Seluruh histori invoice dan catatan penagihan"),
            ("Tipe Laporan", "Analisis deskriptif, diagnostik, prediktif, dan preskriptif cashflow"),
            ("Tanggal Generasi", datetime.now().strftime("%d %B %Y")),
            ("Disusun Oleh", WRITER_FIRM_NAME),
        ]

        for row_index, (label, value) in enumerate(metadata):
            left_cell = metadata_table.rows[row_index].cells[0]
            right_cell = metadata_table.rows[row_index].cells[1]
            left_cell.text = label
            right_cell.text = value
            if left_cell.paragraphs[0].runs:
                left_cell.paragraphs[0].runs[0].bold = True

        doc.add_page_break()

    @staticmethod
    def add_table_of_contents(doc):
        doc.add_heading("Daftar Isi", level=1)
        toc_paragraph = doc.add_paragraph()
        StyleEngine.insert_toc_field(toc_paragraph)

        note = doc.add_paragraph(
            "Catatan: jika daftar isi belum muncul, klik kanan pada area daftar isi lalu pilih Update Field."
        )
        note.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(10)

        doc.add_page_break()


class ReportGenerator:
    SECTION_PASSES = (
        {
            "sections": REPORT_SECTION_SEQUENCE[:1],
            "include_visuals": False,
            "label": "executive_confidence",
        },
        {
            "sections": REPORT_SECTION_SEQUENCE[1:3],
            "include_visuals": True,
            "label": "diagnostic_evidence",
        },
        {
            "sections": REPORT_SECTION_SEQUENCE[3:],
            "include_visuals": False,
            "label": "actions_readiness",
        },
    )

    def __init__(self, kb_instance):
        self.ollama = Client(host=OLLAMA_HOST)
        self.kb = kb_instance
        self.io_pool = concurrent.futures.ThreadPoolExecutor(max_workers=4)

    @staticmethod
    def _normalize_analysis_payload(analysis_payload):
        if isinstance(analysis_payload, dict):
            return analysis_payload
        if not analysis_payload:
            return {}
        if isinstance(analysis_payload, str):
            try:
                parsed = json.loads(analysis_payload)
            except json.JSONDecodeError:
                return {}
            return parsed if isinstance(parsed, dict) else {}
        return {}

    @staticmethod
    def _build_dashboard_visual_markers(analysis_payload):
        payload = ReportGenerator._normalize_analysis_payload(analysis_payload)
        horizon_snapshot = payload.get("horizon_snapshot") or {}
        forecasts = horizon_snapshot.get("forecasts") if isinstance(horizon_snapshot, dict) else None
        if not isinstance(forecasts, dict):
            return []

        markers = []
        for horizon_key in ("short_term", "mid_term", "long_term"):
            forecast = forecasts.get(horizon_key)
            if not isinstance(forecast, dict):
                continue
            dashboard = forecast.get("dashboard_snapshot")
            if not isinstance(dashboard, dict):
                continue
            coverage_bars = (((dashboard.get("coverage_chart") or {}).get("bars")) or [])[:4]
            compact_payload = {
                "horizon_key": dashboard.get("horizon_key") or horizon_key,
                "horizon_label": dashboard.get("horizon_label"),
                "horizon_focus": dashboard.get("horizon_focus"),
                "status": dashboard.get("status"),
                "current_cash": dashboard.get("current_cash"),
                "runway_months": dashboard.get("runway_months"),
                "coverage_ratio": dashboard.get("coverage_ratio"),
                "average_delay_days": dashboard.get("average_delay_days"),
                "balance_projection": dashboard.get("balance_projection_30d") or [],
                "coverage_bars": coverage_bars,
            }
            markers.append(f"[[DASHBOARD:{json.dumps(compact_payload, ensure_ascii=False, separators=(',', ':'))}]]")
        return markers

    @staticmethod
    def _build_operational_snapshot_block(analysis_payload):
        payload = ReportGenerator._normalize_analysis_payload(analysis_payload)
        if not payload:
            return ""

        selected_period = payload.get("selected_period") or {}
        sync_status = payload.get("sync_status") or {}
        financial_sync = sync_status.get("financialData") or {}
        cash_out_sync = sync_status.get("cashOutSource") or {}

        lines = []
        period_label = selected_period.get("label")
        if period_label:
            lines.append(f"- Periode dashboard yang diekspor ke laporan: {period_label}.")

        cash_on_hand = payload.get("cash_on_hand")
        if cash_on_hand is not None:
            lines.append(f"- Cash in hand pada saat review: Rp{int(cash_on_hand):,}.")

        if financial_sync:
            freshness = financial_sync.get("sourceAgeMinutes")
            freshness_label = "belum tersedia"
            if freshness is not None:
                freshness_label = f"{float(freshness):.1f} menit"
            lines.append(
                "- Status sinkronisasi data finansial: "
                f"{financial_sync.get('syncStatus') or '-'}; usia data {freshness_label}; "
                f"record aktif {int(financial_sync.get('recordCount') or 0)}."
            )

        if cash_out_sync:
            if cash_out_sync.get("configured"):
                lines.append(
                    "- Sumber cash out memakai komitmen aktual dengan status "
                    f"{cash_out_sync.get('syncStatus') or '-'} dan "
                    f"{int(cash_out_sync.get('recordCount') or 0)} item aktif."
                )
            else:
                lines.append("- Sumber cash out masih memakai model operating cost bulanan karena feed kewajiban aktual belum dikonfigurasi.")

        horizons = ((payload.get("horizon_snapshot") or {}).get("forecasts")) or {}
        for key in ("short_term", "mid_term", "long_term"):
            horizon_payload = horizons.get(key) or {}
            forecast = horizon_payload.get("forecast") or {}
            health = horizon_payload.get("cashflow_health") or {}
            horizon = horizon_payload.get("time_horizon") or {}
            if not forecast:
                continue
            lines.append(
                "- "
                f"{horizon.get('label') or key}: cash masuk {FinancialAnalyzer._format_currency(int((forecast.get('cash_in') or {}).get('total_predicted_cash_in') or 0))}, "
                f"cash keluar {FinancialAnalyzer._format_currency(int((forecast.get('cash_out') or {}).get('total_cash_out') or 0))}, "
                f"ending cash {FinancialAnalyzer._format_currency(int(forecast.get('ending_cash') or 0))}, "
                f"runway {float(health.get('runway_months') or 0):.1f} bulan, "
                f"coverage {float(health.get('coverage_ratio') or 0):.2f}x, "
                f"fokus {horizon.get('focus') or '-'}."
            )

        return "\n".join(lines).strip()

    @staticmethod
    def _build_user_instruction(notes, active_sections):
        notes = (notes or "").strip()
        report_sections = "\n".join(f"- {section}" for section in active_sections)
        focus_block = notes if notes else "Tidak ada fokus tambahan dari pengguna."
        return (
            "Susun laporan internal yang detail, profesional, dan siap dipakai sebagai bahan diskusi rapat manajemen.\n"
            "Jangan menulis seperti jawaban AI generik; tulis seperti memo manajemen yang berbasis data.\n"
            "Kerjakan hanya section yang diminta pada pass ini dan hentikan output setelah section terakhir selesai.\n"
            "Pastikan setiap bagian di bawah terisi secara jelas:\n"
            f"{report_sections}\n\n"
            "Setiap rekomendasi harus menyebutkan fokus tindakan dan dampak yang diharapkan.\n"
            f"Fokus pengguna:\n{focus_block}"
        )

    @staticmethod
    def _build_section_scope(active_sections):
        section_list = "\n".join(f"- {section}" for section in active_sections)
        section_headings = "\n".join(f"   # {section}" for section in active_sections)
        section_scope = (
            "Generate only the sections listed below for this pass.\n"
            "Start with the first heading listed and stop after the last heading listed.\n"
            "Do not repeat prior sections and do not preview later sections.\n"
            f"{section_list}"
        )
        return section_scope, section_headings

    def _build_report_prompt(self, report_context, notes, analysis_context, macro_osint, active_sections, include_visuals):
        persona = PERSONAS.get("default", "Chief Financial Officer")
        section_scope, section_headings = self._build_section_scope(active_sections)
        structured_context = self._format_structured_context_block(analysis_context)
        return FINANCE_SYSTEM_PROMPT.format(
            persona=persona,
            financial_summary=report_context["financial_summary"],
            management_brief=report_context["management_brief"],
            internal_evidence=report_context["evidence"],
            industry_trends=macro_osint,
            user_focus=(notes or "Tidak ada fokus tambahan."),
            cashflow_context=(structured_context or "Tidak ada konteks forecast terstruktur tambahan."),
            readiness_signals=report_context["readiness_signals"],
            section_scope=section_scope,
            section_headings=section_headings,
            visual_prompt=report_context["visual_prompt"] if include_visuals else "",
        )

    @staticmethod
    def _score_report_completeness(raw_text):
        report_text = str(raw_text or "")
        if not report_text.strip():
            return {
                "score": 0.0,
                "passed": False,
                "components": {},
                "missing": ["Dokumen kosong."],
            }

        components = {}
        missing = []

        required_headings = [f"# {section}" for section in REPORT_SECTION_SEQUENCE]
        present_headings = [heading for heading in required_headings if heading in report_text]
        heading_score = (len(present_headings) / len(required_headings)) * 30
        components["top_level_sections"] = round(heading_score, 1)
        if len(present_headings) < len(required_headings):
            missing_sections = [heading.replace("# ", "") for heading in required_headings if heading not in present_headings]
            missing.append(f"Bagian utama belum lengkap: {', '.join(missing_sections)}.")

        required_subheadings = [
            "### Dampak Bisnis",
            "### Tingkat Keyakinan dan Caveat",
            "### Batasan Data dan Asumsi",
            "### Konteks OSINT Pendukung",
            "### Risiko dan Kontrol",
            "### Prasyarat Implementasi",
            "### Kesiapan Pelaksanaan",
        ]
        present_subheadings = [subheading for subheading in required_subheadings if subheading in report_text]
        subheading_score = (len(present_subheadings) / len(required_subheadings)) * 20
        components["required_subsections"] = round(subheading_score, 1)
        if len(present_subheadings) < len(required_subheadings):
            missing_subsections = [subheading.replace("### ", "") for subheading in required_subheadings if subheading not in present_subheadings]
            missing.append(f"Subbagian wajib belum lengkap: {', '.join(missing_subsections)}.")

        table_score = 0
        scenario_table_pattern = re.compile(
            r"\|\s*Skenario\s*\|\s*Estimasi Arus Kas Masuk\s*\|\s*Gap terhadap Total Invoice\s*\|",
            re.IGNORECASE,
        )
        if scenario_table_pattern.search(report_text):
            table_score += 8
        else:
            missing.append("Tabel skenario arus kas masuk belum ditemukan.")

        priority_table_pattern = re.compile(
            r"\|\s*Prioritas\s*\|\s*Fokus\s*\|\s*Penanggung Jawab\s*\|\s*Isu Utama\s*\|\s*Aksi 30 Hari\s*\|\s*Dampak yang Diharapkan\s*\|"
        )
        if priority_table_pattern.search(report_text):
            table_score += 12
        else:
            missing.append("Tabel prioritas 30 hari belum lengkap.")
        components["tables_and_owners"] = round(table_score, 1)

        enrichment_score = 0
        if "### Konteks OSINT Pendukung" in report_text:
            enrichment_score += 4
        if "[[CHART:" in report_text:
            enrichment_score += 3
        else:
            missing.append("Visual distribusi pembayaran belum masuk.")
        if "[[FLOW:" in report_text:
            enrichment_score += 3
        else:
            missing.append("Visual alur mitigasi belum masuk.")
        components["context_and_visuals"] = round(enrichment_score, 1)

        consistency_score = 10
        contradiction_patterns = [
            r"turun[^.\n]{0,120}memburuk",
            r"menurun[^.\n]{0,120}memburuk",
            r"naik[^.\n]{0,120}membaik",
            r"meningkat[^.\n]{0,120}membaik",
        ]
        if any(re.search(pattern, report_text, re.IGNORECASE) for pattern in contradiction_patterns):
            consistency_score = 0
            missing.append("Narasi tren risiko terdeteksi kontradiktif terhadap arah metrik.")
        components["numeric_consistency"] = round(consistency_score, 1)

        density_score = 0
        if len(report_text.strip()) >= 4500:
            density_score += 4
        elif len(report_text.strip()) >= 3200:
            density_score += 2
        else:
            missing.append("Narasi laporan masih terlalu tipis untuk bahan rapat internal.")

        if report_text.count("\n- ") >= 12:
            density_score += 3
        else:
            missing.append("Rincian bullet operasional masih kurang kaya.")

        if report_text.count("\n1.") >= 1 or report_text.count("\n1. ") >= 1:
            density_score += 1
        else:
            missing.append("Daftar tindakan bernomor belum kuat.")

        if "catatan" in report_text.lower() or "bukti" in report_text.lower():
            density_score += 2
        else:
            missing.append("Rujukan bukti internal belum cukup terlihat.")
        components["narrative_density"] = round(density_score, 1)

        total_score = round(sum(components.values()), 1)
        return {
            "score": total_score,
            "passed": total_score >= REPORT_MIN_COMPLETENESS_SCORE,
            "components": components,
            "missing": missing,
        }

    @staticmethod
    def _is_acceptable_report(raw_text):
        return ReportGenerator._score_report_completeness(raw_text)["passed"]

    @staticmethod
    def _extract_visual_markers(visual_prompt):
        chart_marker = ""
        flow_marker = ""
        for line in str(visual_prompt or "").splitlines():
            stripped_line = line.strip()
            if stripped_line.startswith("[[CHART:"):
                chart_marker = stripped_line
            elif stripped_line.startswith("[[FLOW:"):
                flow_marker = stripped_line
        return chart_marker, flow_marker

    @staticmethod
    def _split_top_level_sections(raw_text):
        matches = list(re.finditer(r"(?m)^# ([^\n]+?)\s*$", raw_text or ""))
        if not matches:
            return []

        sections = []
        for index, match in enumerate(matches):
            section_title = match.group(1).strip()
            section_end = matches[index + 1].start() if index + 1 < len(matches) else len(raw_text)
            section_body = raw_text[match.end():section_end].strip()
            sections.append({"title": section_title, "body": section_body})

        return sections

    @staticmethod
    def _join_top_level_sections(sections):
        blocks = []
        for section in sections:
            section_title = section["title"].strip()
            section_body = section["body"].strip()
            if section_body:
                blocks.append(f"# {section_title}\n{section_body}")
            else:
                blocks.append(f"# {section_title}")
        return "\n\n".join(blocks).strip()

    @staticmethod
    def _inject_subheading_block(section_body, subheading, content, before_subheading=None):
        if not content or not str(content).strip():
            return section_body

        subheading_marker = f"### {subheading}"
        if subheading_marker in section_body:
            return section_body

        new_block = f"{subheading_marker}\n{str(content).strip()}"
        if before_subheading:
            before_match = re.search(rf"(?m)^### {re.escape(before_subheading)}\s*$", section_body)
            if before_match:
                section_prefix = section_body[:before_match.start()].rstrip()
                section_suffix = section_body[before_match.start():].lstrip()
                return f"{section_prefix}\n\n{new_block}\n\n{section_suffix}".strip()

        return f"{section_body.rstrip()}\n\n{new_block}".strip()

    @staticmethod
    def _append_marker_block(section_body, marker):
        marker = str(marker or "").strip()
        if not marker or marker in section_body:
            return section_body
        return f"{section_body.rstrip()}\n\n{marker}".strip()

    @staticmethod
    def _format_structured_context_block(raw_text):
        lines = []
        for raw_line in str(raw_text or "").splitlines():
            cleaned_line = raw_line.strip()
            if not cleaned_line:
                if lines and lines[-1] != "":
                    lines.append("")
                continue
            if cleaned_line.startswith("===") and cleaned_line.endswith("==="):
                cleaned_line = cleaned_line.strip("= ").strip()
            if cleaned_line.endswith(":") and not cleaned_line.startswith("-"):
                lines.append(f"#### {cleaned_line[:-1].strip()}")
            else:
                lines.append(cleaned_line)
        return "\n".join(lines).strip()

    @classmethod
    def _sanitize_generated_report_text(cls, raw_text):
        sanitized = str(raw_text or "")
        sanitized = re.sub(r"(?<!\n)(===\s*[^\n=]+?\s*===)", r"\n\n\1", sanitized)
        sanitized = re.sub(r"\n?===\s*([^\n=]+?)\s*===\n?", lambda match: f"\n\n### {match.group(1).strip()}\n", sanitized)
        sanitized = re.sub(r"[ \t]+\n", "\n", sanitized)
        sanitized = re.sub(r"\n{3,}", "\n\n", sanitized)
        return sanitized.strip()

    def _finalize_report_content(self, raw_text, report_context, macro_osint, analysis_payload=None):
        raw_text = self._sanitize_generated_report_text(raw_text)
        sections = self._split_top_level_sections(raw_text)
        if not sections:
            return raw_text

        chart_marker, flow_marker = self._extract_visual_markers(report_context.get("visual_prompt", ""))
        dashboard_markers = self._build_dashboard_visual_markers(analysis_payload)
        operational_snapshot = self._build_operational_snapshot_block(analysis_payload)
        finalized_sections = []

        for section in sections:
            section_title = section["title"]
            section_body = section["body"]

            if section_title == "Analisis Deskriptif Cashflow":
                section_body = self._append_marker_block(section_body, chart_marker)
            elif section_title == "Analisis Diagnostik Cashflow":
                section_body = self._inject_subheading_block(
                    section_body,
                    "Konteks OSINT Pendukung",
                    macro_osint or "OSINT tidak dipakai karena tidak ada konteks eksternal yang cukup sebanding dengan profil perusahaan.",
                    before_subheading="Risiko dan Kontrol",
                )
            elif section_title == "Analisis Prediktif Cashflow":
                section_body = self._inject_subheading_block(
                    section_body,
                    "Snapshot Dashboard Operasional",
                    operational_snapshot,
                    before_subheading="Skenario 1-2 Kuartal",
                )
                for dashboard_marker in dashboard_markers:
                    section_body = self._append_marker_block(section_body, dashboard_marker)
            elif section_title == "Rekomendasi Preskriptif":
                section_body = self._append_marker_block(section_body, flow_marker)

            finalized_sections.append({"title": section_title, "body": section_body})

        return self._join_top_level_sections(finalized_sections)

    def _build_fallback_report(self, report_context, notes, analysis_context, macro_osint, analysis_payload=None):
        chart_marker, flow_marker = self._extract_visual_markers(report_context.get("visual_prompt", ""))
        dashboard_markers = self._build_dashboard_visual_markers(analysis_payload)
        operational_snapshot = self._build_operational_snapshot_block(analysis_payload)
        focus_block = notes.strip() if notes and notes.strip() else "Tidak ada fokus tambahan dari pengguna."
        structured_context_block = self._format_structured_context_block(analysis_context)

        lines = [
            "# Ringkasan Eksekutif",
            "### Dampak Bisnis",
            "- Laporan ini digunakan untuk membantu manajemen membaca risiko cashflow, memahami prioritas penagihan, mengendalikan tekanan cash out, dan menentukan tindakan yang paling cepat berdampak pada ending cash.",
            "- Fokus pengguna tetap dijaga dalam interpretasi, namun narasi dan prioritas diturunkan dari bukti internal, pola historis, serta konteks pelaksanaan yang tersedia saat ini.",
            "",
            "## Fakta Eksekutif",
            report_context["executive_facts"],
            "",
            "### Tingkat Keyakinan dan Caveat",
            report_context["confidence_summary"],
        ]

        lines.extend(
            [
                "",
                "# Analisis Deskriptif Cashflow",
                "### Snapshot Portofolio dan Konsentrasi Risiko",
                report_context["financial_summary"],
                "",
                "### Batasan Data dan Asumsi",
                report_context["assumptions"],
                "",
                "# Analisis Diagnostik Cashflow",
                "### Pola Hambatan Utama",
                report_context["diagnostic_breakdown"],
                "",
                "### Bukti Internal yang Mewakili",
                report_context["evidence"],
                "",
                "### Konteks OSINT Pendukung",
                macro_osint or "OSINT tidak dipakai karena tidak ada konteks eksternal yang cukup sebanding dengan profil perusahaan.",
                "",
                "### Risiko dan Kontrol",
                report_context["controls"],
                "",
                "### Fokus Pengguna",
                f"- {focus_block}",
            ]
        )

        if structured_context_block:
            lines.extend(
                [
                    "",
                    "### Parameter Forecast dan Ruang Lingkup",
                    structured_context_block,
                ]
            )

        if chart_marker:
            lines.extend(["", chart_marker])

        lines.extend(
            [
                "",
                "# Analisis Prediktif Cashflow",
                "### Dasar Proyeksi",
                "- Proyeksi menggunakan pendekatan risk-adjusted berdasarkan campuran kelas pembayaran historis, sehingga hasil harus dibaca sebagai skenario manajemen, bukan kepastian arus kas masuk.",
                "- Base case mewakili perilaku penagihan yang paling mungkin terjadi bila pola historis bertahan, sedangkan upside dan downside menunjukkan ruang perbaikan atau penurunan.",
                "",
                "### Snapshot Dashboard Operasional",
                operational_snapshot or "- Snapshot dashboard operasional belum tersedia pada saat laporan dibentuk.",
                "",
                "### Skenario 1-2 Kuartal",
                report_context["scenario_table"],
                "",
                "### Implikasi terhadap Arus Kas Masuk dan Keluar",
                report_context["cash_plan_implications"],
                "",
                "# Rekomendasi Preskriptif",
                "### Prinsip Tindakan",
                "1. Dahulukan invoice bernilai besar dengan skor risiko tinggi dan penyebab yang masih bisa dipulihkan dalam 30 hari.",
                "2. Pisahkan treatment untuk isu anggaran, approval, administrasi, likuiditas, sengketa, dan kewajiban jatuh tempo agar tindakan inflow dan outflow tidak tercampur.",
                "3. Gunakan bukti internal dan jadwal tindak lanjut yang terdokumentasi agar eskalasi ke manajemen klien lebih kuat.",
                "",
                "### Prasyarat Implementasi",
                report_context["implementation_prerequisites"],
                "",
                "### Kesiapan Pelaksanaan",
                report_context["organizational_readiness"],
                "",
                "# Prioritas Tindakan 30 Hari",
                "### Tabel Prioritas",
                report_context["priority_table"],
            ]
        )

        lines.extend(
            [
                "",
                "### Catatan Pelaksanaan",
                "- Tetapkan owner utama per akun prioritas dan review statusnya minimal mingguan.",
                "- Gunakan rapat internal untuk memastikan hambatan administratif dan eskalasi ke klien ditutup dengan tenggat yang jelas.",
                "- Konsolidasikan hasil follow-up ke finance collection, account owner, dan sponsor bisnis agar keputusan rapat langsung dapat dieksekusi.",
            ]
        )

        if dashboard_markers:
            dashboard_insert_index = lines.index("### Implikasi terhadap Arus Kas Masuk dan Keluar") + 2
            dashboard_block = ["", "### Visual Dashboard Snapshot", *dashboard_markers, ""]
            lines[dashboard_insert_index:dashboard_insert_index] = dashboard_block

        if flow_marker:
            lines.extend(["", flow_marker])

        return "\n".join(lines)

    def _run_generation_pass(self, report_context, notes, analysis_context, macro_osint, active_sections, include_visuals, label):
        prompt = self._build_report_prompt(
            report_context,
            notes,
            analysis_context,
            macro_osint,
            active_sections,
            include_visuals,
        )
        user_instruction = self._build_user_instruction(notes, active_sections)
        response = self.ollama.chat(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": user_instruction},
            ],
            options={
                "num_ctx": REPORT_NUM_CTX,
                "num_predict": REPORT_NUM_PREDICT,
                "temperature": REPORT_TEMPERATURE,
                "top_p": REPORT_TOP_P,
                "repeat_penalty": REPORT_REPEAT_PENALTY,
            },
        )
        logger.info(
            "Generation pass %s completed with done_reason=%s, eval_count=%s.",
            label,
            response.get("done_reason"),
            response.get("eval_count"),
        )
        return response["message"]["content"]

    def run(self, notes="", analysis_context="", analysis_payload=None):
        logger.info("Starting cashflow intelligence report generation.")

        osint_context = "\n".join(part for part in (notes, analysis_context) if str(part or "").strip())
        global_osint_future = self.io_pool.submit(Researcher.get_macro_finance_trends, osint_context)
        report_context = self.kb.get_report_context(notes)

        try:
            macro_osint = global_osint_future.result(timeout=45)
        except Exception:
            macro_osint = "OSINT tidak dipakai karena konteks eksternal yang cukup sebanding tidak tersedia."

        fallback_used = False
        generated_sections = []
        for section_pass in self.SECTION_PASSES:
            generated_sections.append(
                self._run_generation_pass(
                    report_context,
                    notes,
                    analysis_context,
                    macro_osint,
                    section_pass["sections"],
                    section_pass["include_visuals"],
                    section_pass["label"],
                ).strip()
            )

        generated_content = "\n\n".join(section for section in generated_sections if section).strip()
        generated_content = self._finalize_report_content(generated_content, report_context, macro_osint, analysis_payload=analysis_payload)
        completeness_result = self._score_report_completeness(generated_content)
        logger.info(
            "Report completeness score %.1f/100 before fallback.",
            completeness_result["score"],
        )
        if not completeness_result["passed"]:
            logger.warning("Generated report failed quality gate. Falling back to deterministic management draft.")
            fallback_used = True
            generated_content = self._build_fallback_report(report_context, notes, analysis_context, macro_osint, analysis_payload=analysis_payload)
            generated_content = self._finalize_report_content(generated_content, report_context, macro_osint, analysis_payload=analysis_payload)
            completeness_result = self._score_report_completeness(generated_content)
            logger.info(
                "Report completeness score %.1f/100 after fallback.",
                completeness_result["score"],
            )

        document = Document()
        DocumentBuilder.create_cover(document, DEFAULT_COLOR)
        DocumentBuilder.add_table_of_contents(document)
        self._embed_dashboard_screenshots(document, analysis_payload, DEFAULT_COLOR)
        DocumentBuilder.process_content(
            document,
            generated_content,
            DEFAULT_COLOR,
        )

        dashboard_screenshots_included = self._has_dashboard_screenshots(analysis_payload)
        run_metadata = {
            "fallback_used": fallback_used,
            "quality_gate_passed": completeness_result["passed"],
            "completeness_score": completeness_result["score"],
            "completeness_missing": completeness_result["missing"],
            "osint_available": bool(
                macro_osint
                and "tidak tersedia" not in macro_osint.lower()
                and "tidak ada data osint" not in macro_osint.lower()
                and "osint tidak dipakai" not in macro_osint.lower()
                and "tidak ada sinyal eksternal yang cukup sebanding" not in macro_osint.lower()
            ),
            "visuals_included": any(marker in generated_content for marker in ("[[CHART:", "[[FLOW:", "[[DASHBOARD:")) or dashboard_screenshots_included,
            "dashboard_screenshots_included": dashboard_screenshots_included,
            "report_length": len(generated_content),
        }

        return document, "Inixindo_Cashflow_Intelligence_Report", run_metadata

    @staticmethod
    def _has_dashboard_screenshots(analysis_payload):
        payload = ReportGenerator._normalize_analysis_payload(analysis_payload)
        screenshots = payload.get("dashboard_screenshots")
        return bool(screenshots and isinstance(screenshots, list) and len(screenshots) > 0)

    @staticmethod
    def _embed_dashboard_screenshots(document, analysis_payload, theme_color):
        payload = ReportGenerator._normalize_analysis_payload(analysis_payload)
        screenshots = payload.get("dashboard_screenshots")
        if not screenshots or not isinstance(screenshots, list):
            return

        valid_screenshots = [
            shot for shot in screenshots
            if isinstance(shot, dict) and shot.get("image_base64")
        ]
        if not valid_screenshots:
            return

        heading = document.add_heading("Dashboard Cashflow Snapshot", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*theme_color)

        intro = document.add_paragraph(
            "Visual berikut diambil langsung dari dashboard operasional pada saat laporan diminta. "
            "Setiap horizon menampilkan kondisi kas, runway, coverage ratio, prediksi saldo, "
            "distribusi delay pembayaran, dan daftar akun overdue utama."
        )
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in intro.runs:
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)

        for shot in valid_screenshots:
            try:
                image_bytes = base64.b64decode(shot["image_base64"])
                image_stream = io.BytesIO(image_bytes)
                horizon_label = shot.get("horizon_label") or shot.get("horizon_key") or "Dashboard"

                image_paragraph = document.add_paragraph()
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_paragraph.add_run().add_picture(image_stream, width=Inches(6.8))

                try:
                    caption = document.add_paragraph(
                        f"Dashboard snapshot — {horizon_label}",
                        style="Caption",
                    )
                except KeyError:
                    caption = document.add_paragraph(
                        f"Dashboard snapshot — {horizon_label}"
                    )
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption.runs:
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(100, 100, 100)

            except Exception as exc:
                logger.warning("Failed to embed dashboard screenshot for %s: %s", shot.get("horizon_key"), exc)
                continue

        document.add_page_break()
