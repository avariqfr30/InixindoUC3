import sys
import threading
import unittest
from pathlib import Path

import pandas as pd


WORKSPACE = Path("/Users/avariqfr30/Documents/InixindoUC3/Payment predictor")
sys.path.insert(0, str(WORKSPACE))


class ReportSanitizationTest(unittest.TestCase):
    def test_report_structure_contract_is_shared_with_config_sequence(self):
        from config import REPORT_SECTION_SEQUENCE
        from report_structure import REPORT_STRUCTURE

        self.assertEqual(REPORT_STRUCTURE.section_sequence, tuple(REPORT_SECTION_SEQUENCE))
        self.assertEqual(
            REPORT_STRUCTURE.section_passes[0]["sections"],
            ("Ringkasan Eksekutif",),
        )
        self.assertIn("### Dampak Bisnis", REPORT_STRUCTURE.required_subheadings)
        self.assertIn("priority_30_day", REPORT_STRUCTURE.required_tables)

    def test_report_generator_private_wrappers_delegate_to_focused_helpers(self):
        from core import ReportGenerator
        from report_finalization import ReportFinalizer
        from report_prompting import ReportPromptBuilder
        from report_quality import ReportQualityScorer

        generator = ReportGenerator(None)

        self.assertIsInstance(generator.prompt_builder, ReportPromptBuilder)
        self.assertIsInstance(generator.quality_scorer, ReportQualityScorer)
        self.assertIsInstance(generator.finalizer, ReportFinalizer)

        section_scope, section_headings = generator._build_section_scope(["Ringkasan Eksekutif"])

        self.assertIn("- Ringkasan Eksekutif", section_scope)
        self.assertIn("# Ringkasan Eksekutif", section_headings)
        self.assertIs(
            generator.prompt_builder.structure,
            generator.quality_scorer.structure,
        )

    def test_final_qa_rejects_raw_source_labels_before_docx(self):
        from report_quality import ReportQualityScorer

        result = ReportQualityScorer().final_qa(
            "# Ringkasan Eksekutif\nEndpoint Internal API.\n\n# Analisis Deskriptif Cashflow\n"
        )

        self.assertFalse(result["passes"])
        self.assertIn("raw_source_label", result["categories"])
        self.assertIn("empty_section", result["categories"])

    def test_report_context_prefetch_reuses_focused_notes_cache(self):
        import cashflow_analysis
        from cashflow_analysis import KnowledgeBase

        class FakeFinancialAnalyzer:
            @staticmethod
            def normalize_evidence_text(value):
                return str(value or "").strip()

            @staticmethod
            def apply_silent_assessment(context, notes, runtime_profile=None):
                return {"controls": f"assessment for {notes}"}

        kb = KnowledgeBase.__new__(KnowledgeBase)
        kb.cache_lock = threading.Lock()
        kb.report_context_cache = {"evidence": "bukti dasar", "review_context": {"ready": True}}
        kb.focused_report_context_cache = {}
        kb.data_version = 3
        kb.active_source_key = "production"
        kb.runtime_profile = {}
        kb.query_calls = 0

        def fake_query(notes, max_results=10):
            kb.query_calls += 1
            return f"bukti fokus untuk {notes}"

        kb.query = fake_query
        original_analyzer = cashflow_analysis.FinancialAnalyzer
        try:
            cashflow_analysis.FinancialAnalyzer = FakeFinancialAnalyzer
            first = kb.prefetch_report_context("akun prioritas")
            second = kb.prefetch_report_context("akun prioritas")
            kb.data_version += 1
            third = kb.prefetch_report_context("akun prioritas")
        finally:
            cashflow_analysis.FinancialAnalyzer = original_analyzer

        self.assertEqual(first["status"], "ready")
        self.assertTrue(second["focusedNotes"])
        self.assertTrue(third["reviewContextReady"])
        self.assertEqual(kb.query_calls, 2)

    def test_cashflow_intelligence_desk_adds_guarded_evidence_to_context(self):
        from core import FinancialAnalyzer

        frame = pd.DataFrame(
            [
                {
                    "periode": "2026-01",
                    "partner": "BUMN A",
                    "layanan": "Pelatihan",
                    "kelas pembayaran": "E",
                    "nilai invoice": "Rp 900.000.000",
                    "catatan keterlambatan": "BAST belum lengkap dan approval direksi klien belum turun.",
                },
                {
                    "periode": "2026-02",
                    "partner": "Korporasi B",
                    "layanan": "Sertifikasi",
                    "kelas pembayaran": "B",
                    "nilai invoice": "Rp 250.000.000",
                    "catatan keterlambatan": "Menunggu jadwal pembayaran normal.",
                },
            ]
        )

        context = FinancialAnalyzer.build_report_context(frame, data_mode="internal")

        self.assertIn("agent_evidence_ledger", context)
        self.assertIn("agent_evidence_brief", context)
        self.assertIn("Invoice Evidence Analyst", context["agent_evidence_brief"])
        self.assertIn("Control Reviewer", context["agent_evidence_brief"])
        self.assertGreaterEqual(len(context["agent_evidence_ledger"]), 4)
        self.assertTrue(
            all("allowed_use" in item and "confidence" in item for item in context["agent_evidence_ledger"])
        )
        self.assertIn("Jangan klaim", context["agent_evidence_brief"])

    def test_report_prompt_includes_hidden_intelligence_desk_brief(self):
        from core import ReportGenerator

        generator = ReportGenerator(None)
        report_context = {
            "financial_summary": "- Total invoice: Rp 1.150.000.000",
            "management_brief": "- Fakta manajemen.",
            "evidence": "- Bukti internal.",
            "readiness_signals": "- Siap dengan caveat.",
            "visual_prompt": "",
            "agent_evidence_brief": (
                "## Invoice Evidence Analyst\n"
                "- Klaim: Rp 900.000.000 tertahan pada BUMN A.\n"
                "## Control Reviewer\n"
                "- Jangan klaim penyebab eksternal tanpa OSINT pembanding."
            ),
        }

        prompt = generator._build_report_prompt(
            report_context=report_context,
            notes="",
            analysis_context="",
            macro_osint="-",
            active_sections=["Ringkasan Eksekutif"],
            include_visuals=False,
        )

        self.assertIn("CASHFLOW INTELLIGENCE DESK EVIDENCE", prompt)
        self.assertIn("Invoice Evidence Analyst", prompt)
        self.assertIn("Jangan klaim penyebab eksternal", prompt)

    def test_readiness_assessment_accepts_runtime_profile_without_config_import_coupling(self):
        from core import FinancialAnalyzer

        context = {
            "base_profile": {
                "data_mode": "production",
                "total_invoices": 35,
                "high_risk_invoices": 4,
                "expected_gap_base": 125_000_000,
                "core_fields_available": 6,
                "core_fields_expected": 6,
                "missing_core_fields": [],
                "top_risk_partners": ["BUMN A"],
            },
            "agent_rejected_claims": ["Jangan klaim penyebab eksternal tanpa OSINT pembanding."],
        }

        result = FinancialAnalyzer.apply_silent_assessment(
            context,
            runtime_profile={
                "app_server": "waitress",
                "report_max_concurrent_jobs": 6,
                "waitress_threads": 10,
            },
        )

        self.assertIn("Infrastructure/deployment readiness: 5/5", result["readiness_signals"])
        self.assertIn("queue 6 job", result["readiness_signals"])
        self.assertIn("thread Waitress 10", result["readiness_signals"])
        self.assertIn("Guardrail kualitas laporan", result["controls"])

    def test_internal_note_trimming_removes_ellipsis_artifacts(self):
        from core import FinancialAnalyzer

        note = (
            "Keputusan... masih menunggu dokumen final dan approval direksi untuk pencairan termin "
            "berikutnya agar invoice bisa diproses."
        )

        cleaned = FinancialAnalyzer._trim_note_for_report(note, max_length=120)

        self.assertNotIn("...", cleaned)
        self.assertIn("Keputusan", cleaned)
        self.assertLessEqual(len(cleaned), 120)

    def test_osint_relevance_gate_rejects_non_comparable_signal(self):
        from core import Researcher

        context = "partner pemerintah, bumn, layanan pelatihan dan sertifikasi"
        comparable_entry = {
            "title": "Vendor pelatihan BUMN hadapi siklus approval pembayaran",
            "snippet": "Pembayaran invoice vendor jasa pelatihan di BUMN tertunda karena approval dan dokumen kontrak.",
            "domain": "kontan.co.id",
        }
        unrelated_entry = {
            "title": "Harga minyak global naik tajam",
            "snippet": "Pasar komoditas bergerak karena sentimen Timur Tengah.",
            "domain": "example.com",
        }

        self.assertTrue(Researcher._is_company_comparable_entry(comparable_entry, context))
        self.assertFalse(Researcher._is_company_comparable_entry(unrelated_entry, context))

    def test_osint_filter_ranks_authoritative_comparable_sources(self):
        from core import Researcher

        context = "partner pemerintah, bumn, layanan pelatihan dan invoice termin"
        entries = [
            {
                "title": "Vendor pelatihan menunggu approval pembayaran invoice",
                "snippet": "Pembayaran invoice vendor jasa pelatihan BUMN tertunda karena approval dokumen.",
                "domain": "blog-random.example",
            },
            {
                "title": "Aturan pengadaan dan termin pembayaran penyedia pemerintah",
                "snippet": "Penyedia pelatihan perlu melengkapi BAST agar pembayaran invoice termin dapat diproses.",
                "domain": "lkpp.go.id",
            },
        ]

        filtered = Researcher._filter_company_comparable_entries(entries, context)

        self.assertEqual(filtered[0]["domain"], "lkpp.go.id")
        self.assertGreater(filtered[0]["relevance_score"], filtered[1]["relevance_score"])

    def test_operational_snapshot_includes_cash_in_and_cash_out(self):
        from core import ReportGenerator

        payload = {
            "selected_period": {"label": "1 Januari 2026 - 31 Maret 2026"},
            "cash_on_hand": 500_000_000,
            "sync_status": {
                "financialData": {"syncStatus": "ready", "sourceAgeMinutes": 12.5, "recordCount": 49},
                "cashOutSource": {"configured": False},
            },
            "horizon_snapshot": {
                "forecasts": {
                    "short_term": {
                        "time_horizon": {"label": "Jangka Pendek (0-30 hari)", "focus": "Likuiditas"},
                        "forecast": {
                            "cash_in": {"total_predicted_cash_in": 1_200_000_000},
                            "cash_out": {"total_cash_out": 300_000_000},
                            "ending_cash": 1_400_000_000,
                        },
                        "cashflow_health": {"runway_months": 4.2, "coverage_ratio": 4.0},
                    }
                }
            },
        }

        snapshot = ReportGenerator._build_operational_snapshot_block(payload)

        self.assertIn("cash masuk", snapshot)
        self.assertIn("cash keluar", snapshot)
        self.assertIn("ending cash", snapshot)

    def test_finalized_report_keeps_visual_dashboard_snapshot_subheading(self):
        from core import ReportGenerator

        generator = ReportGenerator(None)
        raw_text = "\n\n".join(
            [
                "# Ringkasan Eksekutif\nRingkas.",
                "# Analisis Deskriptif Cashflow\nDeskriptif.",
                "# Analisis Diagnostik Cashflow\nDiagnostik.",
                "# Analisis Prediktif Cashflow\n### Dasar Proyeksi\nPrediksi inti.",
                "# Rekomendasi Preskriptif\nRekomendasi.",
                "# Prioritas Tindakan 30 Hari\nPrioritas.",
            ]
        )
        analysis_payload = {
            "horizon_snapshot": {
                "forecasts": {
                    "short_term": {
                        "dashboard_snapshot": {
                            "horizon_key": "short_term",
                            "horizon_label": "Jangka Pendek (0-30 hari)",
                            "horizon_focus": "Likuiditas",
                            "status": "AMAN",
                            "current_cash": 500000000,
                            "runway_months": 2.5,
                            "coverage_ratio": 1.8,
                            "average_delay_days": 35,
                            "balance_projection_30d": [],
                            "coverage_chart": {"bars": []},
                        }
                    }
                }
            }
        }
        finalized = generator._finalize_report_content(
            raw_text=raw_text,
            report_context={"visual_prompt": ""},
            macro_osint="-",
            analysis_payload=analysis_payload,
        )

        self.assertIn("### Cuplikan Dasbor Operasional", finalized)
        self.assertIn("[[DASHBOARD:", finalized)

    def test_finalized_report_injects_executive_headlines_before_caveats(self):
        from core import ReportGenerator

        generator = ReportGenerator(None)
        raw_text = "\n\n".join(
            [
                "# Ringkasan Eksekutif\n### Tingkat Keyakinan dan Caveat\nCatatan teknis.",
                "# Analisis Deskriptif Cashflow\n### Batasan Data dan Asumsi\nAsumsi teknis.",
            ]
        )

        finalized = generator._finalize_report_content(
            raw_text=raw_text,
            report_context={
                "executive_headlines": (
                    "- Rp 2,4 miliar invoice berisiko perlu masuk agenda manajemen minggu ini.\n"
                    "- Satu bottleneck approval paling menentukan ending cash 30 hari."
                ),
                "visual_prompt": "",
            },
            macro_osint="-",
        )

        self.assertIn("### Sorotan Utama untuk Manajemen", finalized)
        self.assertLess(
            finalized.index("### Sorotan Utama untuk Manajemen"),
            finalized.index("### Asumsi Proyeksi dan Catatan Batasan"),
        )
        self.assertLess(
            finalized.index("### Sorotan Utama untuk Manajemen"),
            finalized.index("# Analisis Deskriptif Cashflow"),
        )

    def test_docx_table_generation_uses_compact_formatted_tables(self):
        from docx import Document
        from core import DocumentBuilder

        document = Document()
        markdown_text = (
            "| Prioritas | Fokus | Dampak |\n"
            "| --- | --- | --- |\n"
            "| 1 | Invoice BUMN | Cash in lebih cepat |\n"
        )

        DocumentBuilder.process_content(document, markdown_text)

        self.assertEqual(len(document.tables), 1)
        table = document.tables[0]
        self.assertEqual(table.alignment, 1)
        self.assertTrue(table.rows[0].cells[0].paragraphs[0].runs[0].bold)
        self.assertLessEqual(table.rows[1].cells[0].paragraphs[0].paragraph_format.space_after.pt, 2)

    def test_docx_body_text_is_left_aligned_and_ordered_lists_restart(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from core import DocumentBuilder

        document = Document()
        DocumentBuilder.parse_html_to_docx(
            document,
            (
                "<p>Paragraf ringkas untuk laporan.</p>"
                "<ol><li>Prioritas pertama</li><li>Prioritas kedua</li></ol>"
                "<ol><li>Restart prioritas baru</li><li>Lanjutannya</li></ol>"
            ),
            (204, 0, 0),
        )

        paragraph_texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

        self.assertEqual(document.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertIn("Prioritas pertama", paragraph_texts)
        self.assertIn("Prioritas kedua", paragraph_texts)
        self.assertIn("Restart prioritas baru", paragraph_texts)
        self.assertIn("Lanjutannya", paragraph_texts)
        self.assertFalse(any(text.startswith(("1.", "2.")) for text in paragraph_texts))

        list_paragraphs = [paragraph for paragraph in document.paragraphs if "Prioritas" in paragraph.text or "Restart" in paragraph.text or "Lanjutannya" in paragraph.text]
        self.assertEqual(len(list_paragraphs), 4)
        self.assertTrue(all(paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None for paragraph in list_paragraphs))
        first_list_num_id = list_paragraphs[0]._p.pPr.numPr.numId.val
        second_list_num_id = list_paragraphs[2]._p.pPr.numPr.numId.val
        self.assertNotEqual(first_list_num_id, second_list_num_id)
        self.assertEqual(list_paragraphs[0]._p.pPr.numPr.ilvl.val, 0)
        self.assertEqual(list_paragraphs[0].paragraph_format.left_indent.inches, list_paragraphs[2].paragraph_format.left_indent.inches)

    def test_docx_nested_list_indentation_and_bullet_numbering_are_structured(self):
        from docx import Document
        from core import DocumentBuilder

        document = Document()
        DocumentBuilder.parse_html_to_docx(
            document,
            (
                "<h1>Bab Satu</h1>"
                "<p>Kalimat pembuka <strong>dengan prioritas</strong>:</p>"
                "<ol><li>Langkah utama<ul><li>Catatan detail</li></ul></li><li>Langkah lanjutan</li></ol>"
                "<h1>Bab Dua</h1>"
                "<ol><li>Nomor kembali satu</li></ol>"
            ),
            (204, 0, 0),
        )

        list_paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.strip() in {"Langkah utama", "Catatan detail", "Langkah lanjutan", "Nomor kembali satu"}
        ]

        self.assertEqual([paragraph.text for paragraph in list_paragraphs], ["Langkah utama", "Catatan detail", "Langkah lanjutan", "Nomor kembali satu"])
        self.assertTrue(all(paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None for paragraph in list_paragraphs))
        self.assertEqual(list_paragraphs[0]._p.pPr.numPr.ilvl.val, 0)
        self.assertEqual(list_paragraphs[1]._p.pPr.numPr.ilvl.val, 1)
        self.assertGreater(list_paragraphs[1].paragraph_format.left_indent.inches, list_paragraphs[0].paragraph_format.left_indent.inches)
        self.assertNotEqual(list_paragraphs[0]._p.pPr.numPr.numId.val, list_paragraphs[3]._p.pPr.numPr.numId.val)


if __name__ == "__main__":
    unittest.main()
