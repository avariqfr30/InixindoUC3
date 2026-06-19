import sys
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

PROJECT_DIR = Path(__file__).resolve().parents[1]
if str(PROJECT_DIR) not in sys.path:
    sys.path.insert(0, str(PROJECT_DIR))

from docx_rendering import DocumentBuilder
from report_evidence import PaymentEvidenceBuilder
from report_finalization import ReportFinalizer


class ReaderFacingDocumentContractTests(unittest.TestCase):
    FORBIDDEN_READER_TERMS = [
        "Invoice Evidence Analyst",
        "Control Reviewer",
        "Executive Editor",
        "evidence ledger",
        "endpoint",
        "runtime",
        "queue",
        "thread",
        "demo mode",
        "source-of-truth",
        "Internal API",
        "APIDog",
    ]

    FORBIDDEN_READER_PATTERNS = [
        r"\bagent\b",
        r"\bdesk\b",
        r"\bworkflow\b",
    ]

    UNNECESSARY_ENGLISH_LABELS = [
        "BLUF",
        "Key Findings",
        "Recommendation",
        "Dashboard",
        "Insight",
        "Executive Summary",
        "Caveat",
        "Driver",
        "Timing",
        "Headline",
        "Forecast",
        "Owner",
    ]

    def test_reader_safe_text_removes_source_mode_terms(self):
        raw = "API internal perusahaan memakai endpoint source-of-truth dengan Waitress queue thread dan sync status record aktif."
        clean = DocumentBuilder.reader_safe_text(raw)
        for token in ["API internal", "endpoint", "source-of-truth", "Waitress", "queue", "thread", "sync status", "record aktif"]:
            self.assertNotIn(token, clean)
        self.assertIn("data operasional yang tersedia", clean)

    def test_reader_safe_text_synthesizes_raw_ui_and_api_helpers(self):
        raw = (
            "Nama Perusahaan Klien: ReferenceAccount mencatat source=/api/Resource/dataset "
            "dataset_code=ConsultantProjectExpertHistory. Dirangkum dari sumber APIDog: "
            "Problem, Opportunity, Directive untuk Pain Points dan endpoint sync status."
        )

        clean = DocumentBuilder.reader_safe_text(raw)

        forbidden = [
            "Nama Perusahaan Klien",
            "ReferenceAccount",
            "source=",
            "/api/Resource/dataset",
            "dataset_code",
            "ConsultantProjectExpertHistory",
            "Dirangkum dari sumber",
            "APIDog",
            "Problem, Opportunity, Directive",
            "Pain Points",
            "endpoint",
            "sync status",
        ]
        for token in forbidden:
            self.assertNotIn(token, clean)
        self.assertIn("catatan klien", clean)
        self.assertIn("riwayat pengalaman konsultan", clean)
        self.assertIn("kebutuhan prioritas yang perlu dipertegas", clean)

    def test_toc_has_real_field_and_updates_on_open(self):
        doc = Document()
        DocumentBuilder.add_table_of_contents(doc)

        document_xml = doc._element.xml
        settings_xml = doc.settings.element.xml
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        self.assertIn("Daftar Isi", text)
        self.assertNotIn("Update Field", text)
        self.assertIn('TOC \\o "1-3" \\h \\z \\u', document_xml)
        self.assertIn('w:fldCharType="begin"', document_xml)
        self.assertIn('w:fldCharType="separate"', document_xml)
        self.assertIn('w:fldCharType="end"', document_xml)
        update_fields = doc.settings.element.find(qn("w:updateFields"))
        self.assertIsNotNone(update_fields)
        self.assertEqual(update_fields.get(qn("w:val")), "true")
        self.assertIn('w:updateFields w:val="true"', settings_xml)

    def test_executive_summary_finalization_is_decision_first_and_source_neutral(self):
        finalizer = ReportFinalizer()
        raw = "# Ringkasan Eksekutif\n### Tingkat Keyakinan dan Caveat\nCatatan teknis dari API internal perusahaan dan endpoint."
        context = {
            "executive_headlines": "- Risiko kas terkonsentrasi pada keterlambatan pembayaran akun utama.",
            "executive_facts": "- Saldo kas akhir perlu dijaga.",
            "confidence_summary": "Keyakinan sedang berdasarkan bukti operasional.",
        }
        result = finalizer.finalize(raw, context, "", analysis_payload={})
        self.assertIn("# Ringkasan Eksekutif", result)
        self.assertIn("### Kesimpulan Utama", result)
        self.assertIn("### Keputusan yang Diminta", result)
        for token in ["API internal", "endpoint", "source-of-truth", "Waitress", "queue", "thread", "sync status"]:
            self.assertNotIn(token, result)

    def test_executive_summary_uses_management_interpretation_model(self):
        finalizer = ReportFinalizer()
        raw = "# Ringkasan Eksekutif\nDraft."
        context = {
            "executive_headlines": "- Keputusan kas: eskalasi penagihan akun utama minggu ini.",
            "executive_facts": "- Cash in tertahan pada invoice prioritas.\n- Cash out tetap berjalan.",
            "management_interpretation": {
                "signal": "Cash in tertahan sementara cash out tetap berjalan.",
                "meaning": "Risiko utama adalah timing-control, bukan sekadar nominal invoice.",
                "decision": "Manajemen perlu memilih akun yang dieskalasi minggu ini.",
                "action": "Eskalasi senior dan konfirmasi komitmen pembayaran tertulis.",
                "confidence": "Sedang - data operasional cukup, status komitmen perlu validasi.",
            },
            "visual_prompt": "",
        }

        result = finalizer.finalize(raw, context, "", analysis_payload={})
        summary_body = result.split("# Analisis Deskriptif Cashflow", 1)[0]

        self.assertIn("### Interpretasi Manajemen", summary_body)
        self.assertIn("| Sinyal | Makna | Keputusan | Aksi | Keyakinan |", summary_body)
        self.assertIn("kontrol waktu", summary_body)
        self.assertIn("Eskalasi senior", summary_body)
        self.assertLess(summary_body.index("### Interpretasi Manajemen"), summary_body.index("### Agenda Follow-up Meeting"))

    def test_finalization_strips_hidden_workflow_roles_and_source_mechanics(self):
        finalizer = ReportFinalizer()
        raw = "\n\n".join(
            [
                "# Ringkasan Eksekutif\nAgent workflow dari cashflow intelligence desk memakai Invoice Evidence Analyst, Control Reviewer, Executive Editor, evidence ledger, endpoint, runtime, queue, thread, demo mode, source-of-truth, Internal API, dan APIDog.",
                "# Analisis Deskriptif Cashflow\nDesk reviewer menyebut source-of-truth dari APIDog endpoint.",
                "# Analisis Diagnostik Cashflow\nWorkflow evidence ledger internal.",
            ]
        )
        context = {
            "executive_headlines": "- Risiko kas membutuhkan keputusan penagihan senior minggu ini.",
            "executive_facts": "- Driver utama adalah konsentrasi invoice terlambat.",
            "confidence_summary": "Caveat berasal dari endpoint Internal API/APIDog dan demo mode.",
            "visual_prompt": "",
        }

        result = finalizer.finalize(raw, context, "", analysis_payload={})

        for token in self.FORBIDDEN_READER_TERMS:
            self.assertNotIn(token, result)
        for pattern in self.FORBIDDEN_READER_PATTERNS:
            self.assertNotRegex(result, pattern)

    def test_executive_summary_uses_finance_bluf_order_before_caveats(self):
        finalizer = ReportFinalizer()
        raw = "# Ringkasan Eksekutif\n### Tingkat Keyakinan dan Caveat\nEndpoint Internal API perlu dicek dulu."
        context = {
            "executive_headlines": "- Keputusan kas: eskalasi penagihan akun utama minggu ini untuk menjaga ending cash.",
            "executive_facts": "- Driver utama: invoice bernilai besar terlambat.\n- Driver kedua: cash out tetap berjalan.",
            "confidence_summary": "Asumsi forecast memakai data operasional yang tersedia; caveat dibaca setelah tindakan prioritas.",
            "visual_prompt": "",
        }

        result = finalizer.finalize(raw, context, "", analysis_payload={})
        summary_body = result.split("# Analisis Deskriptif Cashflow", 1)[0]

        self.assertRegex(summary_body, r"^# Ringkasan Eksekutif\n### Kesimpulan Utama")
        self.assertLess(summary_body.index("### Kesimpulan Utama"), summary_body.index("### Keputusan yang Diminta"))
        self.assertLess(summary_body.index("### Keputusan yang Diminta"), summary_body.index("### Interpretasi Manajemen"))
        self.assertLess(summary_body.index("### Interpretasi Manajemen"), summary_body.index("### Alasan Utama"))
        self.assertLess(summary_body.index("### Alasan Utama"), summary_body.index("### Bukti Pendukung"))
        self.assertLess(summary_body.index("### Bukti Pendukung"), summary_body.index("### Catatan Keyakinan dan Batasan"))
        self.assertLess(summary_body.index("### Catatan Keyakinan dan Batasan"), summary_body.index("### Agenda Follow-up Meeting"))
        self.assertNotIn("Endpoint", summary_body)
        self.assertNotIn("Internal API", summary_body)

    def test_executive_summary_actions_are_context_specific_not_static(self):
        finalizer = ReportFinalizer()
        raw = "# Ringkasan Eksekutif\nDraft."
        context = {
            "executive_headlines": "- Rp 2,4 miliar tagihan BUMN A perlu dieskalasi minggu ini.",
            "executive_facts": "- Eksposur terbesar ada pada BUMN A.\n- Layanan paling berisiko adalah Pelatihan SPBE.",
            "base_profile": {
                "top_risk_partners": ["BUMN A"],
                "top_risk_services": ["Pelatihan SPBE"],
                "expected_gap_base": 2400000000,
            },
            "confidence_summary": "Keyakinan sedang berdasarkan data operasional.",
            "visual_prompt": "",
        }

        result = finalizer.finalize(raw, context, "", analysis_payload={})
        summary_body = result.split("# Analisis Deskriptif Cashflow", 1)[0]

        self.assertIn("BUMN A", summary_body)
        self.assertIn("Pelatihan SPBE", summary_body)
        self.assertNotIn("Akun mana yang harus diprioritaskan untuk penagihan senior.", summary_body)

    def test_executive_summary_does_not_repeat_factors_as_management_highlights(self):
        finalizer = ReportFinalizer()
        raw = "# Ringkasan Eksekutif\nDraft lama."
        context = {
            "executive_headlines": "- Keputusan kas: eskalasi penagihan akun utama minggu ini.",
            "executive_facts": "- Fakta utama: tagihan besar terlambat.\n- Fakta kedua: arus kas keluar tetap berjalan.",
            "confidence_summary": "Keyakinan sedang berdasarkan data operasional.",
            "visual_prompt": "",
        }

        result = finalizer.finalize(raw, context, "", analysis_payload={})
        summary_body = result.split("# Analisis Deskriptif Cashflow", 1)[0]
        factors = summary_body.split("### Alasan Utama", 1)[1].split("### Bukti Pendukung", 1)[0].strip()
        evidence = summary_body.split("### Bukti Pendukung", 1)[1].split("### Catatan Keyakinan dan Batasan", 1)[0].strip()

        self.assertNotEqual(factors, evidence)
        self.assertIn("Fakta utama", factors)
        self.assertIn("Keputusan kas", summary_body.split("### Keputusan yang Diminta", 1)[0])

    def test_executive_summary_synthesizes_finished_report_sections(self):
        finalizer = ReportFinalizer()
        raw = "\n\n".join(
            [
                "# Ringkasan Eksekutif\nDraft awal sebelum laporan selesai.",
                "# Analisis Deskriptif Cashflow\nSaldo kas akhir stabil setelah prioritas penagihan dijalankan.",
                "# Analisis Diagnostik Cashflow\nRisiko utama berasal dari tagihan terlambat akun besar.",
                "# Rekomendasi Aksi\nEskalasi penagihan senior perlu dimulai minggu ini.",
            ]
        )
        context = {
            "executive_headlines": "- Keputusan kas: eskalasi penagihan akun utama minggu ini.",
            "executive_facts": "- Fakta utama: tagihan besar terlambat.",
            "confidence_summary": "Keyakinan sedang berdasarkan data operasional.",
            "visual_prompt": "",
        }

        result = finalizer.finalize(raw, context, "", analysis_payload={})
        summary_body = result.split("# Analisis Deskriptif Cashflow", 1)[0]

        self.assertIn("### Ringkasan Isi Laporan", summary_body)
        self.assertIn("Saldo kas akhir stabil", summary_body)
        self.assertIn("tagihan terlambat akun besar", summary_body)
        self.assertLess(summary_body.index("### Alasan Utama"), summary_body.index("### Ringkasan Isi Laporan"))
        self.assertLess(summary_body.index("### Ringkasan Isi Laporan"), summary_body.index("### Bukti Pendukung"))

    def test_finalization_inserts_source_safe_evidence_cards(self):
        finalizer = ReportFinalizer()
        raw = "\n\n".join(
            [
                "# Ringkasan Eksekutif\nDraft.",
                "# Analisis Deskriptif Cashflow\nSaldo kas akhir perlu dijaga.",
                "# Analisis Diagnostik Cashflow\nTagihan besar terlambat.",
            ]
        )
        context = {
            "executive_headlines": "- Keputusan kas: eskalasi penagihan akun utama.",
            "executive_facts": "- Fakta utama: tagihan besar terlambat.",
            "confidence_summary": "Keyakinan sedang.",
            "visual_prompt": "",
            "agent_evidence_ledger": [
                {
                    "claim": "Rp 900 juta tertahan pada akun prioritas.",
                    "allowed_use": "Dipakai untuk prioritas penagihan.",
                    "confidence": "Tinggi",
                    "source": "Invoice Evidence Analyst",
                }
            ],
        }

        result = finalizer.finalize(raw, context, "", analysis_payload={})

        self.assertIn("### Bukti yang Dipakai", result)
        self.assertIn("Rp 900 juta", result)
        for token in ["Invoice Evidence Analyst", "agent", "workflow", "endpoint", "Internal API"]:
            self.assertNotIn(token, result)

    def test_finalization_omits_placeholder_osint_marker(self):
        finalizer = ReportFinalizer()
        raw = "\n\n".join(
            [
                "# Ringkasan Eksekutif\nDraft.",
                "# Analisis Diagnostik Cashflow\n### Risiko dan Kontrol\nKontrol dasar.",
            ]
        )

        result = finalizer.finalize(
            raw,
            {"executive_headlines": "- Keputusan kas dibutuhkan.", "visual_prompt": ""},
            "-",
            analysis_payload={},
        )

        self.assertIn("Konteks OSINT Pendukung", result)
        self.assertIn("belum cukup sebanding", result)
        self.assertNotIn("### Konteks OSINT Pendukung\n-", result)

    def test_payment_evidence_uses_implicit_confidence_without_exposing_confidence_label(self):
        text = PaymentEvidenceBuilder.to_markdown(
            [
                {
                    "claim": "Invoice terlambat bernilai Rp 900 juta perlu diprioritaskan.",
                    "allowed_use": "Dipakai untuk agenda penagihan 30 hari.",
                    "confidence": "Tinggi",
                    "source": "Invoice Evidence Analyst",
                }
            ]
        )

        self.assertIn("Dasar bukti kuat", text)
        self.assertNotIn("Keyakinan:", text)
        self.assertNotIn("confidence", text.lower())

    def test_payment_evidence_builder_keeps_internal_roles_hidden(self):
        text = PaymentEvidenceBuilder.to_markdown(
            [
                {
                    "claim": "Invoice Evidence Analyst menemukan tagihan Rp 900 juta.",
                    "allowed_use": "Control Reviewer menyetujui penggunaan untuk prioritas.",
                    "confidence": "Tinggi",
                }
            ]
        )

        self.assertIn("### Bukti yang Dipakai", text)
        self.assertIn("Rp 900 juta", text)
        self.assertNotIn("Invoice Evidence Analyst", text)
        self.assertNotIn("Control Reviewer", text)

    def test_finalized_report_translates_unnecessary_english_reader_labels(self):
        finalizer = ReportFinalizer()
        raw = "\n\n".join(
            [
                "# Executive Summary",
                "### BLUF",
                "Dashboard insight: cashflow pressure needs management action.",
                "### Key Findings",
                "- Driver utama keterlambatan adalah approval.",
                "- Timing pembayaran perlu dipantau.",
                "### Recommendation",
                "- Escalate akun prioritas minggu ini.",
                "### Caveat",
                "Internal API endpoint masih perlu dicek.",
            ]
        )
        context = {
            "executive_headlines": "- BLUF: keputusan kas dibutuhkan minggu ini.",
            "executive_facts": "- Key Findings: Driver utama adalah invoice besar terlambat.",
            "confidence_summary": "Caveat: timing forecast bergantung pada data operasional.",
            "visual_prompt": "",
        }

        result = finalizer.finalize(raw, context, "", analysis_payload={})

        self.assertIn("# Ringkasan Eksekutif", result)
        self.assertIn("### Kesimpulan Utama", result)
        self.assertIn("### Alasan Utama", result)
        self.assertIn("### Agenda Follow-up Meeting", result)
        self.assertIn("### Catatan Keyakinan dan Batasan", result)
        self.assertIn("Inti Keputusan", result)
        self.assertIn("Temuan Utama", result)
        for label in self.UNNECESSARY_ENGLISH_LABELS:
            self.assertNotRegex(result, rf"\b{label}\b")


if __name__ == "__main__":
    unittest.main()
